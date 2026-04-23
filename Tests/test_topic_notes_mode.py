"""
Regression checks for general law revision-notes routing and DOCX rendering.
"""

from pathlib import Path
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from model_applicable_service import (
    QUERY_CHUNK_CONFIG,
    _cleanup_topic_notes_temp_artifacts,
    detect_all_query_types,
    detect_topic_notes_request,
    detect_topic_notes_session_signal,
    register_topic_notes_cleanup_paths,
    send_message_with_docs,
)
from legal_doc_tools.generate_review_report_docx import build_docx


prompt = (
    "Make revision notes on contract law covering all likely big and small topics, "
    "with case law development and journal arguments, in the same style as the sample docx."
)
notes_mode = detect_topic_notes_request(prompt)
print("notes_mode:", notes_mode)
assert notes_mode["is_topic_notes"] is True
assert notes_mode["scope"] == "full_subject"
assert notes_mode["sample_style_requested"] is True
assert "topic_notes" in detect_all_query_types(prompt)
assert QUERY_CHUNK_CONFIG["topic_notes"] == 24

history = [
    {"role": "user", "text": prompt},
    {"role": "assistant", "text": "Contract Law Principles Notes\n..."},
]
accept_signal = detect_topic_notes_session_signal(
    "done, accept the delivery result",
    history,
    active=True,
)
assert accept_signal["is_acceptance"] is True
assert accept_signal["is_amendment"] is False

amend_signal = detect_topic_notes_session_signal(
    "amend topic 2 and add more cases",
    history,
    active=True,
)
assert amend_signal["is_acceptance"] is False
assert amend_signal["is_amendment"] is True

runtime_dir = Path(".codex_runtime")
runtime_dir.mkdir(exist_ok=True)
cleanup_target = runtime_dir / "topic_notes_sample_cleanup_test.txt"
cleanup_target.write_text("temporary notes sample", encoding="utf-8")
register_topic_notes_cleanup_paths("notes_cleanup_test", [cleanup_target])
removed_count = _cleanup_topic_notes_temp_artifacts("notes_cleanup_test")
assert removed_count >= 1
assert not cleanup_target.exists()

acceptance_target = runtime_dir / "topic_notes_sample_acceptance_test.txt"
acceptance_target.write_text("temporary notes sample", encoding="utf-8")
acceptance_project_id = "notes_acceptance_cleanup_test"
register_topic_notes_cleanup_paths(acceptance_project_id, [acceptance_target])
(accept_text, _accept_meta), _accept_rag = send_message_with_docs(
    api_key="",
    message="delivery accepted",
    documents=[],
    project_id=acceptance_project_id,
    history=history,
    stream=False,
)
assert "temporary runtime notes artifacts have been cleared" in accept_text
assert not acceptance_target.exists()

note_text = """
Contract Law Principles Notes
Refined revision notes on arguable principles, key cases, rule analysis and doctrinal development
Based on and expanded from indexed and uploaded materials.

How to use these notes
- Start by identifying whether the issue is formation, vitiation, construction, discharge, or remedy.
- Use the Add / check section at the end of each topic as your issue-spotting checklist.

Topic map

Quick development timeline

Topic 1: Offer and acceptance
REVISION FRAME

Principle: The real issue is often not whether agreement exists in the abstract, but whether the communication relied on was sufficiently final and properly communicated.

Argument 1
Case law support
Carlill v Carbolic Smoke Ball Co [1893] 1 QB 256
Facts: The company promised to pay users who followed the advertisement and still caught influenza.
Held: The advertisement was capable of amounting to a unilateral offer.
Reasoning: Clear language, seriousness, and performance-based acceptance allowed the court to treat the advert as an offer rather than mere sales puff.
Use in an answer: Use this where the dispute is whether language was sufficiently promissory and acceptance occurred by performance.

Argument 2
Academic / journal support
Mindy Chen-Wishart, Contract Law
Argument: Academic commentary often emphasises that modern formation disputes are really about objective commitment and commercial context rather than formal verbal formulae.
Use in an answer: Use this to support a contextual reading of negotiations and alleged offers.

Conclusion / comparison / development
The development runs from formal offer-and-acceptance analysis toward a more contextual approach, but the court still insists on objective commitment and clear communication.

Add / check in a problem question
- Identify the exact words said or written.
- Ask whether the statement was final or merely an invitation to treat.
""".strip()

with tempfile.TemporaryDirectory() as tmpdir:
    out_path = Path(tmpdir) / "contract_notes.docx"
    build_docx(note_text, out_path)
    doc = Document(out_path)

    styled_paras = [(p.text.strip(), p.style.name) for p in doc.paragraphs if p.text.strip()]
    print("styled_paras:", styled_paras[:12])

    assert ("Contract Law Principles Notes", "Title") in styled_paras
    assert (
        "Refined revision notes on arguable principles, key cases, rule analysis and doctrinal development",
        "Subtitle",
    ) in styled_paras
    assert ("How to use these notes", "Heading 1") in styled_paras
    assert ("Topic 1: Offer and acceptance", "Heading 1") in styled_paras
    assert ("REVISION FRAME", "Topic Kicker") in styled_paras
    assert ("Argument 1", "Mini Heading") in styled_paras
    assert ("Conclusion / comparison / development", "Heading 2") in styled_paras
    assert ("Add / check in a problem question", "Heading 2") in styled_paras
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.sections[0].different_first_page_header_footer is True
    assert len(doc.tables) >= 7
    assert doc.styles["Normal"].font.name == "Calibri"
    assert doc.styles["Title"].font.name == "Calibri"
    assert doc.styles["Subtitle"].font.name == "Calibri"
    assert doc.styles["Heading 1"].font.name == "Calibri"
    assert doc.styles["Heading 2"].font.name == "Calibri"
    assert doc.styles["Mini Heading"].font.name == "Calibri"
    assert doc.styles["Support Heading"].font.name == "Calibri"
    assert doc.styles["Topic Kicker"].font.size.pt >= 10
    assert doc.styles["Normal"].font.size.pt >= 10
    assert doc.styles["Subtitle"].font.size.pt >= 10

    summary_table = doc.tables[1]
    assert summary_table.cell(0, 0).text.strip() == "What this document does"
    assert summary_table.cell(1, 0).text.strip() == "Focus"
    assert "contract law" in summary_table.cell(1, 1).text.strip().lower()

    topic_map_table = doc.tables[4]
    assert topic_map_table.cell(0, 0).text.strip() == "Topic"
    assert topic_map_table.cell(1, 0).text.strip() == "Topic 1"
    assert "Offer and acceptance" in topic_map_table.cell(1, 1).text

    case_para = next(p for p in doc.paragraphs if p.text.strip() == "Carlill v Carbolic Smoke Ball Co [1893] 1 QB 256")
    assert any(run.bold for run in case_para.runs if run.text.strip()), "Case law heading line should be bold"

    principle_para = next(p for p in doc.paragraphs if p.text.startswith("Principle:"))
    assert principle_para.runs[0].bold is True
    assert principle_para.style.name == "Normal"

print("Topic notes routing and DOCX rendering checks passed.")
