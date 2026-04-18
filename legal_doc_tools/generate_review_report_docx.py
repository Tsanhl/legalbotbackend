#!/usr/bin/env python3
"""
Generate a DOCX review report from plain text or markdown-like input.

This script is intentionally lightweight so agents can always emit a report
artifact (`.docx`) as a mandatory deliverable.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

DESKTOP_ROOT = (Path.home() / "Desktop").resolve()

NOTE_HEADING_2_LINES = {
    "Topic map",
    "Quick development timeline",
    "Conclusion / comparison / development",
    "Add / check in a problem question",
}
NOTE_SUPPORT_HEADINGS = {
    "Case law support",
    "Academic / journal support",
    "Academic / journal debate",
    "Journal / academic support",
    "Statutory support",
    "Legislation support",
}
NOTE_LABEL_PREFIXES = (
    "Principle:",
    "Facts:",
    "Held:",
    "Reasoning:",
    "Use in an answer:",
    "Argument:",
    "Statute:",
    "Legislation:",
    "Journal argument:",
    "Counterargument:",
)

NOTE_DARK_BLUE = "1F4E79"
NOTE_TEXT_BLUE = "234A7A"
NOTE_LIGHT_BLUE_FILL = "EDF4FB"
NOTE_LIGHT_BLUE_BORDER = "C9DDF2"
NOTE_LIGHT_BLUE_TEXT = "153E75"
NOTE_YELLOW_FILL = "FFF4DB"
NOTE_YELLOW_BORDER = "E5C07B"
NOTE_YELLOW_TEXT = "7A4B00"
NOTE_GRID = "D0D7DE"
NOTE_BODY_FILL = "F8FAFD"
NOTE_TITLE_RULE = "4E81C3"


def _is_numbered(line: str) -> bool:
    i = 0
    while i < len(line) and line[i].isdigit():
        i += 1
    return i > 0 and i + 1 < len(line) and line[i] == "." and line[i + 1] == " "


def _strip_markdown_inline(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def _strip_numeric_prefix(text: str) -> str:
    return re.sub(r"^\s*\d+\.\s+", "", text)


def _normalize_ledger_item(text: str) -> str:
    t = _strip_markdown_inline(text).strip()
    t = _strip_numeric_prefix(t)
    t = t.strip()
    if re.match(r"^\[\d+\]\.\s+", t):
        return t
    if re.match(r"^\[\d+\]\s+", t):
        return t
    if re.match(r"^Footnote\s+\d+\b", t, flags=re.IGNORECASE):
        return t
    return t


def _looks_like_topic_notes_document(input_text: str) -> bool:
    indicators = 0
    patterns = [
        r"(?im)^\s*how to use these notes\s*$",
        r"(?im)^\s*topic\s+\d+\s*:",
        r"(?im)^\s*revision frame\s*$",
        r"(?im)^\s*add\s*/\s*check in a problem question\s*$",
        r"(?im)^\s*case law support\s*$",
    ]
    for pattern in patterns:
        if re.search(pattern, input_text or ""):
            indicators += 1
    return indicators >= 2


def _set_ascii_font(style, name: str, size: float, *, bold=None, italic=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic


def _ensure_paragraph_style(doc: Document, name: str, *, base: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float = 10.5,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def _set_paragraph_border(paragraph, *, side: str, color: str, size: int = 8, space: int = 2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = p_bdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        p_bdr.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, **edges: Dict[str, Any]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        if not edge_data:
            continue
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def _style_all_cell_borders(table, *, color: str = NOTE_GRID, size: int = 8) -> None:
    border = {"val": "single", "sz": size, "color": color}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                left=border,
                top=border,
                right=border,
                bottom=border,
            )


def _add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    _set_run_font(run, size=10, color="9BA7B4")


def _normalise_note_line(raw: str) -> str:
    return _strip_markdown_inline((raw or "").strip())


def _strip_bullet(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("- "):
        return stripped[2:].strip()
    return stripped


def _shorten(text: str, max_len: int = 120) -> str:
    compact = re.sub(r"\s+", " ", (text or "").strip())
    if len(compact) <= max_len:
        return compact
    cut = compact[: max_len - 1].rstrip()
    if " " in cut:
        cut = cut.rsplit(" ", 1)[0]
    return cut + "…"


def _compact_authority_name(text: str) -> str:
    compact = _strip_bullet(_strip_markdown_inline(text))
    compact = re.sub(r"\s*\[[^\]]+\].*$", "", compact).strip()
    compact = re.sub(r"\s*\([^)]*\d{4}[^)]*\).*$", "", compact).strip()
    compact = re.sub(r"\s+", " ", compact)
    return _shorten(compact, 36)


def _looks_like_case_authority_line(text: str) -> bool:
    compact = _strip_bullet(text)
    low = compact.lower()
    return bool(
        " v " in low
        or low.startswith("re ")
        or low.startswith("in re ")
        or re.search(r"\bcase\s+[ctf]-?\d+/\d+\b", low)
        or re.search(r"\[[12]\d{3}\]", compact)
    )


def _looks_like_statute_reference(text: str) -> bool:
    compact = _strip_bullet(text)
    return bool(
        re.search(r"\b[A-Z][A-Za-z'’\- ]+ Act \d{4}\b", compact)
        or re.search(r"\b[A-Z][A-Za-z'’\- ]+ Regulations? \d{4}\b", compact)
        or re.search(r"\bArticle\s+\d+\b", compact)
        or re.search(r"\bsection\s+\d+\b", compact, flags=re.IGNORECASE)
        or re.search(r"\bs\s*\d+[A-Za-z]?\b", compact)
    )


def _extract_topic_title(topic_heading: str) -> str:
    match = re.match(r"^(Topic\s+\d+)\s*:\s*(.+)$", topic_heading)
    if not match:
        return topic_heading.strip()
    return match.group(2).strip()


def _extract_topic_label(topic_heading: str) -> str:
    match = re.match(r"^(Topic\s+\d+)\s*:", topic_heading)
    return match.group(1) if match else topic_heading.strip()


def _split_timeline_entry(text: str, index: int) -> Tuple[str, str]:
    compact = _strip_bullet(text)
    explicit = re.match(
        r"^([A-Z][A-Za-z0-9 /–-]{1,40}?)(?:\s*[:\-–]\s+)(.+)$",
        compact,
    )
    if explicit and (
        re.search(r"\d", explicit.group(1))
        or any(
            token in explicit.group(1).lower()
            for token in ("century", "modern", "early", "late", "current", "foundational")
        )
    ):
        return explicit.group(1).strip(), explicit.group(2).strip()

    if compact.lower().startswith(("early ", "late ", "modern ", "nineteenth", "twentieth", "contemporary")):
        tokens = compact.split()
        if len(tokens) >= 2:
            return " ".join(tokens[:2]), " ".join(tokens[2:]).strip() or compact

    fallback_periods = [
        "Foundational stage",
        "Doctrinal expansion",
        "Refinement",
        "Modern position",
        "Exam deployment",
    ]
    period = fallback_periods[index] if index < len(fallback_periods) else f"Development {index + 1}"
    return period, compact


def _coalesce_unique(items: List[str], *, limit: int = 4) -> List[str]:
    out: List[str] = []
    seen = set()
    for item in items:
        compact = re.sub(r"\s+", " ", (item or "").strip())
        if not compact:
            continue
        key = compact.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(compact)
        if len(out) >= limit:
            break
    return out


def _parse_topic_notes_input(input_text: str) -> Dict[str, Any]:
    frontmatter: List[str] = []
    how_to_use: List[str] = []
    topic_map: List[str] = []
    timeline: List[str] = []
    topics: List[Dict[str, Any]] = []

    current_section = "frontmatter"
    current_topic: Optional[Dict[str, Any]] = None

    for raw in (input_text or "").splitlines():
        stripped = _normalise_note_line(raw)
        if not stripped:
            continue
        if stripped == "How to use these notes":
            current_section = "how_to_use"
            current_topic = None
            continue
        if stripped == "Topic map":
            current_section = "topic_map"
            current_topic = None
            continue
        if stripped == "Quick development timeline":
            current_section = "timeline"
            current_topic = None
            continue
        if re.match(r"^Topic\s+\d+\s*:", stripped):
            current_section = "topic"
            current_topic = {"title": stripped, "lines": []}
            topics.append(current_topic)
            continue

        if current_section == "frontmatter":
            frontmatter.append(stripped)
        elif current_section == "how_to_use":
            how_to_use.append(stripped)
        elif current_section == "topic_map":
            topic_map.append(stripped)
        elif current_section == "timeline":
            timeline.append(stripped)
        elif current_section == "topic" and current_topic is not None:
            current_topic["lines"].append(stripped)

    for topic in topics:
        topic.update(_analyse_topic(topic["title"], topic["lines"]))

    return {
        "frontmatter": frontmatter,
        "how_to_use": how_to_use,
        "topic_map": topic_map,
        "timeline": timeline,
        "topics": topics,
    }


def _analyse_topic(topic_heading: str, lines: List[str]) -> Dict[str, Any]:
    principle = ""
    statutes: List[str] = []
    authorities: List[str] = []
    checklist: List[str] = []
    use_lines: List[str] = []
    conclusion_lines: List[str] = []
    arguments: List[Dict[str, str]] = []
    first_substantive_line = ""

    current_section = ""
    current_argument: Optional[Dict[str, str]] = None
    inside_support_block = False

    for raw in lines:
        stripped = _normalise_note_line(raw)
        if not stripped or stripped == "REVISION FRAME":
            continue
        if re.match(r"^Argument\s+\d+\s*$", stripped):
            current_section = "argument"
            inside_support_block = False
            current_argument = {"heading": stripped, "thesis": ""}
            arguments.append(current_argument)
            continue
        if stripped in NOTE_SUPPORT_HEADINGS:
            current_section = "support"
            inside_support_block = True
            continue
        if stripped == "Conclusion / comparison / development":
            current_section = "conclusion"
            inside_support_block = False
            continue
        if stripped == "Add / check in a problem question":
            current_section = "check"
            inside_support_block = False
            continue

        compact = _strip_bullet(stripped)
        if not first_substantive_line and not compact.startswith(tuple(prefix.rstrip(":") for prefix in NOTE_LABEL_PREFIXES)):
            first_substantive_line = compact

        if stripped.startswith("Principle:") and not principle:
            principle = stripped.split(":", 1)[1].strip()
        if stripped.startswith(("Statute:", "Legislation:")):
            statutes.append(stripped.split(":", 1)[1].strip())
        elif _looks_like_statute_reference(stripped):
            statutes.append(compact)
        if inside_support_block and _looks_like_case_authority_line(stripped):
            authorities.append(compact)
        if stripped.startswith("Use in an answer:"):
            use_lines.append(stripped.split(":", 1)[1].strip())
        if current_section == "conclusion":
            conclusion_lines.append(compact)
        if current_section == "check" and stripped.startswith("- "):
            checklist.append(compact)
        if current_argument is not None and current_section == "argument" and not current_argument["thesis"]:
            if not any(stripped.startswith(prefix) for prefix in NOTE_LABEL_PREFIXES):
                current_argument["thesis"] = compact

    focus = _extract_topic_title(topic_heading)
    principle_text = principle or first_substantive_line or f"Use the leading authorities to define the core rule in {focus.lower()}."
    hooks = _coalesce_unique(statutes, limit=3)
    trigger_points = _coalesce_unique(checklist or use_lines, limit=4)
    leading = _coalesce_unique([_compact_authority_name(item) for item in authorities], limit=4)

    if not hooks:
        hooks = ["Check the controlling statutory wording, documentary framework, and later authority for this topic."]
    if not trigger_points:
        trigger_points = [f"Any problem involving {focus.lower()}, doctrinal classification, timing, or remedy."]
    if not leading:
        leading = ["Retrieved leading authorities for this topic"]

    if len(arguments) >= 2:
        why_arguable = (
            "The point is arguable because the materials support more than one plausible reading. "
            "Keep the orthodox position separate from the strongest counter-position and test which source does the real work."
        )
    elif conclusion_lines:
        why_arguable = _shorten(" ".join(conclusion_lines), 190)
    else:
        why_arguable = (
            "The issue often turns on classification, scope, chronology, or remedy. "
            "Use the authorities to show why the boundary is contested rather than stating the rule abstractly."
        )

    for index, argument in enumerate(arguments):
        thesis = argument.get("thesis", "").strip()
        if thesis:
            prefix = "Primary position" if index == 0 else ("Counter-position" if index == 1 else f"Alternative position {index + 1}")
            argument["banner"] = f"{prefix}: {_shorten(thesis, 110)}"
        else:
            argument["banner"] = (
                f"Primary position on {focus.lower()}"
                if index == 0
                else f"Counter-position on {focus.lower()}"
            )

    return {
        "focus": focus,
        "principle": principle_text,
        "rule_hooks": hooks,
        "problem_triggers": trigger_points,
        "leading_authorities": leading,
        "why_arguable": why_arguable,
        "arguments": arguments,
    }


def _configure_default_report_styles(doc: Document) -> None:
    for style_name in (
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Number",
    ):
        try:
            doc.styles[style_name].font.name = "Calibri"
            doc.styles[style_name].font.size = Pt(12)
        except KeyError:
            pass


def _configure_topic_notes_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_ascii_font(normal, "Calibri", 10)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18

    title = doc.styles["Title"]
    _set_ascii_font(title, "Calibri", 24, bold=True)
    title.font.color.rgb = RGBColor.from_string(NOTE_TEXT_BLUE)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.styles["Subtitle"]
    _set_ascii_font(subtitle, "Calibri", 11, bold=False)
    subtitle.font.color.rgb = RGBColor.from_string("4A4F57")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    heading1 = doc.styles["Heading 1"]
    _set_ascii_font(heading1, "Calibri", 11.5, bold=True)
    heading1.font.color.rgb = RGBColor.from_string(NOTE_TEXT_BLUE)
    heading1.paragraph_format.space_before = Pt(22)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True

    heading2 = doc.styles["Heading 2"]
    _set_ascii_font(heading2, "Calibri", 11.5, bold=True)
    heading2.font.color.rgb = RGBColor.from_string("4E81C3")
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.keep_with_next = True
    heading2.paragraph_format.keep_together = True

    heading3 = doc.styles["Heading 3"]
    _set_ascii_font(heading3, "Calibri", 11, bold=True)
    heading3.font.color.rgb = RGBColor.from_string(NOTE_TEXT_BLUE)
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.keep_with_next = True

    list_bullet = doc.styles["List Bullet"]
    _set_ascii_font(list_bullet, "Calibri", 10)
    list_bullet.paragraph_format.space_after = Pt(4)

    list_number = doc.styles["List Number"]
    _set_ascii_font(list_number, "Calibri", 10)
    list_number.paragraph_format.space_after = Pt(4)

    topic_kicker = _ensure_paragraph_style(doc, "Topic Kicker")
    _set_ascii_font(topic_kicker, "Calibri", 10, bold=True)
    topic_kicker.font.color.rgb = RGBColor.from_string(NOTE_TEXT_BLUE)
    topic_kicker.paragraph_format.space_before = Pt(0)
    topic_kicker.paragraph_format.space_after = Pt(6)

    mini_heading = _ensure_paragraph_style(doc, "Mini Heading")
    _set_ascii_font(mini_heading, "Calibri", 11, bold=True)
    mini_heading.font.color.rgb = RGBColor.from_string("333333")
    mini_heading.paragraph_format.space_before = Pt(12)
    mini_heading.paragraph_format.space_after = Pt(4)

    support_heading = _ensure_paragraph_style(doc, "Support Heading")
    _set_ascii_font(support_heading, "Calibri", 11, bold=True)
    support_heading.font.color.rgb = RGBColor.from_string(NOTE_TEXT_BLUE)
    support_heading.paragraph_format.space_before = Pt(7)
    support_heading.paragraph_format.space_after = Pt(4)


def _apply_default_report_run_fonts(doc: Document) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
                run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _configure_topic_notes_page(doc: Document, title_text: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(title_text)
    _set_run_font(header_run, size=10, bold=True, color="9BA7B4")
    _set_paragraph_border(header_para, side="bottom", color=NOTE_GRID, size=6, space=1)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_border(footer_para, side="top", color=NOTE_GRID, size=6, space=1)
    page_label = footer_para.add_run("Page ")
    _set_run_font(page_label, size=10, color="9BA7B4")
    _add_page_field(footer_para)


def _write_single_run_paragraph(
    paragraph,
    text: str,
    *,
    size: float = 10,
    bold: bool = False,
    italic: bool = False,
    color: str = "222222",
) -> None:
    if paragraph.text:
        paragraph.clear()
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def _add_cover_bar(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.6)
    _set_cell_shading(cell, NOTE_DARK_BLUE)
    _set_cell_border(
        cell,
        left={"val": "dashed", "sz": 6, "color": NOTE_GRID},
        top={"val": "dashed", "sz": 6, "color": NOTE_GRID},
        right={"val": "dashed", "sz": 6, "color": NOTE_GRID},
        bottom={"val": "dashed", "sz": 6, "color": NOTE_GRID},
    )
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(" ")
    _set_run_font(run, size=6, color="FFFFFF")


def _set_cell_text(
    cell,
    text: str,
    *,
    size: float = 10.0,
    bold: bool = False,
    color: str = "222222",
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.1
    para.clear()
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_block_gap(doc: Document, *, size: float = 5.5) -> None:
    paragraph = doc.add_paragraph("")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(" ")
    _set_run_font(run, size=size, color="FFFFFF")


def _add_summary_table(doc: Document, subject_label: str) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.9), Inches(4.9))

    rows = [
        ("What this document does", "What you get"),
        ("Focus", f"Answer-ready notes for contentious and high-yield {subject_label.lower()} issues."),
        ("Structure", "Each topic follows the same frame: principle -> competing sides -> cases -> development -> add/check."),
        ("Use", "Designed for revision, problem questions and structured comparison of authority over time."),
    ]

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = widths[col_index]
            if row_index == 0:
                _set_cell_shading(cell, NOTE_DARK_BLUE)
                _set_cell_text(cell, value, size=10, bold=True, color="FFFFFF")
            else:
                _set_cell_shading(cell, NOTE_BODY_FILL if col_index == 1 else "FFFFFF")
                _set_cell_text(
                    cell,
                    value,
                    size=10,
                    bold=(col_index == 0),
                    color=NOTE_TEXT_BLUE if col_index == 0 else "222222",
                )

    _style_all_cell_borders(table, color=NOTE_GRID, size=8)
    _add_block_gap(doc, size=5.0)


def _add_notice_box(
    doc: Document,
    label: str,
    body: str,
    *,
    fill: str = NOTE_YELLOW_FILL,
    border: str = NOTE_YELLOW_BORDER,
    label_color: str = NOTE_YELLOW_TEXT,
    body_color: str = "333333",
    size: float = 10,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    _set_cell_shading(cell, fill)
    _set_cell_border(
        cell,
        left={"val": "single", "sz": 8, "color": border},
        top={"val": "single", "sz": 8, "color": border},
        right={"val": "single", "sz": 8, "color": border},
        bottom={"val": "single", "sz": 8, "color": border},
    )
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    lead = para.add_run(f"{label}. ")
    _set_run_font(lead, size=size, bold=True, color=label_color)
    tail = para.add_run(body)
    _set_run_font(tail, size=size, color=body_color)
    _add_block_gap(doc, size=5.5)


def _add_topic_map_table(doc: Document, topics: List[Dict[str, Any]], fallback_lines: List[str]) -> None:
    rows: List[Tuple[str, str, str]] = [("Topic", "Focus", "Leading authorities")]
    if topics:
        for topic in topics:
            rows.append(
                (
                    _extract_topic_label(topic["title"]),
                    topic["focus"],
                    "; ".join(topic["leading_authorities"]),
                )
            )
    else:
        for index, line in enumerate(fallback_lines, start=1):
            rows.append((f"Topic {index}", _shorten(_strip_bullet(line), 80), "Retrieved authorities"))

    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(0.9), Inches(3.7), Inches(2.05))

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = widths[col_index]
            if row_index == 0:
                _set_cell_shading(cell, NOTE_DARK_BLUE)
                _set_cell_text(cell, value, size=10, bold=True, color="FFFFFF")
            else:
                _set_cell_shading(cell, "FFFFFF" if col_index == 0 else NOTE_BODY_FILL)
                _set_cell_text(
                    cell,
                    value,
                    size=10,
                    bold=(col_index == 0),
                    color=NOTE_TEXT_BLUE if col_index == 0 else "222222",
                )

    _style_all_cell_borders(table, color=NOTE_GRID, size=8)
    _add_block_gap(doc, size=5.0)


def _add_timeline_table(doc: Document, lines: List[str]) -> None:
    entries = [_split_timeline_entry(line, idx) for idx, line in enumerate(lines)]
    if not entries:
        entries = [("Development", "Use the retrieved materials to trace the doctrine from its starting point to the current position.")]

    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.1), Inches(5.55))

    header = ("Period", "Development")
    for col_index, value in enumerate(header):
        cell = table.cell(0, col_index)
        cell.width = widths[col_index]
        _set_cell_shading(cell, NOTE_DARK_BLUE)
        _set_cell_text(cell, value, size=10, bold=True, color="FFFFFF")

    for row_index, (period, development) in enumerate(entries, start=1):
        period_cell = table.cell(row_index, 0)
        development_cell = table.cell(row_index, 1)
        period_cell.width = widths[0]
        development_cell.width = widths[1]
        _set_cell_shading(period_cell, "FFFFFF")
        _set_cell_shading(development_cell, NOTE_BODY_FILL)
        _set_cell_text(period_cell, period, size=10, bold=True, color=NOTE_TEXT_BLUE)
        _set_cell_text(development_cell, development, size=10, color="222222")

    _style_all_cell_borders(table, color=NOTE_GRID, size=8)
    _add_block_gap(doc, size=5.0)


def _add_topic_summary_table(doc: Document, topic: Dict[str, Any]) -> None:
    rows = [
        ("Core principle", topic["principle"]),
        ("Key statutory / rule hooks", "; ".join(topic["rule_hooks"])),
        ("Problem-question triggers", "; ".join(topic["problem_triggers"])),
        ("Leading authorities", "; ".join(topic["leading_authorities"])),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.85), Inches(4.95))

    for row_index, (label, value) in enumerate(rows):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)
        label_cell.width = widths[0]
        value_cell.width = widths[1]
        _set_cell_shading(label_cell, "FFFFFF")
        _set_cell_shading(value_cell, NOTE_BODY_FILL)
        _set_cell_text(label_cell, label, size=10, bold=True, color=NOTE_TEXT_BLUE)
        _set_cell_text(value_cell, value, size=10, color="222222")

    _style_all_cell_borders(table, color=NOTE_GRID, size=8)
    _add_block_gap(doc, size=5.0)


def _add_position_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.65)
    _set_cell_shading(cell, NOTE_LIGHT_BLUE_FILL)
    _set_cell_border(
        cell,
        left={"val": "single", "sz": 8, "color": NOTE_LIGHT_BLUE_BORDER},
        top={"val": "single", "sz": 8, "color": NOTE_LIGHT_BLUE_BORDER},
        right={"val": "single", "sz": 8, "color": NOTE_LIGHT_BLUE_BORDER},
        bottom={"val": "single", "sz": 8, "color": NOTE_LIGHT_BLUE_BORDER},
    )
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    _set_run_font(run, size=11, bold=True, color=NOTE_LIGHT_BLUE_TEXT)
    _add_block_gap(doc, size=4.5)


def _add_topic_notes_prefixed_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if text.startswith("Principle:"):
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(9)
    elif text.startswith(("Facts:", "Held:", "Reasoning:", "Use in an answer:", "Argument:")):
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(7)
    for prefix in NOTE_LABEL_PREFIXES:
        if text.startswith(prefix):
            lead = paragraph.add_run(prefix)
            _set_run_font(lead, size=10, bold=True, color=NOTE_TEXT_BLUE)
            tail = text[len(prefix) :].lstrip()
            if tail:
                tail_run = paragraph.add_run(f" {tail}")
                _set_run_font(tail_run, size=10, color="222222")
            return
    run = paragraph.add_run(text)
    _set_run_font(run, size=10, color="222222")


def _render_topic_notes_topic_body(doc: Document, topic: Dict[str, Any]) -> None:
    doc.add_paragraph(topic["title"], style="Heading 1")
    doc.add_paragraph("REVISION FRAME", style="Topic Kicker")
    _add_topic_summary_table(doc, topic)
    _add_notice_box(
        doc,
        "Why this topic is arguable",
        topic["why_arguable"],
        fill=NOTE_YELLOW_FILL,
        border=NOTE_YELLOW_BORDER,
        label_color=NOTE_YELLOW_TEXT,
        body_color="333333",
    )

    in_support_section = False
    argument_index = 0
    for raw in topic["lines"]:
        stripped = _normalise_note_line(raw)
        if not stripped or stripped == "REVISION FRAME":
            continue
        if re.match(r"^Argument\s+\d+\s*$", stripped):
            if argument_index < len(topic["arguments"]):
                _add_position_box(doc, topic["arguments"][argument_index]["banner"])
            doc.add_paragraph(stripped, style="Mini Heading")
            argument_index += 1
            in_support_section = False
            continue
        if stripped in NOTE_HEADING_2_LINES:
            doc.add_paragraph(stripped, style="Heading 2")
            in_support_section = False
            continue
        if stripped in NOTE_SUPPORT_HEADINGS:
            doc.add_paragraph(stripped, style="Support Heading")
            in_support_section = True
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(_strip_bullet(stripped), style="List Bullet")
            in_support_section = False
            continue
        if any(stripped.startswith(prefix) for prefix in NOTE_LABEL_PREFIXES):
            _add_topic_notes_prefixed_paragraph(doc, stripped)
            continue
        if _is_numbered(stripped):
            doc.add_paragraph(_strip_numeric_prefix(stripped), style="List Number")
            continue
        if in_support_section:
            para = doc.add_paragraph(style="Normal")
            run = para.add_run(stripped)
            _set_run_font(run, size=10, bold=True, color="222222")
            continue
        doc.add_paragraph(stripped, style="Normal")


def _render_topic_notes_docx(doc: Document, input_text: str) -> None:
    _configure_topic_notes_styles(doc)
    parsed = _parse_topic_notes_input(input_text)
    title_text = parsed["frontmatter"][0] if parsed["frontmatter"] else "Law Principles Notes"
    subtitle_text = (
        parsed["frontmatter"][1]
        if len(parsed["frontmatter"]) >= 2
        else "Refined revision notes on arguable principles, key cases, rule analysis and doctrinal development"
    )
    subject_label = re.sub(r"\s+principles notes$", "", title_text, flags=re.IGNORECASE).strip() or "law"
    doc.core_properties.title = title_text
    _configure_topic_notes_page(doc, title_text)

    title_para = doc.add_paragraph(title_text, style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_border(title_para, side="bottom", color=NOTE_TITLE_RULE, size=8, space=3)
    subtitle_para = doc.add_paragraph(subtitle_text, style="Subtitle")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_cover_bar(doc)
    _add_summary_table(doc, subject_label)
    _add_notice_box(
        doc,
        "Study note only",
        "On a live file or in an assessed answer, always re-check the primary text, statutory wording, later authority, and any factual chronology before relying on a point.",
    )
    doc.add_page_break()

    doc.add_paragraph("How to use these notes", style="Heading 1")
    _add_notice_box(
        doc,
        "Suggested answer pattern",
        "State the principle first. If the point is arguable, run both sides separately. Under each side, identify the argument, then use the case law in a strict order: facts -> held -> reasoning -> how to use the case. End with a short comparison showing how the law developed.",
    )
    for line in parsed["how_to_use"]:
        if line.startswith("- "):
            doc.add_paragraph(_strip_bullet(line), style="List Bullet")
        else:
            doc.add_paragraph(line, style="Normal")

    doc.add_paragraph("Topic map", style="Heading 2")
    _add_topic_map_table(doc, parsed["topics"], parsed["topic_map"])

    doc.add_paragraph("Quick development timeline", style="Heading 2")
    _add_timeline_table(doc, parsed["timeline"])
    _add_notice_box(
        doc,
        "Revision tip",
        "When a question asks how the law developed, do not just list cases. Explain the sequence: starting principle -> later clarification -> modern refinement -> what that means in an answer.",
    )

    for topic in parsed["topics"]:
        _render_topic_notes_topic_body(doc, topic)


def build_docx(input_text: str, out_path: Path) -> None:
    doc = Document()
    if _looks_like_topic_notes_document(input_text):
        _render_topic_notes_docx(doc, input_text)
        doc.save(out_path)
        return

    in_ledger = False
    for raw in input_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            heading = _strip_markdown_inline(stripped[3:].strip())
            doc.add_heading(heading, level=2)
            in_ledger = "verification ledger" in heading.lower()
            continue
        if stripped.startswith("### "):
            heading = _strip_markdown_inline(stripped[4:].strip())
            doc.add_heading(heading, level=3)
            in_ledger = "verification ledger" in heading.lower()
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(_strip_markdown_inline(stripped[2:].strip()), style="List Bullet")
            continue
        if _is_numbered(stripped):
            item = _strip_numeric_prefix(stripped)
            item = _normalize_ledger_item(item) if in_ledger else _strip_markdown_inline(item)
            doc.add_paragraph(item, style="List Number")
            continue
        doc.add_paragraph(_strip_markdown_inline(line))

    _configure_default_report_styles(doc)
    _apply_default_report_run_fonts(doc)
    doc.save(out_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate DOCX review report from text/markdown."
    )
    parser.add_argument("--input", required=True, help="Path to .md/.txt report input")
    parser.add_argument("--out", required=True, help="Path to output .docx")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    out_path = Path(args.out).expanduser().resolve()

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if out_path.parent != DESKTOP_ROOT:
        raise ValueError(
            "Output report must be written directly to Desktop root "
            f"({DESKTOP_ROOT}), not inside a folder: {out_path}"
        )
    if out_path.suffix.lower() != ".docx":
        raise ValueError(f"Output report must be a .docx file: {out_path}")

    text = input_path.read_text(encoding="utf-8")
    build_docx(text, out_path)
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
