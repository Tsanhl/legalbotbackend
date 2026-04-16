from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


ITALIC_PATTERN = re.compile(r"\*([^*]+)\*")


def add_markdown_runs(paragraph, text: str) -> None:
    pos = 0
    for match in ITALIC_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(1))
        run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python3 scripts/markdown_to_docx.py <input.md> <output.docx>")
        return 1

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    text = src.read_text(encoding="utf-8").strip()
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]

    if not blocks:
        print("Input file is empty")
        return 1

    doc = Document()
    title = blocks[0].strip()
    doc.add_heading(title, level=0)

    for block in blocks[1:]:
        if block == "References":
            doc.add_heading(block, level=1)
            continue

        lines = block.splitlines()
        joined = " ".join(line.strip() for line in lines if line.strip())
        paragraph = doc.add_paragraph()
        add_markdown_runs(paragraph, joined)

    doc.save(dst)
    print(dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
