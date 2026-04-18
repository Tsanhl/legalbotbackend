"""
Backend answer-runtime helpers for Legal AI.

This module keeps the deterministic answer-shaping and continuation logic that
used to live in the Streamlit app, without any UI dependency.
"""
import json
import base64
import os
import re
import bisect
import math
import shutil
import tempfile
from collections import Counter
from difflib import SequenceMatcher
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
import uuid

# Import services
from knowledge_base import load_law_resource_index, get_knowledge_base_summary
from word_count_rules import complete_word_count_floor, complete_word_count_window
from model_applicable_service import (
    LONG_ESSAY_THRESHOLD,
    PROBLEM_FINAL_CONCLUSION_TITLE,
    PROBLEM_REMEDIES_LIABILITY_TITLE,
    SUPPORTED_LLM_PROVIDERS,
    _is_problem_final_conclusion_title,
    _is_problem_remedies_liability_title,
    detect_provider_from_api_key,
    get_default_model_for_provider,
    get_provider_display_name,
    get_provider_model_presets,
    get_provider_model_placeholder,
    initialize_knowledge_base, 
    send_message_with_docs as _provider_send_message_with_docs,
    reset_session,
    resolve_llm_provider,
    resolve_model_name_for_provider,
    resolve_provider_api_key,
    encode_file_to_base64,
    detect_long_essay,
    detect_unit_structure_policy_violation,
    extract_word_targets_from_prompt,
    get_allowed_authorities_from_rag,
    sanitize_output_against_allowlist,
    strip_internal_reasoning
)
try:
    from model_applicable_service import (
        _backend_request_requires_mandatory_rag,
        detect_topic_notes_request,
        _infer_retrieval_profile,
        _subissue_queries_for_unit,
        register_topic_notes_cleanup_paths,
    )
except Exception:
    _backend_request_requires_mandatory_rag = None
    detect_topic_notes_request = None
    _infer_retrieval_profile = None
    register_topic_notes_cleanup_paths = None
    _subissue_queries_for_unit = None

try:
    from legal_doc_tools.workflow import (
        DOCX_MIME as LEGAL_DOCX_MIME,
        run_auto_legal_doc_amend_workflow,
        wants_legal_doc_amend,
        wants_local_legal_doc_amend,
    )
    LEGAL_DOC_WORKFLOW_AVAILABLE = True
except Exception as e:
    print(f"Legal DOCX amend workflow not available: {e}")
    LEGAL_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    LEGAL_DOC_WORKFLOW_AVAILABLE = False

    def wants_legal_doc_amend(message, documents):
        return False

    def wants_local_legal_doc_amend(message):
        return False

BACKEND_ANSWER_OUTPUT_CHAT = "chat"
BACKEND_ANSWER_OUTPUT_MARKDOWN = "markdown"
BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT = "markdown_artifact"
BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT = "docx_artifact"
BACKEND_COMPLETE_ANSWER_REGEN_MAX_ATTEMPTS = 1
BACKEND_ANSWER_ARTIFACT_DESKTOP_ROOT = (Path.home() / "Desktop").resolve()
BACKEND_ANSWER_ONE_OFF_HINTS = (
    "one_off",
    "one-off",
    "tmp",
    "temp",
    "scratch",
    "draft",
    "helper",
    "context_dump",
    "prompt_dump",
    "question_pack",
)
BACKEND_ANSWER_ONE_OFF_SUFFIXES = {
    ".txt",
    ".md",
    ".json",
    ".docx",
    ".csv",
    ".log",
    ".tmp",
    ".prompt",
}


def resolve_backend_answer_output_mode(
    prompt_text: str = "",
    output_mode: Optional[str] = None,
) -> str:
    """
    Resolve the backend delivery preference for complete answers.

    The canonical runtime path returns answer text directly in chat/API.
    If the user explicitly asks for Markdown or `.md`, we still return direct
    text, but treat it as markdown-compatible output rather than as a file-write
    instruction.
    """
    explicit = re.sub(r"[^a-z]", "", (output_mode or "").strip().lower())
    explicit_map = {
        "chat": BACKEND_ANSWER_OUTPUT_CHAT,
        "api": BACKEND_ANSWER_OUTPUT_CHAT,
        "text": BACKEND_ANSWER_OUTPUT_CHAT,
        "direct": BACKEND_ANSWER_OUTPUT_CHAT,
        "markdown": BACKEND_ANSWER_OUTPUT_MARKDOWN,
        "md": BACKEND_ANSWER_OUTPUT_MARKDOWN,
        "markdownartifact": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "markdownfile": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "mdartifact": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "mdfile": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "docx": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "word": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "worddoc": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "worddocument": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "documentfile": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
    }
    if explicit in explicit_map:
        return explicit_map[explicit]

    low = (prompt_text or "").lower()
    if (
        ".md" in low
        or "markdown" in low
        or "mark down" in low
    ):
        return BACKEND_ANSWER_OUTPUT_MARKDOWN
    return BACKEND_ANSWER_OUTPUT_CHAT


def resolve_backend_answer_delivery_mode(
    prompt_text: str = "",
    output_mode: Optional[str] = None,
) -> str:
    """
    Resolve the final delivery mode for backend complete answers.

    Distinguishes between:
    - direct chat/API text,
    - direct markdown-compatible text,
    - a saved markdown artifact in the project/workspace, and
    - a saved DOCX artifact on Desktop.
    """
    explicit = re.sub(r"[^a-z]", "", (output_mode or "").strip().lower())
    explicit_map = {
        "chat": BACKEND_ANSWER_OUTPUT_CHAT,
        "api": BACKEND_ANSWER_OUTPUT_CHAT,
        "text": BACKEND_ANSWER_OUTPUT_CHAT,
        "direct": BACKEND_ANSWER_OUTPUT_CHAT,
        "markdown": BACKEND_ANSWER_OUTPUT_MARKDOWN,
        "md": BACKEND_ANSWER_OUTPUT_MARKDOWN,
        "markdownartifact": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "markdownfile": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "mdartifact": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "mdfile": BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        "docx": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "word": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "worddoc": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "worddocument": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
        "documentfile": BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
    }
    if explicit in explicit_map:
        return explicit_map[explicit]

    low = (prompt_text or "").lower()
    artifact_verbs = ("save", "create", "export", "write", "generate", "put")
    mentions_file_location = any(term in low for term in ("file", "project", "workspace", "desktop", "folder"))

    if any(term in low for term in (".docx", "word document", "word doc")):
        if mentions_file_location or any(verb in low for verb in artifact_verbs):
            return BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT

    if ".md" in low:
        if mentions_file_location or any(verb in low for verb in artifact_verbs):
            return BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT
    if any(term in low for term in ("markdown file", "md file")):
        if mentions_file_location or any(verb in low for verb in artifact_verbs):
            return BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT

    return resolve_backend_answer_output_mode(prompt_text, output_mode=output_mode)


def _backend_answer_output_instruction(output_mode: str) -> str:
    if output_mode == BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT:
        return (
            "Return the final answer directly as clean markdown-compatible text suitable for DOCX rendering. "
            "Do NOT create or save the `.docx` file inside the model response; the runtime will create the artifact."
        )
    if output_mode == BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT:
        return (
            "Return the final answer directly as clean markdown-compatible text. "
            "Do NOT create or save the `.md` file inside the model response; the runtime will create the artifact."
        )
    if output_mode == BACKEND_ANSWER_OUTPUT_MARKDOWN:
        return (
            "Return the final answer directly in markdown-compatible text in the response body. "
            "Do NOT create or save a `.md` file unless the caller separately asks for an artifact."
        )
    return (
        "Return the final answer directly in the response body for chat/API delivery. "
        "Do NOT create or save a file unless the caller separately asks for an artifact."
    )


def _suggest_complete_answer_artifact_stem(prompt_text: str, project_id: str = "") -> str:
    candidates: List[str] = []
    for raw_line in (prompt_text or "").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue
        low = line.lower()
        if re.fullmatch(r"\d+\s+words?\.?", low):
            continue
        if low.startswith("in your answer") or low.startswith("advise the parties"):
            continue
        candidates.append(line)
        if len(candidates) >= 3:
            break

    seed = candidates[0] if candidates else (project_id or "backend_answer")
    seed = re.sub(r"[^\w\s-]", " ", seed, flags=re.UNICODE)
    seed = re.sub(r"[\s_-]+", "_", seed).strip("_").lower()
    return (seed[:80] or "backend_answer").strip("_") or "backend_answer"


def _resolve_complete_answer_artifact_path(
    *,
    delivery_mode: str,
    prompt_text: str,
    project_id: str,
    artifact_path: Optional[str] = None,
) -> Path:
    expected_suffix = ".docx" if delivery_mode == BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT else ".md"
    default_root = (
        BACKEND_ANSWER_ARTIFACT_DESKTOP_ROOT
        if delivery_mode == BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT
        else Path.cwd()
    )

    if artifact_path:
        path = Path(str(artifact_path)).expanduser()
        if not path.is_absolute():
            path = default_root / path
    else:
        stem = _suggest_complete_answer_artifact_stem(prompt_text, project_id=project_id)
        path = default_root / f"{stem}_backend_answer{expected_suffix}"

    if path.suffix.lower() != expected_suffix:
        path = path.with_suffix(expected_suffix)
    return path.resolve()


def _path_is_within(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


def _backend_answer_name_has_one_off_hint(path: Path) -> bool:
    low = path.name.lower().replace(" ", "_")
    return any(hint in low for hint in BACKEND_ANSWER_ONE_OFF_HINTS)


def _normalize_complete_answer_cleanup_paths(paths: Optional[List[Any]]) -> List[Path]:
    normalized: List[Path] = []
    for path in paths or []:
        if path is None:
            continue
        try:
            normalized.append(Path(str(path)).expanduser().resolve())
        except Exception:
            continue
    return normalized


def _is_complete_answer_cleanup_candidate(path: Path) -> bool:
    resolved = path.expanduser().resolve()
    project_root = Path.cwd().resolve()
    desktop_root = BACKEND_ANSWER_ARTIFACT_DESKTOP_ROOT.expanduser().resolve()
    home_root = Path.home().resolve()

    if resolved in {project_root, desktop_root, home_root}:
        return False

    try:
        temp_root = Path(tempfile.gettempdir()).resolve()
    except Exception:
        temp_root = Path("/tmp").resolve()
    temp_roots = {temp_root, Path("/tmp").resolve()}

    if any(_path_is_within(resolved, root) for root in temp_roots):
        return True

    relevant_nodes = [resolved, *resolved.parents]
    if resolved.is_dir():
        return any(
            node != project_root
            and _path_is_within(node, project_root)
            and _backend_answer_name_has_one_off_hint(node)
            for node in relevant_nodes
        )

    if resolved.suffix.lower() not in BACKEND_ANSWER_ONE_OFF_SUFFIXES:
        return False

    return any(
        node != project_root
        and _path_is_within(node, project_root)
        and _backend_answer_name_has_one_off_hint(node)
        for node in relevant_nodes
    )


def _cleanup_complete_answer_one_off_artifacts(
    paths: Optional[List[Any]],
    *,
    protected_paths: Optional[List[Any]] = None,
) -> int:
    seen: set[Path] = set()
    removed = 0
    protected_nodes = {
        Path(str(path)).expanduser().resolve()
        for path in (protected_paths or [])
        if path is not None
    }
    protected_dirs = {
        Path.home().resolve(),
        BACKEND_ANSWER_ARTIFACT_DESKTOP_ROOT.expanduser().resolve(),
        Path.cwd().resolve(),
    }

    for resolved in _normalize_complete_answer_cleanup_paths(paths):
        if resolved in seen:
            continue
        seen.add(resolved)
        if resolved in protected_nodes:
            continue
        if not _is_complete_answer_cleanup_candidate(resolved):
            continue
        if resolved.is_dir():
            if resolved in protected_dirs:
                continue
            try:
                shutil.rmtree(resolved, ignore_errors=False)
                removed += 1
            except OSError:
                continue
            continue
        try:
            resolved.unlink(missing_ok=True)
            removed += 1
        except OSError:
            continue
    return removed


def _render_complete_answer_artifact_text(answer_text: str) -> str:
    text = _strip_generation_artifacts(answer_text or "")
    text = _normalize_output_style(text)
    text = _restore_paragraph_separation(text)
    return text.strip()


def _write_complete_answer_docx(text: str, output_path: Path) -> None:
    try:
        from legal_doc_tools.generate_review_report_docx import build_docx

        build_docx(text, output_path)
        return
    except Exception:
        from docx import Document

        doc = Document()
        for block in [b.strip() for b in re.split(r"\n\s*\n", text or "") if b.strip()]:
            if re.match(r"(?im)^part\s+[ivxlcdm0-9]+\s*:", block):
                doc.add_heading(block, level=1)
                continue
            if re.match(r"(?im)^[a-d]\.\s+", block):
                doc.add_heading(block, level=2)
                continue
            lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
            if len(lines) == 1 and lines[0].startswith("- "):
                doc.add_paragraph(lines[0][2:].strip(), style="List Bullet")
                continue
            doc.add_paragraph("\n".join(lines) if lines else block)
        doc.save(output_path)


def write_complete_answer_artifact(
    answer_text: str,
    *,
    delivery_mode: str,
    prompt_text: str = "",
    project_id: str = "",
    artifact_path: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Persist a complete answer artifact from the already-generated backend answer.

    This must only run after the canonical backend answer pipeline finishes, so
    chat/API, markdown files, and DOCX files all derive from the same verified
    answer text.
    """
    if delivery_mode not in {
        BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
    }:
        return None

    output_path = _resolve_complete_answer_artifact_path(
        delivery_mode=delivery_mode,
        prompt_text=prompt_text,
        project_id=project_id,
        artifact_path=artifact_path,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    text = _render_complete_answer_artifact_text(answer_text)

    if delivery_mode == BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT:
        body = text if text.endswith("\n") else f"{text}\n"
        output_path.write_text(body, encoding="utf-8")
    else:
        _write_complete_answer_docx(text, output_path)

    return {
        "mode": delivery_mode,
        "path": str(output_path),
    }


def _attach_complete_answer_artifact_meta(
    response_meta: Any,
    artifact_info: Optional[Dict[str, Any]],
) -> Any:
    if not artifact_info:
        return response_meta
    wrapped = {"backend_answer_artifact": artifact_info}
    if response_meta is None or response_meta == "":
        return [wrapped]
    if isinstance(response_meta, dict):
        merged = dict(response_meta)
        merged["backend_answer_artifact"] = artifact_info
        return merged
    if isinstance(response_meta, list):
        cleaned = [
            item for item in response_meta
            if not (isinstance(item, dict) and "backend_answer_artifact" in item)
        ]
        return [*cleaned, wrapped]
    return [response_meta, wrapped]


def _extract_complete_answer_artifact_meta(response_meta: Any) -> Optional[Dict[str, Any]]:
    if isinstance(response_meta, dict):
        info = response_meta.get("backend_answer_artifact")
        return info if isinstance(info, dict) else None
    if isinstance(response_meta, list):
        for item in response_meta:
            if isinstance(item, dict) and isinstance(item.get("backend_answer_artifact"), dict):
                return item["backend_answer_artifact"]
    return None


def _looks_like_backend_complete_answer_request(
    message: str,
    documents: Optional[List[Dict[str, Any]]] = None,
) -> bool:
    low = (message or "").lower()
    if not low.strip():
        return False
    if wants_local_legal_doc_amend(message) or wants_legal_doc_amend(message, documents or []):
        return False
    if _backend_request_requires_mandatory_rag is not None:
        try:
            if not _backend_request_requires_mandatory_rag(message):
                return False
        except Exception:
            pass
    if _extract_word_targets(message):
        return True
    return any(
        marker in low
        for marker in [
            "essay question",
            "problem question",
            "complete answer",
            "full answer",
            "critically evaluate",
            "critically discuss",
            "in your answer",
            "advise ",
        ]
    )


def _iter_recent_user_prompts(messages: List[Dict[str, Any]]) -> List[str]:
    prompts: List[str] = []
    for msg in reversed(messages or []):
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        prompts.append(txt)
    return prompts


def _resolve_complete_answer_word_window(
    prompt_text: str,
    messages: List[Dict[str, Any]],
    *,
    enforce_long_response_split: bool,
) -> Optional[Tuple[int, int]]:
    """
    Resolve the strict answer window for backend complete answers.

    Backend direct answers use the FULL requested target (99-100%), not the
    website multi-part per-chunk window. Split mode retains the existing
    continuation-aware behaviour.
    """
    if enforce_long_response_split:
        return _resolve_word_window_from_history(prompt_text, messages)

    current_targets = _extract_word_targets(prompt_text or "")
    if current_targets and (not _is_continuation_command(prompt_text or "")):
        return complete_word_count_window(sum(int(n) for n in current_targets))

    if _is_continuation_command(prompt_text or ""):
        return _resolve_word_window_from_history(prompt_text, messages)

    for txt in _iter_recent_user_prompts(messages or []):
        anchor_targets = _extract_word_targets(txt)
        if anchor_targets:
            return complete_word_count_window(sum(int(n) for n in anchor_targets))
    return None


def _extract_profile_prompt_map_requirements(prompt_text: str) -> List[Dict[str, Any]]:
    if _infer_retrieval_profile is None:
        return []
    try:
        profile = _infer_retrieval_profile(prompt_text or "")
    except Exception:
        return []
    asks = [str(a or "").strip() for a in (profile.get("prompt_map_asks") or []) if str(a or "").strip()]
    if not asks:
        return []

    stop = {
        "the", "and", "for", "with", "that", "this", "from", "into", "what", "when",
        "where", "which", "whether", "should", "would", "could", "under", "between",
        "about", "their", "they", "them", "your", "answer", "consider", "evaluate",
        "discuss", "critically", "advise", "role", "impact", "extent", "basis",
    }
    out: List[Dict[str, Any]] = []
    for ask in asks[:8]:
        words = [
            tok for tok in re.findall(r"\b[a-z][a-z0-9\-]{3,}\b", ask.lower())
            if tok not in stop
        ]
        keywords: List[str] = []
        for tok in words:
            if tok not in keywords:
                keywords.append(tok)
            if len(keywords) >= 8:
                break
        if keywords:
            out.append({"label": ask, "keywords": keywords})
    return out


def _missing_profile_prompt_map_asks(answer_text: str, prompt_text: str) -> List[str]:
    low_txt = (answer_text or "").lower()
    if not low_txt.strip():
        return []
    missing: List[str] = []
    for item in _extract_profile_prompt_map_requirements(prompt_text):
        kws = item.get("keywords") or []
        hits = sum(1 for kw in kws[:8] if re.search(rf"\b{re.escape(kw)}\b", low_txt))
        needed = 1 if len(kws) <= 3 else 2
        if hits < needed:
            missing.append(item.get("label") or "")
    return [m for m in missing if m]


def _complete_answer_sentence_support_issues(answer_text: str) -> List[str]:
    """
    Deterministic sentence-support verification for complete answers.

    This mirrors the amend workflow's stricter expectation that argumentative
    propositions should be individually supported, while using lightweight text
    heuristics suitable for plain-answer output.
    """
    txt = _strip_generation_artifacts(answer_text or "").strip()
    if not txt:
        return []

    sentence_candidates = re.split(r"(?<=[.!?])\s+", txt)
    missing: List[str] = []

    legal_signal_terms = {
        "act", "acts", "article", "articles", "authority", "authorities", "breach",
        "breaches", "claim", "claims", "conclusion", "court", "courts", "damages",
        "defence", "duty", "duties", "exercise", "grave", "habitual", "harm", "law",
        "legal", "liable", "liability", "likely", "negligence", "remedy", "remedies",
        "residence", "retention", "return", "right", "rights", "risk", "settled",
        "settlement", "should", "statute", "statutory", "undertakings", "wrongful",
        "custody", "objections", "resident", "residence", "jurisdiction", "means",
        "requires", "prevent", "prevents", "permits", "therefore", "because",
        "arguably", "better", "view", "must", "would", "may", "under",
    }

    case_pat = re.compile(
        r"\b([A-Z][A-Za-z0-9'’.\-]*(?:\s+[A-Za-z][A-Za-z0-9'’.\-]*){0,10})\s+v\.?\s+"
        r"([A-Z][A-Za-z0-9'’.\-]*(?:\s+[A-Za-z][A-Za-z0-9'’.\-]*){0,10})\b"
    )
    inline_oscola_pat = re.compile(
        r"\([^()\n]{3,260}(?:\[[12]\d{3}\]|Act\s+\d{4}|Article\s+\d+(?:\(\d+\))?|section\s+\d+|s\.?\s*\d+)",
        flags=re.IGNORECASE,
    )
    harvard_pat = re.compile(
        r"\((?:[A-Z][A-Za-z&.\-'\s]{1,80}|[A-Z][A-Za-z][^()\n]{0,100}),\s*(?:\d{4}[a-z]?|n\.d\.|no date)(?:,\s*(?:p{1,2}\.|para\.?)\s*[^)]+)?\)",
    )

    def _is_heading_like(s: str) -> bool:
        stripped = (s or "").strip()
        return bool(
            re.match(
                r"(?im)^(?:question\s+\d+\s*:|part\s+[ivxlcdm0-9]+\s*:|[A-D]\.\s+\w+|\(end of answer\)|will continue\b)",
                stripped,
            )
        )

    def _looks_argumentative(s: str) -> bool:
        stripped = (s or "").strip()
        if not stripped or _is_heading_like(stripped):
            return False
        if stripped.startswith("(") and stripped.endswith(")"):
            return False
        words = re.findall(r"\b[A-Za-z][A-Za-z'’/-]*\b", stripped)
        if len(words) < 8:
            return False
        lower_words = {w.lower() for w in words}
        if lower_words & legal_signal_terms:
            return True
        if case_pat.search(stripped):
            return True
        if re.search(r"\b[A-Z][A-Za-z ,&()'-]+ Act \d{4}\b", stripped):
            return True
        if re.search(r"\bArticle\s+\d+(?:\(\d+\))?\b", stripped, flags=re.IGNORECASE):
            return True
        if re.search(r"\b(?:section|s\.?)\s*\d+[a-z]?(?:\(\d+\))?\b", stripped, flags=re.IGNORECASE):
            return True
        return False

    def _has_inline_support(s: str) -> bool:
        stripped = (s or "").strip()
        return bool(inline_oscola_pat.search(stripped) or harvard_pat.search(stripped))

    for sent in sentence_candidates:
        s = (sent or "").strip()
        if not _looks_argumentative(s):
            continue
        if _has_inline_support(s):
            continue
        missing.append(re.sub(r"\s+", " ", s)[:120])

    if not missing:
        return []

    preview = "; ".join(missing[:3])
    return [
        "Argumentative sentence-support verification failed: "
        f"{len(missing)} argumentative sentence(s) lack immediate inline authority support. "
        f"Examples: {preview}"
    ]


def _strict_complete_answer_issues(
    answer_text: str,
    prompt_text: str,
    messages: List[Dict[str, Any]],
    *,
    enforce_long_response_split: bool,
) -> List[str]:
    """
    Apply the deterministic amend-style quality gate to complete-answer output.

    This is intentionally stricter than the provider's built-in guidance: it
    checks structure, prompt-map coverage, history-aware continuity, and the
    backend-only 99-100% answer word window.
    """
    issues: List[str] = []
    unit_mode = _current_unit_mode_from_history(prompt_text, messages or [])
    is_problem_mode = bool(unit_mode.get("is_problem_mode"))
    issues.extend(
        _essay_quality_issues(
            answer_text,
            prompt_text,
            is_short_single_essay=_is_short_single_essay_prompt(prompt_text),
            is_problem_mode=is_problem_mode,
        )
    )
    issues.extend(_history_aware_structure_issues(answer_text, prompt_text, messages or []))
    issues.extend(_direct_complete_answer_structure_issues(answer_text, prompt_text, messages or []))
    issues.extend(_complete_answer_sentence_support_issues(answer_text))

    word_window = _resolve_complete_answer_word_window(
        prompt_text,
        messages or [],
        enforce_long_response_split=enforce_long_response_split,
    )
    if word_window:
        min_words, max_words = word_window
        actual_words = _count_words(answer_text or "")
        if actual_words < min_words:
            issues.append(
                f"Answer is below the strict complete-answer word window ({actual_words} words; need at least {min_words})."
            )
        if actual_words > max_words:
            issues.append(
                f"Answer exceeds the strict complete-answer word window ({actual_words} words; cap is {max_words})."
            )

    missing_prompt_map = _missing_profile_prompt_map_asks(answer_text, prompt_text)
    if missing_prompt_map:
        issues.append(
            "Prompt-map asks appear under-covered: " + "; ".join(missing_prompt_map[:4])
        )

    # Keep issue ordering stable for deterministic tests.
    deduped: List[str] = []
    seen = set()
    for issue in issues:
        key = re.sub(r"\s+", " ", (issue or "").strip().lower())
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append((issue or "").strip())
    return deduped


def _direct_complete_answer_structure_issues(
    answer_text: str,
    prompt_text: str,
    messages: List[Dict[str, Any]],
) -> List[str]:
    """
    Apply a direct-answer scaffold check when the response is a fresh backend
    complete answer rather than a website-style continuation with anchored
    history state.
    """
    if _is_continuation_command(prompt_text or ""):
        return []

    state = _expected_unit_structure_state_from_history(prompt_text, messages or [])
    if state:
        return []

    if _is_problem_flow(prompt_text, messages or []):
        unit_kind = "problem"
    elif _is_essay_flow(prompt_text, messages or []) or _is_long_form_analysis_flow(prompt_text, messages or []):
        unit_kind = "essay"
    else:
        return []

    violates, reason = detect_unit_structure_policy_violation(
        answer_text,
        unit_kind=unit_kind,
        require_question_heading=False,
        expected_question_number=None,
        is_same_topic_continuation=False,
        expected_part_number=1,
        starts_new_question=True,
        enforce_single_top_level_part=False,
    )
    if violates and reason:
        return [f"Direct complete-answer structure violation: {reason}."]
    return []


def _build_complete_answer_rewrite_prompt(
    *,
    user_message: str,
    issues: List[str],
    output_mode: str,
    word_window: Optional[Tuple[int, int]],
) -> str:
    window_line = ""
    if word_window:
        window_line = (
            f"Keep the total answer inside the strict complete-answer window: "
            f"{int(word_window[0])}-{int(word_window[1])} words."
        )
    issue_lines = "\n".join(f"- {issue}" for issue in (issues or [])[:10])
    return "\n".join(
        [
            user_message.strip(),
            "",
            "[BACKEND STRICT COMPLETE-ANSWER REWRITE]",
            _backend_answer_output_instruction(output_mode),
            "Chat/API delivery does NOT relax the required Part-numbered answer scaffold.",
            "Regenerate the full answer from the beginning as one complete backend answer.",
            "Use indexed RAG and the shared backend code-guide instructions already active for this request.",
            "Ensure every argumentative sentence is individually supported by an immediate inline authority citation, or rewrite it more cautiously so it does not overstate support.",
            "Mirror every explicit question limb and prompt-map ask distinctly.",
            "Delete irrelevant, repetitive, or drifted material instead of padding around it.",
            "Keep the answer structurally final: no website split markers, no continuation markers, no partial-draft language.",
            window_line,
            "Fix every listed issue:",
            issue_lines,
        ]
    ).strip()


def send_complete_answer_with_docs(
    api_key: str,
    message: str,
    documents: List[Dict],
    project_id: str,
    history: List[Dict] = None,
    stream: bool = False,
    provider: str = "auto",
    model_name: Optional[str] = None,
    enforce_long_response_split: Optional[bool] = None,
    output_mode: Optional[str] = None,
    artifact_path: Optional[str] = None,
    cleanup_paths: Optional[List[Any]] = None,
    strict_complete_answer_verification: bool = True,
) -> Tuple[Any, Optional[str]]:
    """
    Canonical backend answer entrypoint.

    Defaults to direct/backend delivery:
    - mandatory legal RAG remains in the underlying provider path;
    - shared code-guide injection remains in the underlying provider path;
    - website multipart splitting is OFF by default here; and
    - complete-answer outputs get one deterministic backend verification/retry.
    """
    enforce_split = bool(enforce_long_response_split) if enforce_long_response_split is not None else False
    delivery_mode = resolve_backend_answer_delivery_mode(message, output_mode=output_mode)
    resolved_output_mode = delivery_mode
    history = history or []
    notes_request_active = bool(
        callable(detect_topic_notes_request)
        and detect_topic_notes_request(message or "").get("is_topic_notes")
    )

    if notes_request_active and callable(register_topic_notes_cleanup_paths) and cleanup_paths:
        register_topic_notes_cleanup_paths(project_id, cleanup_paths)

    if stream and delivery_mode in {
        BACKEND_ANSWER_OUTPUT_MARKDOWN_ARTIFACT,
        BACKEND_ANSWER_OUTPUT_DOCX_ARTIFACT,
    }:
        raise ValueError("Streaming is not supported when creating markdown or DOCX answer artifacts.")

    def _finalize_complete_answer_return(
        final_text: str,
        final_meta: Any,
        final_rag_context: Optional[str],
    ) -> Tuple[Any, Optional[str]]:
        artifact_info = write_complete_answer_artifact(
            final_text,
            delivery_mode=delivery_mode,
            prompt_text=message,
            project_id=project_id,
            artifact_path=artifact_path,
        )
        protected_cleanup_paths: List[Any] = []
        if artifact_info and artifact_info.get("path"):
            protected_cleanup_paths.append(artifact_info["path"])
        _cleanup_complete_answer_one_off_artifacts(
            cleanup_paths,
            protected_paths=protected_cleanup_paths,
        )
        final_meta = _attach_complete_answer_artifact_meta(final_meta, artifact_info)
        return (final_text, final_meta), final_rag_context

    response, rag_context = _provider_send_message_with_docs(
        api_key,
        message,
        documents,
        project_id,
        history=history,
        stream=stream,
        provider=provider,
        model_name=model_name,
        enforce_long_response_split=enforce_split,
    )

    if stream:
        return response, rag_context

    response_text = ""
    response_meta: Any = []
    if isinstance(response, tuple) and len(response) >= 1:
        response_text = response[0] or ""
        response_meta = response[1] if len(response) >= 2 else []
    else:
        response_text = response or ""

    if not strict_complete_answer_verification:
        return _finalize_complete_answer_return(response_text, response_meta, rag_context)

    if not _looks_like_backend_complete_answer_request(message, documents):
        return _finalize_complete_answer_return(response_text, response_meta, rag_context)

    low_message = (message or "").lower()
    if "[backend strict complete-answer rewrite]" in low_message:
        return _finalize_complete_answer_return(response_text, response_meta, rag_context)

    issues = _strict_complete_answer_issues(
        response_text,
        message,
        history,
        enforce_long_response_split=enforce_split,
    )
    if not issues:
        return _finalize_complete_answer_return(response_text, response_meta, rag_context)

    retry_text = response_text
    retry_meta = response_meta
    best_text = response_text
    best_meta = response_meta
    best_issues = issues
    word_window = _resolve_complete_answer_word_window(
        message,
        history,
        enforce_long_response_split=enforce_split,
    )

    for _attempt in range(BACKEND_COMPLETE_ANSWER_REGEN_MAX_ATTEMPTS):
        rewrite_message = _build_complete_answer_rewrite_prompt(
            user_message=message,
            issues=best_issues,
            output_mode=resolved_output_mode,
            word_window=word_window,
        )
        retry_response, rag_context = _provider_send_message_with_docs(
            api_key,
            rewrite_message,
            documents,
            project_id,
            history=history,
            stream=False,
            provider=provider,
            model_name=model_name,
            enforce_long_response_split=enforce_split,
        )
        if isinstance(retry_response, tuple) and len(retry_response) >= 1:
            retry_text = retry_response[0] or ""
            retry_meta = retry_response[1] if len(retry_response) >= 2 else []
        else:
            retry_text = retry_response or ""
            retry_meta = []

        retry_issues = _strict_complete_answer_issues(
            retry_text,
            message,
            history,
            enforce_long_response_split=enforce_split,
        )
        if len(retry_issues) < len(best_issues):
            best_text = retry_text
            best_meta = retry_meta
            best_issues = retry_issues
        if not retry_issues:
            break

    if word_window:
        best_text = _truncate_to_word_cap(best_text, int(word_window[1]), int(word_window[0]))
    return _finalize_complete_answer_return(best_text, best_meta, rag_context)


def send_message_with_docs(
    api_key: str,
    message: str,
    documents: List[Dict],
    project_id: str,
    history: List[Dict] = None,
    stream: bool = False,
    provider: str = "auto",
    model_name: Optional[str] = None,
    enforce_long_response_split: Optional[bool] = None,
    output_mode: Optional[str] = None,
    artifact_path: Optional[str] = None,
    cleanup_paths: Optional[List[Any]] = None,
    strict_complete_answer_verification: bool = True,
) -> Tuple[Any, Optional[str]]:
    """
    Backwards-compatible alias to the canonical backend direct-answer entrypoint.
    """
    return send_complete_answer_with_docs(
        api_key,
        message,
        documents,
        project_id,
        history=history,
        stream=stream,
        provider=provider,
        model_name=model_name,
        enforce_long_response_split=enforce_long_response_split,
        output_mode=output_mode,
        artifact_path=artifact_path,
        cleanup_paths=cleanup_paths,
        strict_complete_answer_verification=strict_complete_answer_verification,
    )


def send_complete_answer_with_output(
    api_key: str,
    message: str,
    documents: List[Dict],
    project_id: str,
    history: List[Dict] = None,
    stream: bool = False,
    provider: str = "auto",
    model_name: Optional[str] = None,
    enforce_long_response_split: Optional[bool] = None,
    output_mode: Optional[str] = None,
    artifact_path: Optional[str] = None,
    cleanup_paths: Optional[List[Any]] = None,
    strict_complete_answer_verification: bool = True,
) -> Tuple[Any, Optional[str], Optional[Dict[str, Any]]]:
    """
    Canonical delivery wrapper for backend complete answers.

    No matter whether the caller wants chat text, direct markdown-compatible
    text, a saved `.md`, or a Desktop `.docx`, generation still runs through
    `send_complete_answer_with_docs(...)` first so backend answer shaping, guide
    injection, and mandatory RAG stay identical.
    """
    delivery_mode = resolve_backend_answer_delivery_mode(message, output_mode=output_mode)
    response, rag_context = send_complete_answer_with_docs(
        api_key=api_key,
        message=message,
        documents=documents,
        project_id=project_id,
        history=history,
        stream=stream,
        provider=provider,
        model_name=model_name,
        enforce_long_response_split=enforce_long_response_split,
        output_mode=delivery_mode,
        artifact_path=artifact_path,
        cleanup_paths=cleanup_paths,
        strict_complete_answer_verification=strict_complete_answer_verification,
    )

    if stream:
        return response, rag_context, None

    response_text = ""
    response_meta: Any = []
    if isinstance(response, tuple) and len(response) >= 1:
        response_text = response[0] or ""
        response_meta = response[1] if len(response) >= 2 else []
    else:
        response_text = response or ""

    artifact_info = _extract_complete_answer_artifact_meta(response_meta)
    return (response_text, response_meta), rag_context, artifact_info


def send_message_with_output(
    api_key: str,
    message: str,
    documents: List[Dict],
    project_id: str,
    history: List[Dict] = None,
    stream: bool = False,
    provider: str = "auto",
    model_name: Optional[str] = None,
    enforce_long_response_split: Optional[bool] = None,
    output_mode: Optional[str] = None,
    artifact_path: Optional[str] = None,
    cleanup_paths: Optional[List[Any]] = None,
    strict_complete_answer_verification: bool = True,
) -> Tuple[Any, Optional[str], Optional[Dict[str, Any]]]:
    """
    Backwards-compatible alias to the canonical backend delivery wrapper.
    """
    return send_complete_answer_with_output(
        api_key=api_key,
        message=message,
        documents=documents,
        project_id=project_id,
        history=history,
        stream=stream,
        provider=provider,
        model_name=model_name,
        enforce_long_response_split=enforce_long_response_split,
        output_mode=output_mode,
        artifact_path=artifact_path,
        cleanup_paths=cleanup_paths,
        strict_complete_answer_verification=strict_complete_answer_verification,
    )

# RAG Service for document content retrieval
try:
    from rag_service import get_rag_service, RAGService
    RAG_AVAILABLE = True
except (ImportError, Exception) as e:
    print(f"RAG service not available: {e}")
    RAG_AVAILABLE = False

# Fixed rewrite policy (not user-toggleable):
# - hard-failure post-generation rewrites are globally capped to ONE retry
# - cap applies across citation/OSCOLA/quality/conclusion repair passes
AUTO_HARD_FAILURE_REGEN_MAX_ATTEMPTS = 2
AUTO_CITATION_REWRITE_MAX_ATTEMPTS = 1
# Intermediate-part underlength completion is a bounded continuation patch
# (not a full rewrite), kept to one attempt for latency control.
AUTO_INTERMEDIATE_UNDERLENGTH_FIX_MAX_ATTEMPTS = 2
# Retrieval/debug context visibility toggle.
ENABLE_RAG_DEBUG_UI = os.getenv("ENABLE_RAG_DEBUG_UI", "0").strip().lower() in {"1", "true", "yes"}
PROVIDER_SELECT_OPTIONS: List[str] = ["auto", *list(SUPPORTED_LLM_PROVIDERS)]
CUSTOM_MODEL_SELECT_LABEL = "Custom model..."

def _provider_option_label(option: str) -> str:
    normalized = (option or "").strip().lower()
    if normalized == "auto":
        return "Auto detect"
    return get_provider_display_name(normalized)

def _normalize_output_style(text: str) -> str:
    """
    Normalize formatting for consistency:
    - Remove decorative separator lines (e.g., repeated box-drawing characters).
    - Collapse multiple blank lines to a single blank line.
    """
    raw = (text or "").replace("\r\n", "\n").replace("\u00a0", " ")
    if not raw.strip():
        return raw

    sep_line = re.compile(r"^\s*[═─—\-_=]{8,}\s*$")
    lines = []
    for ln in raw.splitlines():
        if sep_line.match(ln):
            continue
        lines.append(ln.rstrip())

    normalized = "\n".join(lines).strip()
    # Collapse oversized vertical gaps, including whitespace-only blank lines.
    normalized = re.sub(r"(?:\n[^\S\n]*){3,}", "\n\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized

def _restore_paragraph_separation(text: str) -> str:
    """
    Recover readable paragraph structure when model output arrives as a single dense block.
    - Ensure line breaks before section headers like "Part I:", "Part II:", "ESSAY QUESTION:", etc.
    - Keep existing spacing when already well-formed.
    """
    t = (text or "").replace("\u00a0", " ").strip()
    if not t:
        return t
    # Insert breaks before common structural headers when they are glued inline.
    t = re.sub(r'(?<!\n)(\bPart\s+[IVXLC]+\s*:)', r'\n\n\1', t)
    t = re.sub(r'(?<!\n)(\b(?:ESSAY QUESTION|PROBLEM QUESTION)\s*:)', r'\n\n\1', t, flags=re.IGNORECASE)
    t = re.sub(r'(?<!\n)(\bPart\s+\d+\s*:)', r'\n\n\1', t, flags=re.IGNORECASE)
    # Universal heading-spacing rule: every structural heading gets exactly one blank line after it.
    t = re.sub(r'(?im)^((?:Part\s+[IVXLCDM0-9]+|Question\s+\d+|Title)\s*:[^\n]+)\n(?!\n)', r'\1\n\n', t)

    paragraphs = re.split(r"\n\s*\n", t)
    merged = []

    def _is_structural_heading(p: str) -> bool:
        s = (p or "").strip()
        return bool(re.match(r"(?im)^(?:Part\s+[IVXLCDM0-9]+|Question\s+\d+|Title)\s*:[^\n]+$", s))

    def _is_list_like(p: str) -> bool:
        s = (p or "").strip()
        return bool(re.match(r"(?im)^(?:[-*•]|\(?\d+\)?[.)]|[A-Z][.)])\s+", s))

    def _looks_incomplete(p: str) -> bool:
        s = (p or "").strip()
        if not s or _is_structural_heading(s) or _is_list_like(s):
            return False
        if re.search(r'[.!?]["\')\]]*\s*$', s):
            return False
        return True

    def _is_citation_only_block(p: str) -> bool:
        s = (p or "").strip()
        if not s or _is_structural_heading(s) or _is_list_like(s):
            return False
        if not s.startswith("("):
            return False
        inner = re.sub(r'^\(\s*', '', s)
        inner = re.sub(r'\s*\)[.;:]*\s*$', '', inner)
        if not inner:
            return False
        if len(re.findall(r"\b\w+\b", inner)) > 45:
            return False
        return bool(
            re.search(r"\bact\s+\d{4}\b", inner, flags=re.IGNORECASE)
            or re.search(r"\barticle\s+\d+(?:\(\d+\))?\b", inner, flags=re.IGNORECASE)
            or re.search(r"\b(?:section|s)\.?\s*\d+[a-z]?(?:\(\d+\))?\b", inner, flags=re.IGNORECASE)
            or re.search(r"\b[A-Z][A-Za-z0-9'’.\-]+(?:\s+[A-Z][A-Za-z0-9'’.\-]+){0,8}\s+v\.?\s+[A-Z][A-Za-z0-9'’.\-]+", inner)
            or re.search(r"\[[12]\d{3}\]", inner)
        )

    for para in paragraphs:
        p = (para or "").strip()
        if not p:
            continue
        if merged and _is_citation_only_block(p) and not _is_structural_heading(merged[-1]) and not _is_list_like(merged[-1]):
            merged[-1] = merged[-1].rstrip() + " " + p.lstrip()
        elif merged and _looks_incomplete(merged[-1]) and not _is_structural_heading(p) and not _is_list_like(p):
            merged[-1] = merged[-1].rstrip() + " " + p.lstrip()
        else:
            merged.append(p)
    t = "\n\n".join(merged)
    # Collapse oversized vertical gaps, including whitespace-only lines.
    t = re.sub(r'(?:\n[^\S\n]*){3,}', '\n\n', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    t = _split_oversized_paragraphs(t)
    return t.strip()


def _split_oversized_paragraphs(text: str, max_words: int = 180) -> str:
    """
    Split very large prose paragraphs at sentence boundaries to avoid dense,
    unreadable wall-of-text output.
    """
    raw = (text or "").strip()
    if not raw:
        return raw

    blocks = re.split(r"\n\s*\n", raw)
    rebuilt: List[str] = []

    def _is_structural(block: str) -> bool:
        s = (block or "").strip()
        if not s:
            return True
        return bool(
            re.match(r"(?im)^(?:Question\s+\d+|Part\s+[IVXLCDM0-9]+|Title)\s*:[^\n]+$", s)
            or re.match(r"(?im)^[A-D]\.\s*(?:Issue|Rule|Application|Conclusion)\b", s)
        )

    for block in blocks:
        b = (block or "").strip()
        if not b:
            continue
        if _is_structural(b):
            rebuilt.append(b)
            continue

        words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", b)
        if len(words) <= max_words + 25:
            rebuilt.append(b)
            continue

        sentences = re.split(r"(?<=[.!?])\s+", b)
        if len(sentences) < 3:
            rebuilt.append(b)
            continue

        current: List[str] = []
        current_words = 0
        split_blocks: List[str] = []
        for sent in sentences:
            s = (sent or "").strip()
            if not s:
                continue
            s_words = len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", s))
            if current and current_words >= max_words:
                split_blocks.append(" ".join(current).strip())
                current = [s]
                current_words = s_words
            else:
                current.append(s)
                current_words += s_words
        if current:
            split_blocks.append(" ".join(current).strip())

        rebuilt.extend(split_blocks if len(split_blocks) >= 2 else [b])

    return "\n\n".join(rebuilt).strip()


def _build_routing_debug(prompt_text: str) -> Dict[str, Any]:
    """
    Best-effort local preview of topic routing and planned subqueries for the UI debug panel.
    """
    prompt = (prompt_text or "").strip()
    if not prompt or _infer_retrieval_profile is None or _subissue_queries_for_unit is None:
        return {}
    try:
        profile = _infer_retrieval_profile(prompt) or {}
        subqueries = _subissue_queries_for_unit("Prompt", prompt) or []
        return {
            "topic": profile.get("topic") or "",
            "prompt_map_asks": list(profile.get("prompt_map_asks") or [])[:8],
            "subqueries": [label for label, _ in subqueries][:8],
        }
    except Exception:
        return {}

def _count_active_question_units(prompt_text: str) -> int:
    """
    Estimate how many substantive question units are in the active prompt.
    Used by title policy so multi-question answers never get a global Title line.
    """
    src = (prompt_text or "").strip()
    if not src:
        return 1
    try:
        plan = detect_long_essay(src)
        deliverables = plan.get("deliverables") or []
        qs = {
            int(q)
            for d in deliverables
            for q in (d.get("question_indices") or [d.get("question_index", 0)])
            if int(q or 0) > 0
        }
        if qs:
            return max(1, len(qs))
    except Exception:
        pass

    numbered_heads = re.findall(
        r"(?im)^\s*\d+\.\s+.*?\b(?:essay|problem)\s+question\b",
        src,
    )
    if numbered_heads:
        return len(numbered_heads)
    return 1

def _roman_to_int(roman: str) -> int:
    vals = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100}
    s = (roman or "").upper().strip()
    total = 0
    prev = 0
    for ch in reversed(s):
        v = vals.get(ch, 0)
        if v < prev:
            total -= v
        else:
            total += v
            prev = v
    return total

def _int_to_roman(num: int) -> str:
    if num <= 0:
        return "I"
    table = [
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")
    ]
    out = []
    n = num
    for val, sym in table:
        while n >= val:
            out.append(sym)
            n -= val
    return "".join(out) or "I"

def _enforce_part_numbered_conclusion_heading(text: str) -> str:
    """
    Ensure conclusion-like tail headings are Part-numbered.
    Example: "Conclusion and Advice" -> "Part VI: Conclusion and Advice"
    """
    raw = (text or "")
    if not raw.strip():
        return raw

    body = re.sub(r"\(End of Answer\)\s*$", "", raw, flags=re.IGNORECASE).rstrip()
    continuation_line = None
    m_cont = re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", body)
    if m_cont:
        continuation_line = "Will Continue to next part, say continue"
        body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).rstrip()

    # Find highest existing Part roman numeral.
    part_nums = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*.+$", body):
        part_nums.append(_roman_to_int(m.group(1)))
    max_part = max(part_nums) if part_nums else 0

    if max_part <= 0:
        # No Part structure present; do not force-convert.
        rebuilt = body
        if continuation_line:
            rebuilt = rebuilt + "\n\n" + continuation_line
        return rebuilt

    # Replace bare tail headings.
    heading_pat = re.compile(
        r"(?im)^\s*(Conclusion(?:\s+and\s+Advice)?|Advice(?:\s+and\s+Conclusion)?|Final\s+Advice)\s*$"
    )
    next_part = max_part + 1

    def repl(m):
        nonlocal next_part
        heading = m.group(1).strip()
        label = f"Part {_int_to_roman(next_part)}: {heading}"
        next_part += 1
        return label

    body = heading_pat.sub(repl, body)

    rebuilt = body.rstrip()
    if continuation_line:
        rebuilt = rebuilt + "\n\n" + continuation_line
    return rebuilt

def _next_part_conclusion_heading(text: str) -> str:
    """
    Build a conclusion heading that matches current Part numbering, if present.
    Falls back to a plain heading when no Part structure exists.
    """
    raw = (text or "")
    if not raw.strip():
        return "Conclusion and Advice"
    body = re.sub(r"\(End of Answer\)\s*$", "", raw, flags=re.IGNORECASE)
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    part_nums = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*.+$", body):
        part_nums.append(_roman_to_int(m.group(1)))
    if not part_nums:
        return "Conclusion and Advice"
    return f"Part {_int_to_roman(max(part_nums) + 1)}: Conclusion and Advice"

def _next_part_essay_conclusion_heading(text: str) -> str:
    """
    Build an essay conclusion heading that matches current Part numbering.
    Falls back to a plain 'Conclusion' heading when no Part structure exists.
    """
    raw = (text or "")
    if not raw.strip():
        return "Conclusion"
    body = re.sub(r"\(End of Answer\)\s*$", "", raw, flags=re.IGNORECASE)
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    part_nums = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*.+$", body):
        part_nums.append(_roman_to_int(m.group(1)))
    if not part_nums:
        return "Conclusion"
    return f"Part {_int_to_roman(max(part_nums) + 1)}: Conclusion"

def _next_part_problem_terminal_headings(text: str) -> Tuple[str, str]:
    raw = (text or "")
    body = re.sub(r"\(End of Answer\)\s*$", "", raw, flags=re.IGNORECASE)
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    part_nums: List[int] = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+|\d+)\s*:\s*.+$", body):
        token = (m.group(1) or "").strip()
        try:
            part_nums.append(int(token) if token.isdigit() else _roman_to_int(token))
        except Exception:
            continue
    next_num = (max(part_nums) + 1) if part_nums else 1
    return (
        f"Part {_int_to_roman(next_num)}: {PROBLEM_REMEDIES_LIABILITY_TITLE}",
        f"Part {_int_to_roman(next_num + 1)}: {PROBLEM_FINAL_CONCLUSION_TITLE}",
    )

def _ensure_problem_terminal_sections_within_cap(answer_text: str, max_words: int) -> str:
    body = (answer_text or "").strip()
    if not body:
        return body

    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    body = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)(?:Remedies\s+and\s+Liability|Liability\s+and\s+Remedies)\b",
        rf"\1{PROBLEM_REMEDIES_LIABILITY_TITLE}",
        body,
    )
    body = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)(?:Conclusion(?:\s+and\s+Advice)?|Overall\s+Conclusion(?:\s+and\s+Advice)?)\b",
        rf"\1{PROBLEM_FINAL_CONCLUSION_TITLE}",
        body,
    )
    body = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)Final\s+Advice\b",
        rf"\1{PROBLEM_FINAL_CONCLUSION_TITLE}",
        body,
    )

    remedies_re = re.compile(
        rf"(?im)^\s*Part\s+([IVXLCDM0-9]+)\s*:\s*{re.escape(PROBLEM_REMEDIES_LIABILITY_TITLE)}\b"
    )
    final_re = re.compile(
        rf"(?im)^\s*Part\s+([IVXLCDM0-9]+)\s*:\s*{re.escape(PROBLEM_FINAL_CONCLUSION_TITLE)}\b"
    )
    has_remedies = bool(remedies_re.search(body))
    final_matches = list(final_re.finditer(body))
    has_final = bool(final_matches)

    remedies_body = (
        "The remedial position should now be stated directly: identify the strongest liability outcome, "
        "the principal remedy or response available, and any practical limit on relief that follows from the analysis above."
    )
    final_body = (
        "Accordingly, the stronger answer is the one that identifies the most likely liability outcome, "
        "the key doctrinal reason for it, and the practical advice that follows for the parties. "
        "On that basis, the claimant's best route and the defendant's main exposure should be stated expressly."
    )

    if has_remedies and has_final:
        return _truncate_to_word_cap(body, max_words, 1) if max_words > 0 else body

    if (not has_remedies) and has_final:
        final_match = final_matches[-1]
        before = body[:final_match.start()].rstrip()
        after = body[final_match.end():].strip()
        remedies_heading, final_heading = _next_part_problem_terminal_headings(before)
        terminal_block = (
            f"{remedies_heading}\n\n{remedies_body}\n\n"
            f"{final_heading}\n\n{after or final_body}"
        )
        return _append_conclusion_within_cap(before, terminal_block, max_words)

    if has_remedies and (not has_final):
        _, final_heading = _next_part_problem_terminal_headings(body)
        terminal_block = f"{final_heading}\n\n{final_body}"
        return _append_conclusion_within_cap(body, terminal_block, max_words)

    remedies_heading, final_heading = _next_part_problem_terminal_headings(body)
    terminal_block = (
        f"{remedies_heading}\n\n{remedies_body}\n\n"
        f"{final_heading}\n\n{final_body}"
    )
    return _append_conclusion_within_cap(body, terminal_block, max_words)

def _enforce_end_of_answer(text: str) -> str:
    """
    Enforce a clean ending:
    - If the response is an intermediate multi-part output (ends with a 'Will Continue...' line),
      DO NOT include any '(End of Answer)' marker.
    - Otherwise, ensure EXACTLY ONE '(End of Answer)' at the end (remove any duplicates/legacy markers).
    """
    raw = _normalize_output_style(text).strip()
    if not raw:
        return "(End of Answer)"

    # If an explicit end marker already appears, discard everything after the first marker.
    # This prevents any leaked debug/context tail from surviving.
    first_end = re.search(r"\(End of Answer\)", raw, flags=re.IGNORECASE)
    if first_end:
        raw = raw[:first_end.start()].rstrip()

    # Never allow retrieval/debug dumps to leak into the main answer text.
    leak_markers = [
        "[RAG CONTEXT - INTERNAL - DO NOT OUTPUT]",
        "[END RAG CONTEXT]",
        "RETRIEVED LEGAL CONTEXT (from indexed documents)",
        "END OF RETRIEVED CONTEXT",
        "📚 RAG Retrieved Content (Debug)",
        "RAG Retrieved Content (Debug)",
        "Context Length:",
        "Allowed Authorities (preview):",
        "[ALL RETRIEVED DOCUMENTS]",
        "[END ALL RETRIEVED DOCUMENTS]",
        "No obvious primary authorities",
        "Removed 1 non-retrieved authority mention",
        "Removed 2 non-retrieved authority mention",
        "Removed 3 non-retrieved authority mention",
    ]
    leak_positions = [raw.find(m) for m in leak_markers if m in raw]
    if leak_positions:
        raw = raw[: min(leak_positions)].rstrip()
        if not raw:
            return "(End of Answer)"

    continue_patterns = [
        r"will\s+continue\s+to\s+next\s+part,\s*say\s+continue",
        r"will\s+continue\s+to\s+next\s+part",
        r"say\s+continue\s*$",
    ]
    has_continuation = any(re.search(p, raw, flags=re.IGNORECASE) for p in continue_patterns)
    has_end_marker = bool(re.search(r"\(End of Answer\)", raw, flags=re.IGNORECASE))

    # If BOTH "(End of Answer)" and "Will Continue" appear, the answer is COMPLETE.
    # The "Will Continue" is erroneous and must be stripped. "(End of Answer)" takes priority.
    if has_end_marker and has_continuation:
        has_continuation = False  # treat as final answer

    # Remove all end markers (including legacy ones) everywhere to prevent duplicates.
    cleaned = re.sub(r"\(End of Answer\)\s*", "", raw, flags=re.IGNORECASE)
    cleaned = re.sub(r"\(End of Essay\)\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\(End of Problem Question\)\s*", "", cleaned, flags=re.IGNORECASE)
    # Always strip erroneous "Will Continue" lines from final answers
    if not has_continuation:
        cleaned = re.sub(r"(?i)\n*will\s+continue\s+to\s+next\s+part.*$", "", cleaned, flags=re.MULTILINE)
    cleaned = cleaned.strip()

    if has_continuation:
        # If the model produced multiple "Will Continue..." lines, keep only one at the end.
        lines = [ln.rstrip() for ln in cleaned.splitlines() if ln.strip()]
        # Remove all existing continuation lines first, then append a single canonical line.
        lines = [ln for ln in lines if not re.search(r"will\s+continue\s+to\s+next\s+part", ln, flags=re.IGNORECASE)]
        lines.append("Will Continue to next part, say continue")
        return "\n\n".join(lines).strip()

    return cleaned + "\n\n(End of Answer)"

def _strip_generation_artifacts(text: str) -> str:
    """
    Remove debug/meta artefacts that sometimes leak from generation/rewrite steps.
    """
    t = (text or "")
    if not t.strip():
        return t
    # Hard-cut when known debug headers leak into answer body.
    cut_markers = [
        "📚 RAG Retrieved Content (Debug)",
        "Allowed Authorities (preview):",
        "[RAG CONTEXT - INTERNAL - DO NOT OUTPUT]",
        "[ALL RETRIEVED DOCUMENTS]",
    ]
    positions = [t.find(m) for m in cut_markers if m in t]
    if positions:
        t = t[:min(positions)]
    # Remove standalone debug lines that can survive partial cuts.
    t = re.sub(r'(?im)^\s*Removed\s+\d+\s+non-retrieved authority mention\(s\).*$', '', t)
    t = re.sub(r'(?im)^\s*Context Length:\s*\d+\s*characters\s*$', '', t)
    t = re.sub(r'(?im)^\s*No obvious primary authorities.*$', '', t)
    t = re.sub(r'\n{3,}', '\n\n', t).strip()
    return t

def _extract_word_targets(prompt_text: str) -> List[int]:
    """
    Extract explicit per-question word count targets from the user's prompt.

    Uses left-to-right order for multi-question prompts (Q1 count, Q2 count, etc.).
    """
    parsed = extract_word_targets_from_prompt(prompt_text or "", min_words=300)
    return [int(n) for n in (parsed.get("active_targets") or []) if int(n) >= 300]

def _is_continuation_command(text: str) -> bool:
    """
    Detect short continuation commands such as "continue", "part 2", etc.
    """
    normalized = re.sub(r"\s+", " ", (text or "").strip().lower())
    if not normalized:
        return False
    patterns = [
        r'^(?:please\s+)?continue(?:\s+now)?[.!]?$',
        r'^(?:please\s+)?next\s+part[.!]?$',
        r'^(?:please\s+)?go\s+on[.!]?$',
        r'^(?:please\s+)?keep\s+going[.!]?$',
        r'^(?:please\s+)?carry\s+on[.!]?$',
        r'^(?:please\s+)?proceed\s+now[.!]?$',
        r'^(?:please\s+)?part\s*\d+[.!]?$',
        r'^(?:please\s+)?continue\s+with\s+part\s*\d+[.!]?$',
        r'^(?:please\s+)?write\s+part\s*\d+[.!]?$',
        r'^(?:please\s+)?give\s+me\s+part\s*\d+[.!]?$',
        r'^(?:please\s+)?write\s+the\s+rest[.!]?$',
        r'^(?:please\s+)?finish\s+the\s+essay[.!]?$',
        r'^(?:please\s+)?complete\s+the\s+essay[.!]?$',
    ]
    return any(re.fullmatch(p, normalized) for p in patterns)

def _extract_authority_hints_from_prompt(prompt_text: str, limit: int = 40) -> List[str]:
    """
    Extract authority-like tokens from the user prompt so core authorities explicitly
    provided by the user are not removed by the strict RAG allow-list sanitizer.
    """
    text = (prompt_text or "")
    if not text:
        return []

    seen = set()
    out: List[str] = []

    def add(item: str):
        s = (item or "").strip()
        s = re.sub(r"^[\s,;:]+|[\s,;:.]+$", "", s)
        s = re.sub(
            r"^(?:apply|using|under|consider|evaluate|assess|analyse|analyze|discuss|critically\s+discuss|and)\s+",
            "",
            s,
            flags=re.IGNORECASE,
        )
        s = re.sub(r"\s+", " ", s).strip()
        s = re.sub(r"\s+(?:and|or|to|for|of|in|on|at)$", "", s, flags=re.IGNORECASE)
        if not s:
            return
        if len(s) > 180:
            return
        key = s.lower()
        if key in seen:
            return
        seen.add(key)
        out.append(s)

    # Statutes.
    for m in re.finditer(r"\b([A-Z][A-Za-z ,&()'-]+ Act \d{4})\b", text):
        add(m.group(1))

    # Treaty articles.
    for m in re.finditer(r"\b(Article\s+\d+(?:\(\d+\))?\s+T[FE]U)\b", text, flags=re.IGNORECASE):
        add(m.group(1))
    for m in re.finditer(r"\b(Article\s+\d+(?:\(\d+\))?\s+(?:ECHR|ICCPR))\b", text, flags=re.IGNORECASE):
        add(m.group(1))

    # Classic "X v Y" case names (with optional citation tail).
    party_token = r"(?:[A-Z][A-Za-z0-9&'’-]+|of|and|the|for|in|on|at|to|de|la|le|du|van|von|plc|ltd|llp|co|inc)"
    for m in re.finditer(
        rf"\b([A-Z][A-Za-z0-9&'’-]+(?:\s+{party_token}){{0,10}}\s+v\.?\s+[A-Z][A-Za-z0-9&'’-]+(?:\s+{party_token}){{0,10}}(?:\s*\[[12][0-9]{{3}}\][^\n.;)]{{0,80}})?)\b",
        text,
    ):
        add(m.group(1))

    # Short-form labels explicitly stated as "... <Name> case ..." in the prompt.
    # This helps preserve required aliases such as "Belmarsh case" through strict sanitization.
    stop_aliases = {
        "essay", "question", "problem", "part", "introduction", "conclusion",
        "international", "human", "rights", "law",
    }
    for m in re.finditer(r"\b([A-Z][A-Za-z-]{3,}(?:\s+[A-Z][A-Za-z-]{3,}){0,2})\s+case\b", text):
        alias = (m.group(1) or "").strip()
        if alias and alias.lower() not in stop_aliases:
            add(alias)

    # Preserve short-form authority names used in exam prompts:
    # "Bronner criteria", "Magill / IMS Health", "Post Danmark I", etc.
    for m in re.finditer(
        r"\b([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,2})\s+(?:criteria|test|doctrine|precedent|exception|framework)\b",
        text,
        flags=re.IGNORECASE,
    ):
        candidate = (m.group(1) or "").strip()
        if re.match(r"(?i)^(?:the|this|that|a|an)\b", candidate):
            continue
        add(candidate)
    for m in re.finditer(r"\b([A-Z][A-Za-z]+)\s*/\s*([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)?)\b", text):
        add(m.group(1))
        add(m.group(2))
    for m in re.finditer(
        r"\b([A-Z][A-Za-z'’.\-]+(?:\s+(?:of|and|the|for|in|on|at|to|de|la|le|du|van|von|[A-Z][A-Za-z'’.\-]+)){0,8})\s*\[(19|20)\d{2}\]\b",
        text,
    ):
        add(m.group(1))
    for k in [
        "Bronner", "Magill", "IMS Health", "Microsoft", "Deutsche Telekom",
        "TeliaSonera", "Post Danmark I", "Post Danmark", "Slovak Telekom",
        "AKZO", "Intel", "Google Shopping",
        "Infopaq", "Painer", "University of London Press",
        "Nova Productions", "Mazooma Games", "Thaler v Perlmutter",
        "Miller I", "Miller II", "Cherry v Advocate General for Scotland",
        "De Keyser", "Fire Brigades Union", "Case of Proclamations",
        "GCHQ", "Council of Civil Service Unions", "Burmah Oil",
    ]:
        if re.search(rf"\b{re.escape(k.lower())}\b", text.lower()):
            add(k)

    return out[:limit]

def _count_words(text: str) -> int:
    cleaned = text or ""
    cleaned = re.sub(r"(?im)^\s*(ESSAY|PROBLEM QUESTION|Q\d+)\s*:.*$", "", cleaned)
    cleaned = re.sub(r"(?im)^\s*[═=]{3,}\s*$", "", cleaned)
    cleaned = re.sub(r"\(End of Answer\)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", cleaned)
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9'’&./-]*", cleaned)
    return len(tokens)

def _assistant_message_counts_as_part(msg_text: str) -> bool:
    """
    Return True only for substantive assistant outputs that should advance
    long-response part numbering.
    """
    txt = (msg_text or "").strip()
    if not txt:
        return False

    low = txt.lower()
    non_part_markers = [
        "long response detected",
        "long multi-topic response detected",
        "type **'part 1'** or **'continue'** to begin",
        "ready to start?",
        "please respond with either",
        "retrieving sources",
        "thinking...",
    ]
    if any(m in low for m in non_part_markers):
        return False

    if re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", txt):
        return True
    if re.search(r"\(End of Answer\)", txt, flags=re.IGNORECASE):
        return True

    words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", txt)
    return len(words) >= 120

def _looks_like_pasted_generation(text: str) -> bool:
    """
    Detect pasted model output/debug transcripts so they are not used as
    long-response anchors.
    """
    t = (text or "").lower()
    if not t:
        return False
    markers = [
        "📚 rag retrieved content (debug)",
        "[rag context - internal - do not output]",
        "allowed authorities (preview):",
        "will continue to next part, say continue",
        "(end of answer)",
        "long multi-topic response detected",
        "word count plan:",
        "ready to start? type",
        "part 1",
    ]
    hits = sum(1 for m in markers if m in t)
    if hits >= 2:
        return True
    if len(t) > 2500 and hits >= 1:
        if ("part 1" in t and "continue" in t) or ("you\npart 1" in t) or ("part 2" in t and "will continue" in t):
            return True
    return False

def _has_visible_conclusion(answer_text: str) -> bool:
    """
    Heuristic check for a closing conclusion/advice in the tail of an answer.
    Used only as a quality guard for final parts.
    """
    text = (answer_text or "").strip()
    if not text:
        return False
    # Intermediate part marker means this is intentionally not final.
    if re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", text):
        return True

    body = re.sub(r"\(End of Answer\)\s*$", "", text, flags=re.IGNORECASE).strip()
    if not body:
        return False

    # Fast-path for an explicit conclusion heading, but only if it is the final
    # structural section rather than an earlier misplaced "Conclusion" followed
    # by later Parts.
    concl_matches = list(re.finditer(
        r"(?im)^\s*(?:Part\s+(?:[IVXLC]+|\d+)\s*:\s*)?(?:Conclusion(?:\s+and\s+Advice)?|Advice(?:\s+and\s+Conclusion)?|Final\s+Advice|Final\s+Conclusion)\s*$",
        body,
    ))
    if concl_matches:
        last_concl = concl_matches[-1]
        tail_after = body[last_concl.end():]
        if not re.search(r"(?im)^\s*Part\s+(?:[IVXLC]+|\d+)\s*:\s*", tail_after):
            return True

    # Check only the final structural block; broader tail scans can produce false
    # positives where an earlier conclusion heading is followed by later Parts.
    part_heads = list(re.finditer(r"(?im)^\s*Part\s+(?:[IVXLC]+|\d+)\s*:\s*", body))
    if part_heads:
        tail = body[part_heads[-1].start():].lower()
    else:
        tail = body[max(0, len(body) - 900):].lower()
    patterns = [
        r"\bin conclusion\b",
        r"\bto conclude\b",
        r"\bin summary\b",
        r"\bon balance\b",
        r"\bfinal outcome\b",
        r"\bfinal advice\b",
        r"\bfinal conclusion\b",
        r"\bconclusion and advice\b",
    ]
    return any(re.search(p, tail) for p in patterns)

def _append_conclusion_within_cap(answer_text: str, conclusion_block: str, max_words: int) -> str:
    """
    Append a required conclusion block while respecting max_words by trimming earlier body first.
    This avoids losing the conclusion to end-truncation.
    """
    body = (answer_text or "").strip()
    concl = (conclusion_block or "").strip()
    if not concl:
        return _truncate_to_word_cap(body, max_words, 1) if max_words > 0 else body
    if max_words <= 0:
        return (body + "\n\n" + concl).strip()

    # Remove terminal markers before recomposition; they are re-added later.
    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()

    concl_words = max(1, _count_words(concl))
    # Reserve at least the conclusion length (+ small cushion).
    body_cap = max(1, max_words - concl_words - 6)
    trimmed_body = _truncate_to_word_cap(body, body_cap, 1)
    combined = (trimmed_body.rstrip() + "\n\n" + concl).strip()

    if _count_words(combined) > max_words:
        overflow = _count_words(combined) - max_words
        tighter_cap = max(1, body_cap - overflow - 4)
        trimmed_body = _truncate_to_word_cap(body, tighter_cap, 1)
        combined = (trimmed_body.rstrip() + "\n\n" + concl).strip()

    if _count_words(combined) > max_words:
        combined = _truncate_to_word_cap(combined, max_words, 1)
    return combined

def _is_short_single_essay_prompt(prompt_text: str) -> bool:
    """
    True when prompt is a single essay request with explicit word target <= 2000.
    """
    targets = _extract_word_targets(prompt_text or "")
    if len(targets) != 1:
        return False
    if int(targets[0]) > 2000:
        return False
    low = (prompt_text or "").lower()
    if "problem question" in low:
        return False
    # Accept explicit essay prompts and "any topic" short-form prompts.
    if "essay" in low:
        return True
    return any(k in low for k in ["any topic", "any subject", "any area"])

def _is_short_single_problem_prompt(prompt_text: str) -> bool:
    """
    True when prompt is a single problem/advice request with explicit word target <= 2000.
    """
    targets = _extract_word_targets(prompt_text or "")
    if len(targets) != 1:
        return False
    if int(targets[0]) > 2000:
        return False
    low = (prompt_text or "").lower()
    if "essay" in low and "problem question" not in low and "advise" not in low:
        return False
    return ("problem question" in low) or ("advise" in low)

def _short_essay_effective_cap(target_words: int) -> int:
    """
    Short essays now target the user-requested ceiling directly.
    The min/max enforcement layer separately keeps output in the 99-100% window.
    """
    t = max(1, int(target_words or 0))
    return max(350, t)

def _normalize_short_essay_output(answer_text: str) -> str:
    """
    Normalize structure for short essays:
    - remove leaked debug/context tails
    - prevent late restart with a second "Part I"
    - keep one final conclusion section at the end
    """
    text = _strip_generation_artifacts(answer_text or "").strip()
    if not text:
        return text

    # Hard-cut anything after an explicit ending marker.
    m_end = re.search(r"\((?:End of Answer|End of Essay|End of Problem Question)\)", text, flags=re.IGNORECASE)
    if m_end:
        text = text[:m_end.start()].rstrip()

    # Remove continuation marker from short single-shot answers.
    text = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", text).strip()

    # If answer restarts with another Part I after already reaching later parts, cut the restart.
    part_i_hits = list(re.finditer(r"(?im)^\s*Part\s+I\s*:\s*", text))
    if len(part_i_hits) >= 2:
        second_i = part_i_hits[1].start()
        prefix = text[:second_i]
        if re.search(r"(?im)^\s*Part\s+(?:IV|V|VI|VII|VIII|IX|X)\s*:\s*", prefix):
            text = prefix.rstrip()

    # Normalize essay conclusion labels without forcing a hardcoded Part V.
    text = re.sub(
        r"(?im)^(\s*Part\s+[IVXLC]+\s*:)\s*Conclusion\s+and\s+Advice\b",
        r"\1 Conclusion",
        text,
    )

    plain_concl_match = re.search(r"(?im)^\s*Conclusion(?:\s+and\s+Advice)?\b", text)
    if plain_concl_match and not re.search(r"(?im)^\s*Part\s+[IVXLC]+\s*:\s*Conclusion\b", text):
        before = text[:plain_concl_match.start()].rstrip()
        after = text[plain_concl_match.end():].lstrip()
        heading = _next_part_essay_conclusion_heading(before)
        text = f"{before}\n\n{heading}\n\n{after}".strip()

    # If a conclusion exists, keep only the first conclusion block and drop any later re-started Parts.
    concl_match = re.search(r"(?im)^\s*Part\s+[IVXLC]+\s*:\s*Conclusion\b", text)
    if concl_match:
        tail = text[concl_match.end():]
        next_part = re.search(r"(?im)^\s*Part\s+[IVXLC]+\s*:\s*", tail)
        if next_part:
            text = text[:concl_match.end() + next_part.start()].rstrip()

    return text.strip()

def _normalize_short_problem_output(answer_text: str) -> str:
    """
    Normalize short single-shot problem answers:
    - remove leaked debug/context tails and placeholder numeric citations
    - strip misplaced generic conclusion blocks inserted mid-sentence
    - prevent sparse/jumping Part numbering such as Part I -> Part V
    """
    text = _strip_generation_artifacts(answer_text or "").strip()
    if not text:
        return text

    m_end = re.search(r"\((?:End of Answer|End of Essay|End of Problem Question)\)", text, flags=re.IGNORECASE)
    if m_end:
        text = text[:m_end.start()].rstrip()

    text = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", text).strip()
    text = re.sub(r"\(\s*\[\s*\d+\s*\]\s*\)", "", text)
    text = re.sub(r"\[\s*\d+\s*\]", "", text)
    text = re.sub(r"[^\S\n]{2,}", " ", text).strip()

    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    if blocks:
        cleaned_blocks: List[str] = []

        def _is_conclusion_heading_block(block: str) -> bool:
            return bool(
                re.match(
                    r"(?im)^(?:Part\s+[IVXLCDM0-9]+\s*:\s*)?(?:Conclusion(?:\s+and\s+Advice)?|Final\s+Conclusion)\b",
                    (block or "").strip(),
                )
            )

        def _is_generic_filler(block: str) -> bool:
            b = (block or "").strip().lower()
            return (
                b.startswith("on balance, the authorities and statutory framework support the integrated analysis set out above")
                or b.startswith("the outcome should turn on structured application of the governing legal tests")
            )

        def _is_substantive_follow_on(block: str) -> bool:
            b = (block or "").strip()
            if not b:
                return False
            return bool(
                re.match(r"(?im)^(?:[A-D]\.\s*(?:Issue|Rule|Application|Conclusion)\b|Part\s+[IVXLCDM0-9]+\s*:)", b)
                or b[:1].islower()
                or len(re.findall(r"\b\w+\b", b)) >= 6
            )

        i = 0
        while i < len(blocks):
            cur = blocks[i]
            prev = cleaned_blocks[-1] if cleaned_blocks else ""
            nxt = blocks[i + 1] if i + 1 < len(blocks) else ""
            nxt2 = blocks[i + 2] if i + 2 < len(blocks) else ""

            prev_incomplete = bool(
                prev
                and not re.search(r'[.!?]["\')\]]?\s*$', prev)
                and len(re.findall(r"\b\w+\b", prev)) >= 2
            )

            if _is_conclusion_heading_block(cur):
                later_analysis = _is_substantive_follow_on(nxt2) or (
                    _is_substantive_follow_on(nxt) and not _is_generic_filler(nxt)
                )
                if prev_incomplete or (_is_generic_filler(nxt) and later_analysis):
                    if _is_generic_filler(nxt):
                        i += 2
                    else:
                        i += 1
                    continue

            cleaned_blocks.append(cur)
            i += 1

        text = _restore_paragraph_separation("\n\n".join(cleaned_blocks))

    text = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)(?:Remedies\s+and\s+Liability|Liability\s+and\s+Remedies)\b",
        rf"\1{PROBLEM_REMEDIES_LIABILITY_TITLE}",
        text,
    )
    text = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)(?:Conclusion(?:\s+and\s+Advice)?|Overall\s+Conclusion(?:\s+and\s+Advice)?)\b",
        rf"\1{PROBLEM_FINAL_CONCLUSION_TITLE}",
        text,
    )
    text = re.sub(
        r"(?im)^(\s*Part\s+[IVXLCDM0-9]+\s*:\s*)Final\s+Advice\b",
        rf"\1{PROBLEM_FINAL_CONCLUSION_TITLE}",
        text,
    )

    plain_concl_match = re.search(
        r"(?im)^\s*(?:Conclusion(?:\s+and\s+Advice)?|Overall\s+Conclusion(?:\s+and\s+Advice)?|Final\s+Advice)\b",
        text,
    )
    if plain_concl_match and not re.search(
        rf"(?im)^\s*Part\s+[IVXLCDM0-9]+\s*:\s*{re.escape(PROBLEM_FINAL_CONCLUSION_TITLE)}\b",
        text,
    ):
        before = text[:plain_concl_match.start()].rstrip()
        after = text[plain_concl_match.end():].lstrip()
        _, final_heading = _next_part_problem_terminal_headings(before)
        text = f"{before}\n\n{final_heading}\n\n{after}".strip()

    final_concl_match = re.search(
        rf"(?im)^\s*Part\s+[IVXLCDM0-9]+\s*:\s*{re.escape(PROBLEM_FINAL_CONCLUSION_TITLE)}\b",
        text,
    )
    if final_concl_match:
        tail = text[final_concl_match.end():]
        next_part = re.search(r"(?im)^\s*Part\s+[IVXLCDM0-9]+\s*:\s*", tail)
        if next_part:
            text = text[:final_concl_match.end() + next_part.start()].rstrip()

    part_matches = list(re.finditer(r"(?im)^\s*Part\s+([IVXLCDM0-9]+)\s*:\s*(.+?)\s*$", text))
    if len(part_matches) >= 2:
        part_nums: List[int] = []
        for m in part_matches:
            token = (m.group(1) or "").strip()
            try:
                part_nums.append(int(token) if token.isdigit() else _roman_to_int(token))
            except Exception:
                part_nums.append(0)
        if any(
            (part_nums[idx] <= 0)
            or (idx > 0 and (part_nums[idx] <= part_nums[idx - 1] or part_nums[idx] > part_nums[idx - 1] + 1))
            for idx in range(len(part_nums))
        ):
            seq = {"n": 0}
            def _repl(m: re.Match[str]) -> str:
                seq["n"] += 1
                return f"Part {_int_to_roman(seq['n'])}: {(m.group(2) or '').strip()}"
            text = re.sub(
                r"(?im)^\s*Part\s+([IVXLCDM0-9]+)\s*:\s*(.+?)\s*$",
                _repl,
                text,
            )

    return text.strip()

def _trim_regressive_part_restart(
    answer_text: str,
    prompt_text: str,
    messages: List[Dict[str, Any]],
) -> str:
    """
    For continuation parts in single-question long flows, trim accidental restarted tails
    (for example Part V -> Part II -> Part I).
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    state = _expected_part_state_from_history(prompt_text, messages)
    if not state or int(state.get("current_part") or 1) <= 1:
        return txt

    def _trim_continuation_edge_overlap(current_text: str) -> str:
        prev_text = ""
        for msg in reversed(messages or []):
            if msg.get("role") != "assistant":
                continue
            prev_text = (msg.get("text") or "").strip()
            if prev_text:
                break
        if not prev_text:
            return current_text

        def _strip_markers(raw: str) -> str:
            out = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", raw or "").strip()
            out = re.sub(r"\(End of Answer\)\s*$", "", out, flags=re.IGNORECASE).strip()
            return out

        def _split_heading_prefix(raw: str) -> Tuple[str, str]:
            s = (raw or "").lstrip()
            pos = 0
            heading_re = re.compile(r"(?im)^(?:Title|Question\s+\d+|Part\s+[IVXLCDM0-9]+)\s*:[^\n]+$")
            while True:
                m = heading_re.match(s, pos)
                if not m:
                    break
                pos = m.end()
                while pos < len(s) and s[pos] in " \t\r\n":
                    pos += 1
            return s[:pos], s[pos:]

        cur_full = _strip_markers(current_text)
        prev_full = _strip_markers(prev_text)
        prefix, cur_body = _split_heading_prefix(cur_full)
        if not cur_body or not prev_full:
            return current_text

        prev_tokens = list(re.finditer(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", prev_full))
        cur_tokens = list(re.finditer(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", cur_body))
        if len(prev_tokens) < 10 or len(cur_tokens) < 10:
            return current_text

        prev_words = [m.group(0).lower() for m in prev_tokens]
        cur_words = [m.group(0).lower() for m in cur_tokens]
        max_overlap = min(45, len(prev_words), len(cur_words))
        trim_chars = 0
        for k in range(max_overlap, 9, -1):
            if prev_words[-k:] == cur_words[:k]:
                trim_chars = cur_tokens[k - 1].end()
                break
        if trim_chars <= 0:
            return current_text

        trimmed_body = cur_body[trim_chars:].lstrip(" \t\r\n,;:.-")
        if not trimmed_body:
            return current_text
        rebuilt = (prefix + trimmed_body).strip()
        if re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", current_text):
            rebuilt = rebuilt.rstrip() + "\n\nWill Continue to next part, say continue"
        elif re.search(r"\(End of Answer\)\s*$", current_text, flags=re.IGNORECASE):
            rebuilt = rebuilt.rstrip() + "\n\n(End of Answer)"
        return rebuilt

    txt = _trim_continuation_edge_overlap(txt)

    # Resolve latest anchor request and skip multi-unit/by-section flows where restarts may be intentional.
    anchor_text = ""
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        u = (msg.get("text") or "").strip()
        if not u or _looks_like_pasted_generation(u):
            continue
        if _extract_word_targets(u):
            anchor_text = u
            break
    if not anchor_text:
        return txt

    plan = detect_long_essay(anchor_text)
    split_mode = (plan.get("split_mode") or "").strip().lower()
    if split_mode in {"by_units", "by_section"}:
        return txt

    # If answer explicitly starts new section types, avoid destructive trimming.
    if re.search(r"(?im)^\s*(ESSAY QUESTION|PROBLEM QUESTION)\b", txt):
        return txt

    body = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()

    part_hits = list(re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", body))
    if len(part_hits) < 2:
        return txt

    cut_pos = None
    prev_num = _roman_to_int(part_hits[0].group(1))
    seen_ge_ii = prev_num >= 2
    for m in part_hits[1:]:
        num = _roman_to_int(m.group(1))
        if num >= 2:
            seen_ge_ii = True
        if num < prev_num:
            cut_pos = m.start()
            break
        prev_num = num

    if cut_pos is None and seen_ge_ii:
        part_i_positions = [m.start() for m in part_hits if _roman_to_int(m.group(1)) == 1]
        if len(part_i_positions) >= 2:
            cut_pos = part_i_positions[1]

    if cut_pos is None:
        return txt

    trimmed = body[:cut_pos].rstrip()
    return trimmed if trimmed else txt

def _is_essay_prompt(prompt_text: str) -> bool:
    """
    Detect essay-style prompts (excluding problem questions).
    """
    low = (prompt_text or "").lower()
    return ("essay" in low) and ("problem question" not in low)

def _is_essay_flow(prompt_text: str, messages: List[Dict[str, Any]]) -> bool:
    """
    Detect essay flow even during continuation turns where prompt may be just "continue".
    """
    if _is_essay_prompt(prompt_text):
        return True
    for msg in reversed(messages or []):
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt:
            continue
        if _looks_like_pasted_generation(txt):
            continue
        if _is_essay_prompt(txt) and _extract_word_targets(txt):
            return True
    return False

def _is_problem_flow(prompt_text: str, messages: List[Dict[str, Any]]) -> bool:
    """
    Detect problem-question/advice flow for conclusion style control.
    """
    low = (prompt_text or "").lower()
    if "problem question" in low:
        return True
    if ("advise" in low) and ("essay question" not in low):
        return True
    for msg in reversed(messages or []):
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        t = txt.lower()
        if "problem question" in t:
            return True
        if ("advise" in t) and ("essay question" not in t):
            return True
    return False

def _is_long_form_analysis_flow(prompt_text: str, messages: List[Dict[str, Any]]) -> bool:
    """
    Detect non-problem long-form doctrinal analysis prompts (e.g., "Explain X under Y law").
    This captures quality-rewrite candidates that are not explicitly labelled as "essay".
    """
    if _is_essay_flow(prompt_text, messages):
        return True
    if _is_problem_flow(prompt_text, messages):
        return False

    indicators = ["explain", "analyse", "analyze", "discuss", "evaluate", "assess", "critically"]
    low = (prompt_text or "").lower()
    targets = _extract_word_targets(prompt_text or "")
    if any(k in low for k in indicators):
        if not targets:
            return True
        if max(targets) >= 1200:
            return True

    for msg in reversed(messages or []):
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        t = txt.lower()
        t_targets = _extract_word_targets(txt)
        if any(k in t for k in indicators):
            if not t_targets:
                return True
            if max(t_targets) >= 1200:
                return True
    return False

def _extract_conclusion_section_text(answer_text: str) -> str:
    """
    Extract conclusion section body text (best-effort) from an essay response.
    """
    body = (answer_text or "").strip()
    if not body:
        return ""
    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()
    m = re.search(
        r"(?im)^\s*(?:Part\s+(?:[IVXLC]+|\d+)\s*:\s*)?(?:Conclusion(?:\s+and\s+Advice)?|Final\s+Conclusion)\b.*$",
        body,
    )
    if not m:
        # Tail fallback for implicit conclusions
        return body[max(0, len(body) - 700):].strip()
    tail = body[m.end():]
    nxt = re.search(r"(?im)^\s*Part\s+(?:[IVXLC]+|\d+)\s*:\s*", tail)
    if nxt:
        return tail[:nxt.start()].strip()
    return tail.strip()

def _extract_introduction_section_text(answer_text: str) -> str:
    """
    Extract Part I: Introduction body text (best-effort) from a structured answer.
    """
    body = (answer_text or "").strip()
    if not body:
        return ""
    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()
    m = re.search(r"(?im)^\s*Part\s+I\s*:\s*Introduction\b.*$", body)
    if not m:
        return ""
    tail = body[m.end():]
    nxt = re.search(r"(?im)^\s*Part\s+(?:[IVXLC]+|\d+)\s*:\s*", tail)
    if nxt:
        return tail[:nxt.start()].strip()
    return tail.strip()

def _extract_explicit_issue_requirements(prompt_text: str) -> List[Dict[str, Any]]:
    """
    Extract explicit issue asks (for example a)/b)/c)/d) and bullet requirements)
    from the prompt, returning lightweight keyword sets for coverage checks.
    """
    txt = (prompt_text or "")
    if not txt:
        return []

    lines = [ln.strip(" \t•-") for ln in txt.splitlines() if ln.strip()]
    req_lines: List[str] = []

    # Primary: explicit a)/b)/c)/d) asks.
    for ln in lines:
        if re.match(r"(?i)^[a-d]\)\s+", ln):
            req_lines.append(re.sub(r"(?i)^[a-d]\)\s+", "", ln).strip())

    # Secondary: requirement bullets under "You are required to" / "Your answer should".
    if not req_lines:
        collect = False
        for raw in txt.splitlines():
            s = raw.strip()
            low = s.lower()
            if ("you are required to" in low) or ("your answer should" in low) or ("your answer should explore" in low) or ("your answer should consider" in low):
                collect = True
                continue
            if collect:
                if not s:
                    continue
                if re.match(r"(?i)^(part\s+\d+|question\s*\d+|factual scenario|fact pattern)\b", s):
                    collect = False
                    continue
                if s.startswith(("•", "-", "*")):
                    req_lines.append(s.lstrip("•-* ").strip())
                elif re.match(r"(?i)^[a-d]\)\s+", s):
                    req_lines.append(re.sub(r"(?i)^[a-d]\)\s+", "", s).strip())
                elif len(req_lines) >= 1:
                    # Stop when the bullet/requirement block ends.
                    collect = False

    if not req_lines:
        return []

    stop = {
        "the", "and", "for", "with", "from", "that", "this", "which", "where", "when",
        "your", "answer", "should", "consider", "explore", "analyse", "analyze", "assess",
        "discuss", "critically", "evaluate", "whether", "can", "be", "is", "are", "of", "to",
        "law", "legal", "international", "state", "states",
    }
    out: List[Dict[str, Any]] = []
    seen = set()
    for raw in req_lines:
        key = raw.lower().strip()
        if not key or key in seen:
            continue
        seen.add(key)
        toks = re.findall(r"[a-z][a-z0-9\-]{3,}", key)
        kws: List[str] = []
        for t in toks:
            if t in stop:
                continue
            if t not in kws:
                kws.append(t)
            if len(kws) >= 5:
                break
        if kws:
            out.append({"text": raw, "keywords": kws})
    return out

def _strip_draft_continuation_opener(answer_text: str) -> str:
    """
    Remove draft/meta continuation openers from user-visible text.
    Example: "Continuing the analysis of ...".
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    body = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()

    parts = re.split(r"\n\s*\n", body)
    def _is_meta_opener(block: str) -> bool:
        b = (block or "").strip().lower()
        b = re.sub(r"^[\-\u2022\*\s]+", "", b).strip()
        return (
            b.startswith("continuing the analysis")
            or b.startswith("continuing the examination")
            or b.startswith("building directly upon")
            or b.startswith("transitioning from")
            or b.startswith("to continue")
            or b.startswith("in continuation")
        )

    if parts:
        first = (parts[0] or "").strip().lower()
        starts_meta = _is_meta_opener(first)
        if starts_meta and not re.match(r"(?i)^part\s+[ivxlc\d]+\s*:", first):
            body = "\n\n".join(parts[1:]).strip()
            parts = re.split(r"\n\s*\n", body)

    # Secondary case: heading first, then a continuation-style draft opener.
    if len(parts) >= 2:
        first_block = (parts[0] or "").strip()
        second_block = (parts[1] or "").strip()
        if re.match(r"(?im)^part\s+[ivxlc\d]+\s*:\s*", first_block) and _is_meta_opener(second_block):
            remaining = parts[2:]
            if remaining:
                body = "\n\n".join([parts[0]] + remaining).strip()
            # If no remaining parts after removal, keep the meta-opener rather than leaving an empty heading
    return body


def _strip_answer_wrappers(answer_text: str) -> str:
    """
    Remove legacy multi-question wrappers like "Answer 1:" when the user-facing
    heading should simply be "Question 1: ...".
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    txt = re.sub(
        r"(?im)^\s*Answer\s+\d+\s*:\s*(?=Question\s+\d+\s*:)",
        "",
        txt,
    )
    txt = re.sub(
        r"(?im)^\s*Answer\s+\d+\s*:\s*$\n(?=\s*Question\s+\d+\s*:)",
        "",
        txt,
    )
    return txt.strip()


def _normalize_question_heading_part_resets(answer_text: str) -> str:
    """
    When a response contains explicit "Question N: ..." wrappers, keep the
    question->Part linkage clean and ensure each question block can restart its
    own Part sequence cleanly when needed.
    """
    txt = _strip_answer_wrappers(answer_text or "")
    if not txt or not re.search(r"(?im)^\s*Question\s+\d+\s*:\s*", txt):
        return txt

    # Remove stray carry-over Part headings that appear immediately before a fresh
    # "Question N:" block with no intervening body text.
    txt = re.sub(
        r"(?im)^\s*Part\s+[IVXLC\d]+\s*:\s*[^\n]+(?:\n\s*){1,3}(?=Question\s+\d+\s*:)",
        "",
        txt,
    ).strip()

    q_re = re.compile(r"(?im)^\s*(Question\s+\d+\s*:\s*.+?)\s*$")
    matches = list(q_re.finditer(txt))
    if not matches:
        return txt

    rebuilt: List[str] = []
    cursor = 0
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(txt)
        prefix = txt[cursor:start]
        if prefix:
            rebuilt.append(prefix)

        heading = (m.group(1) or "").strip()
        block_body = txt[m.end():end].strip()
        if not block_body:
            if "".join(rebuilt).strip():
                rebuilt.append("\n\n")
            rebuilt.append(heading)
            cursor = end
            continue

        first_part = re.search(r"(?im)^\s*Part\s+([IVXLC]+|\d+)\s*:\s*(.+?)\s*$", block_body)
        if not first_part:
            block_body = f"Part I: Introduction\n\n{block_body}".strip()

        if "".join(rebuilt).strip():
            rebuilt.append("\n\n")
        rebuilt.append(f"{heading}\n\n{block_body}")
        cursor = end

    rebuilt.append(txt[cursor:])
    normalized = "".join(rebuilt).strip()
    normalized = re.sub(r"(?s)(\S)\n(?=Question\s+\d+\s*:)", r"\1\n\n", normalized)
    normalized = re.sub(r"\n{3,}(?=Question\s+\d+\s*:)", "\n\n", normalized)
    return normalized


def _empty_heading_structure_issues(answer_text: str) -> List[str]:
    """
    Detect empty structural stubs that should trigger a rewrite:
    - bare lettered headings like "A."
    - titled lettered headings with no body beneath them
    - Part headings with no content at all beneath them
    """
    txt = (answer_text or "").strip()
    if not txt:
        return []

    lines = txt.splitlines()
    non_empty: List[Tuple[int, str]] = [
        (idx, (line or "").strip())
        for idx, line in enumerate(lines)
        if (line or "").strip()
    ]
    if not non_empty:
        return []

    def _is_non_substantive_line(line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return True
        return bool(
            re.match(r"(?im)^will\s+continue\s+to\s+next\s+part,\s*say\s+continue\s*$", s)
            or re.match(r"(?im)^\(end\s+of\s+answer\)\s*$", s)
            or re.match(r"(?im)^📚\s*rag\s+retrieved\s+content", s)
            or re.match(r"(?im)^\[rag\s+context", s)
        )

    def _heading_kind(line: str) -> Optional[str]:
        s = (line or "").strip()
        if re.match(r"(?i)^question\s+\d+\s*:", s):
            return "question"
        if re.match(r"(?i)^part\s+[ivxlcdm\d]+\s*:", s):
            return "part"
        if re.match(r"^[A-D]\.\s*", s):
            return "letter"
        return None

    def _is_title_like_line(line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return False
        if re.match(r"(?i)^title\s*:", s):
            return True
        if (
            "\n" not in s
            and (
                re.search(r"(?i)\b(?:essay|problem)\s+question\b", s)
                or re.search(r"(?i)\(\s*words?\s*\)\s*$", s)
            )
            and not re.match(r"(?i)^question\s+\d+\s*:", s)
        ):
            return True
        return False

    headings: List[Tuple[int, str, str]] = []
    for idx, line in non_empty:
        kind = _heading_kind(line)
        if not kind:
            continue
        title = ""
        if kind == "part":
            title = re.sub(r"(?i)^part\s+[ivxlcdm\d]+\s*:\s*", "", line).strip()
        elif kind == "letter":
            title = re.sub(r"^[A-D]\.\s*", "", line).strip()
        else:
            title = line
        headings.append((idx, kind, title))

    issues: List[str] = []
    for i, (line_idx, kind, title) in enumerate(headings):
        next_idx = headings[i + 1][0] if (i + 1) < len(headings) else len(lines)
        next_kind = headings[i + 1][1] if (i + 1) < len(headings) else None
        between = [(ln or "").strip() for ln in lines[line_idx + 1:next_idx] if (ln or "").strip()]

        prev_nonempty = ""
        for prev_i in range(line_idx - 1, -1, -1):
            prev_line = (lines[prev_i] or "").strip()
            if prev_line:
                prev_nonempty = prev_line
                break

        if kind == "question":
            continue

        if kind == "letter" and not title:
            issues.append("Contains a bare lettered subheading like 'A.' with no title or analysis.")
            continue

        non_heading_body = [
            ln for ln in between
            if _heading_kind(ln) is None and not _is_non_substantive_line(ln)
        ]

        if kind == "letter" and not non_heading_body:
            issues.append("Contains an empty lettered subsection heading with no substantive content beneath it.")
        elif kind == "part" and not non_heading_body and next_kind in {None, "part", "question"}:
            issues.append("Contains a Part heading with no substantive content beneath it.")

        if (
            prev_nonempty
            and _heading_kind(prev_nonempty) is None
            and not _is_title_like_line(prev_nonempty)
            and not _is_non_substantive_line(prev_nonempty)
            and not re.search(r'[.!?:]["\')\]]?\s*$', prev_nonempty)
            and len(re.findall(r"\b\w+\b", prev_nonempty)) >= 2
        ):
            issues.append("Text appears cut off immediately before a new structural heading.")
            break

    return issues

def _ensure_clean_terminal_sentence(answer_text: str, is_intermediate: bool) -> str:
    """
    Ensure the body ends with a complete sentence.
    Deterministic fallback prevents truncated tails like "This rigid framework".
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt
    if not _is_abrupt_answer_ending(txt):
        return txt

    body = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    if not body:
        return txt

    truncated_terminal_word = bool(
        re.search(r"(?i)\b(and|or|to|of|for|with|under|against|between|by|from|the|a|an)\s*[.!?](?:[\"')\]]+)?\s*$", body)
    )

    # Prefer clipping to the latest complete sentence to avoid injecting
    # fallback text directly after a truncated fragment.
    sentence_end_re = re.compile(r'[.!?](?:["\')\]]+)?(?=\s|$)')
    ends = [m.end() for m in sentence_end_re.finditer(body)]
    sentence_spans = list(sentence_end_re.finditer(body))
    if sentence_spans:
        last_sentence = body[sentence_spans[-2].end():sentence_spans[-1].end()].strip() if len(sentence_spans) >= 2 else body[:sentence_spans[-1].end()].strip()
        if len(sentence_spans) >= 2 and len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", last_sentence)) <= 3:
            ends = ends[:-1]
    if truncated_terminal_word and ends:
        ends = ends[:-1]
    if ends:
        clip_at = ends[-1]
        # Keep only if the retained body is still substantive.
        clipped = body[:clip_at].strip()
        if _count_words(clipped) >= 8:
            return clipped

    if truncated_terminal_word:
        while re.search(
            r"(?i)\b(and|or|to|of|for|with|under|against|between|by|from|the|a|an)\s*(?:[.!?](?:[\"')\]]+)?)?\s*$",
            body,
        ):
            new_body = re.sub(
                r"(?i)\b\w+\s+(?:and|or|to|of|for|with|under|against|between|by|from|the|a|an)\s*(?:[.!?](?:[\"')\]]+)?)?\s*$",
                "",
                body,
            ).rstrip(" ,;:-.")
            if new_body == body:
                break
            body = new_body
        if body and not re.search(r'[.!?](?:["\')\]]+)?\s*$', body):
            body = re.sub(r"\b\w+\s*$", "", body).rstrip(" ,;:-.")

    fallback = (
        "Further analysis follows in the next part."
        if is_intermediate else
        "Accordingly, this is the final advised position on the issues raised."
    )
    # Place fallback as a separate paragraph for readability.
    return f"{body.rstrip()}.\n\n{fallback}".strip()

def _enforce_expected_part_heading(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> str:
    """
    Enforce continuation stability:
    - remove draft continuation opener,
    - ensure continuation parts begin with the expected Part heading.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    txt = _strip_draft_continuation_opener(txt)
    txt = _strip_answer_wrappers(txt)
    txt = _normalize_question_heading_part_resets(txt)
    txt = _enforce_expected_question_heading(txt, prompt_text, messages)
    state = _expected_part_state_from_history(prompt_text, messages)
    unit_state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if not state:
        return txt

    expected_part = int(state.get("current_part") or 1)
    expected_internal_part = _expected_internal_part_heading_from_history(prompt_text, messages)
    expected_heading_num = (
        int(unit_state.get("expected_part_number") or 0)
        if unit_state and int(unit_state.get("expected_part_number") or 0) > 0
        else (expected_internal_part if (expected_internal_part and expected_internal_part > 0) else expected_part)
    )
    if expected_heading_num <= 1:
        # Fresh long-response part must begin at Part I deterministically.
        if re.search(r"(?im)^\s*(ESSAY QUESTION|PROBLEM QUESTION)\b", txt):
            return txt
        expect_intro_heading = bool((not unit_state) or unit_state.get("starts_new_question"))
        m_head_i = re.search(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt)
        if m_head_i:
            actual_i = _roman_to_int(m_head_i.group(1))
            if actual_i != 1:
                txt = re.sub(
                    r"(?im)^\s*Part\s+[IVXLC]+\s*:",
                    "Part I:",
                    txt,
                    count=1,
                ).strip()
            if expect_intro_heading:
                txt = re.sub(
                    r"(?im)^(\s*Part\s+I\s*:\s*).*$",
                    r"\1Introduction",
                    txt,
                    count=1,
                ).strip()
            return txt
        # If the model omitted a heading, add a deterministic Part I heading.
        if expect_intro_heading:
            return f"Part I: Introduction\n\n{txt}".strip()
        return f"Part I: Further Analysis\n\n{txt}".strip()

    # Keep explicit question wrappers untouched.
    if re.search(r"(?im)^\s*(ESSAY QUESTION|PROBLEM QUESTION)\b", txt):
        return txt

    m_head = re.search(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt)
    if m_head:
        # Drop/trim leaked draft bridge prose before the first Part heading.
        if m_head.start() > 0:
            prefix = (txt[:m_head.start()] or "").strip()
            q_only_prefix = False
            if unit_state and unit_state.get("require_question_heading"):
                prefix_wo_q = re.sub(r"(?im)^\s*Question\s+\d+\s*:\s*.+\s*$", "", prefix).strip()
                q_only_prefix = bool(prefix) and (not prefix_wo_q)
            if re.match(
                r"(?is)^(?:[-\u2022*]\s*)?(?:continuing|building|transitioning|in continuation)\b",
                prefix,
            ):
                txt = txt[m_head.start():].lstrip()
                m_head = re.search(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt)
                if not m_head:
                    return txt
            elif prefix and (not q_only_prefix):
                stripped_prefix = prefix.strip()
                title_like_prefix = bool(
                    stripped_prefix
                    and ("\n" not in stripped_prefix)
                    and (
                        re.search(r"(?i)\b(?:essay|problem)\s+question\b", stripped_prefix)
                        or re.search(r"(?i)\(\s*words?\s*\)\s*$", stripped_prefix)
                    )
                    and (not re.match(r"(?i)^question\s+\d+\s*:", stripped_prefix))
                    and (not re.match(r"(?i)^title\s*:", stripped_prefix))
                )
                if title_like_prefix:
                    txt = txt[m_head.start():].lstrip()
                    m_head = re.search(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt)
                    if not m_head:
                        return txt
                    prefix = ""
                # Continuation/body prose must not appear before the first Part heading.
                # Preserve the prose, but force the expected heading to the top of the response.
                if prefix:
                    forced_lead = f"Part {_int_to_roman(expected_heading_num)}: Further Analysis"
                    if state.get("is_final"):
                        forced_lead = f"Part {_int_to_roman(expected_heading_num)}: Final Analysis"
                    txt = f"{forced_lead}\n\n{txt}".strip()
                    m_head = re.search(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt)
                    if not m_head:
                        return txt
        actual = _roman_to_int(m_head.group(1))
        # Force the first heading to the expected in-answer part number when it drifts
        # (for example restarts at Part I or jumps to Part X).
        if actual != expected_heading_num:
            txt = re.sub(
                r"(?im)^\s*Part\s+[IVXLC]+\s*:",
                f"Part {_int_to_roman(expected_heading_num)}:",
                txt,
                count=1,
            ).strip()

        # Continuation parts must never use an "Introduction" or generic
        # "Continued Analysis" heading. Relabel deterministically.
        first_heading = re.search(r"(?im)^\s*Part\s+[IVXLC]+\s*:\s*(.+?)\s*$", txt)
        if first_heading and expected_heading_num > 1:
            heading_tail = (first_heading.group(1) or "").strip().lower()
            allow_intro_here = bool(unit_state and unit_state.get("starts_new_question"))
            if heading_tail.startswith("continued analysis") or (
                heading_tail.startswith("introduction") and not allow_intro_here
            ):
                prompt_low = (prompt_text or "").lower()
                is_problem_like = bool(
                    re.search(r"(?i)\bproblem question\b|\badvise\b", prompt_low)
                )
                if state.get("is_final"):
                    relabel = "Conclusion and Advice" if is_problem_like else "Conclusion"
                else:
                    relabel = "Further Analysis"
                txt = re.sub(
                    r"(?im)^(\s*Part\s+[IVXLC]+\s*:\s*).*$",
                    rf"\1{relabel}",
                    txt,
                    count=1,
                ).strip()
        return txt

    # No heading present: add deterministic heading keyed to part state.
    if unit_state and unit_state.get("starts_new_question"):
        tail_label = "Introduction"
    else:
        tail_label = "Final Analysis and Conclusion" if state.get("is_final") else "Further Analysis"
    heading = f"Part {_int_to_roman(expected_heading_num)}: {tail_label}"
    return f"{heading}\n\n{txt}".strip()


def _renumber_internal_part_sequence(answer_text: str) -> str:
    """
    Renumber multiple Part headings sequentially inside a single-question
    response block to avoid duplicates such as Part II ... Part II.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt
    if len(re.findall(r"(?im)^\s*Question\s+\d+\s*:\s*.+$", txt)) > 1:
        return txt

    matches = list(re.finditer(r"(?im)^(\s*Part\s+)([IVXLCDM]+|\d+)(\s*:\s*.+?)$", txt))
    if len(matches) < 2:
        return txt

    first_token = (matches[0].group(2) or "").strip()
    try:
        current_num = int(first_token) if first_token.isdigit() else _roman_to_int(first_token)
    except Exception:
        return txt
    if current_num <= 0:
        return txt

    rebuilt: List[str] = []
    last_idx = 0
    next_num = current_num
    for m in matches:
        rebuilt.append(txt[last_idx:m.start()])
        rebuilt.append(f"{m.group(1)}{_int_to_roman(next_num)}{m.group(3)}")
        last_idx = m.end()
        next_num += 1
    rebuilt.append(txt[last_idx:])
    return "".join(rebuilt).strip()

def _has_substantial_duplicate_blocks(answer_text: str) -> bool:
    """
    Detect substantial repeated long blocks without mutating output text.
    Used only for quality diagnostics.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return False

    blocks = [b.strip() for b in re.split(r"\n\s*\n", txt) if (b or "").strip()]
    long_sigs: List[str] = []

    for b in blocks:
        if re.match(r"(?im)^(title\s*:|part\s+[ivxlc\d]+\s*:|essay question|problem question)\b", b):
            continue
        words = re.findall(r"\b\w+\b", b.lower())
        if len(words) < 18:
            continue
        sig = re.sub(r"[^a-z0-9\s]", " ", b.lower())
        sig = re.sub(r"\s+", " ", sig).strip()
        long_sigs.append(sig)

    if len(long_sigs) < 3:
        return False

    dup_hits = 0
    for i in range(len(long_sigs)):
        for j in range(i + 1, len(long_sigs)):
            a = long_sigs[i]
            b = long_sigs[j]
            if a == b:
                dup_hits += 1
                continue
            if len(a) > 120 and len(b) > 120:
                common = len(set(a.split()) & set(b.split()))
                denom = max(1, min(len(set(a.split())), len(set(b.split()))))
                if (common / denom) >= 0.88:
                    dup_hits += 1
                    continue
                if SequenceMatcher(None, a[:1200], b[:1200]).ratio() >= 0.90:
                    dup_hits += 1

    # Signal only when repetition is material, not incidental.
    return dup_hits >= 2

def _collapse_adjacent_duplicate_headings(answer_text: str) -> str:
    """
    Collapse exact adjacent duplicate structural heading lines, e.g.
    "Part II: X" followed immediately by the same "Part II: X".
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    lines = txt.splitlines()
    rebuilt: List[str] = []
    i = 0
    heading_re = re.compile(
        r"(?i)^\s*(?:question\s+\d+\s*:|part\s+[ivxlcdm\d]+\s*:|[A-D]\.\s+).+$"
    )
    while i < len(lines):
        current = lines[i]
        current_stripped = current.strip()
        if current_stripped and heading_re.match(current_stripped):
            j = i + 1
            blank_run: List[str] = []
            while j < len(lines) and not lines[j].strip():
                blank_run.append(lines[j])
                j += 1
            if j < len(lines) and lines[j].strip() == current_stripped:
                rebuilt.append(current)
                rebuilt.append("")  # preserve one blank line after the kept heading
                i = j + 1
                while i < len(lines) and not lines[i].strip():
                    i += 1
                continue
        rebuilt.append(current)
        i += 1
    return "\n".join(rebuilt).strip()

def _strip_empty_headings(answer_text: str) -> str:
    """
    Remove Part/lettered headings that have no substantive content beneath them.
    This prevents bare headings from appearing in final output.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    heading_re = re.compile(r"(?im)^(?:Part\s+[IVXLCDM0-9]+\s*:|[A-D]\.\s*\S)")
    blocks = re.split(r"\n\s*\n", txt)
    if len(blocks) <= 1:
        return txt

    result: List[str] = []
    i = 0
    while i < len(blocks):
        block = (blocks[i] or "").strip()
        if not block:
            i += 1
            continue
        # Check if block is a standalone heading (single line, matches heading pattern)
        block_lines = [ln for ln in block.splitlines() if ln.strip()]
        if len(block_lines) == 1 and heading_re.match(block_lines[0].strip()):
            # Check if the next block is also a heading or end-of-text
            next_i = i + 1
            while next_i < len(blocks) and not blocks[next_i].strip():
                next_i += 1
            if next_i >= len(blocks):
                # Heading at the very end with no content after it — drop it
                i += 1
                continue
            next_block = blocks[next_i].strip()
            next_lines = [ln for ln in next_block.splitlines() if ln.strip()]
            if len(next_lines) >= 1 and heading_re.match(next_lines[0].strip()):
                # Next block is also a heading — this heading has no body, drop it
                i += 1
                continue
        result.append(block)
        i += 1

    return "\n\n".join(result).strip()

def _final_output_integrity_cleanup(answer_text: str) -> str:
    """
    Last deterministic cleanup to keep output presentable and citation-safe:
    - remove leaked debug tails
    - remove empty/orphaned parentheticals like "()" and "( )"
    - unwrap accidental double-parenthetical wrappers
    - trim dangling lead-ins ("as noted in .", "in (YEAR ...)")
    - remove empty Part/lettered headings with no content
    """
    txt = _strip_generation_artifacts(answer_text or "").strip()
    if not txt:
        return txt

    txt = _strip_draft_continuation_opener(txt)
    txt = _strip_answer_wrappers(txt)
    txt = _normalize_question_heading_part_resets(txt)
    txt = _collapse_adjacent_duplicate_headings(txt)
    txt = _strip_empty_headings(txt)

    # Drop obvious draft scaffolding if it leaks into the final response.
    txt = re.sub(r"(?im)^\s*(?:draft|first draft|rough draft|working draft)\s*:?\s*$", "", txt)
    txt = re.sub(r"(?im)^\s*(?:draft answer to rewrite|internal draft)\s*:?\s*$", "", txt)

    # Empty / punctuation-only parentheticals.
    txt = re.sub(r"\(\s*\)", "", txt)
    txt = re.sub(r"\(\s*[,.;:]+\s*\)", "", txt)
    # Remove obvious non-OSCOLA parenthetical labels.
    txt = re.sub(r"\(\s*source\s*\d+[^)]*\)", "", txt, flags=re.IGNORECASE)
    txt = re.sub(r"\(\s*[^)]*law\s+trove[^)]*\)", "", txt, flags=re.IGNORECASE)
    txt = re.sub(r"\(\s*[^)]*\.pdf[^)]*\)", "", txt, flags=re.IGNORECASE)
    txt = re.sub(r"\(\s*key\s+case[^)]*\)", "", txt, flags=re.IGNORECASE)

    def _looks_like_oscola_reference(s: str) -> bool:
        low = (s or "").lower()
        if not low:
            return False
        if re.search(r"\bact\s+\d{4}\b", low):
            return True
        if re.search(r"\barticle\s+\d+(?:\(\d+\))?\b", low):
            return True
        if re.search(r"\b(?:s|section)\.?\s*\d+[a-z]?(?:\(\d+\))?\b", low):
            return True
        if (" v " in low or " v. " in low) and (
            re.search(r"\[[12]\d{3}\]", s)
            or re.search(r"\((?:19|20)\d{2}\)", s)
            or re.search(r"\b(?:AC|QB|Ch|WLR|All\s+ER|Lloyd'?s\s+Rep|UKSC|EWCA|EWHC)\b", s)
        ):
            return True
        return False

    def _is_pinpoint_tail(s: str) -> bool:
        return bool(re.match(r"(?i)^\s*(?:\[\d+\]|at\s+|para(?:graph)?s?\.?\s*\d+|p(?:p)?\.?\s*\d+|n\.?\s*\d+|s\.?\s*\d+[a-z]?(?:\(\d+\))?)", s or ""))

    def _normalize_parenthetical(m: re.Match) -> str:
        inner = (m.group(1) or "").strip()
        if not inner:
            return ""
        low = inner.lower()
        if ".pdf" in low or "law trove" in low:
            return ""

        # Strip stray footnote-like numerals fused to a case name, e.g.
        # "(Fairchild v Glenhaven Funeral Services Ltd10)".
        if (
            re.search(r"\bv\.?\b", inner, flags=re.IGNORECASE)
            and not re.search(r"\[[12]\d{3}\]|\((?:19|20)\d{2}\)|\b(?:AC|QB|Ch|WLR|All\s+ER|UKSC|EWCA|EWHC)\b", inner)
        ):
            inner = re.sub(r"(?<=[A-Za-z])\d{1,3}\s*$", "", inner).strip()

        # Trim narrative tails accidentally inserted inside citation brackets.
        if _looks_like_oscola_reference(inner):
            if "," in inner:
                head, tail = inner.split(",", 1)
                if _looks_like_oscola_reference(head) and (not _is_pinpoint_tail(tail)):
                    inner = head.strip()
            inner = re.sub(
                r"(?i)\s+(?:by which|to which|whereby|which|who|because|as\s+noted|as\s+stated|as\s+held|for\s+this\s+reason)\b.*$",
                "",
                inner,
            ).strip(" ,;:")
        if not inner:
            return ""
        return f"({inner})"

    txt = re.sub(r"\(([^()\n]{1,320})\)", _normalize_parenthetical, txt)

    # Unwrap accidental double wrappers: "((...))" -> "(...)"
    txt = re.sub(r"\(\s*\(([^()]{1,260})\)\s*\)", r"(\1)", txt)

    # Remove report-only citations without case name, but do not strip the report part
    # from a full in-text citation that already includes "X v Y" immediately before it.
    report_only_patterns = [
        (
            r"\(\s*\(?\s*"
            r"(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s*\)?\s*"
            r"(?:\d{1,4}\s+)?(?:LR\s+\d{1,4}\s+)?"
            r"(?:AC|QB|Ch|WLR|All\s+ER|Lloyd'?s\s+Rep|ECR|Bus\s+LR|ECC|KB|P)\b"
            r"[^()]{0,60}\)"
        ),
        r"\(\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s+[A-Z][A-Za-z. ]{1,30}\s+\d{1,4}(?:\s+[A-Z]{1,4}\s+\d{1,4})?\s*\)",
        r"\(\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s+\d{1,4}\s+[A-Z]{1,8}(?:\s+[A-Z]{1,8}){0,2}\s+\d{1,5}\s*\)",
    ]
    case_name_nearby = r"\b[A-Z][A-Za-z0-9'’.\-]+(?:\s+[A-Z][A-Za-z0-9'’.\-]+){0,6}\s+v\.?\s+[A-Z][A-Za-z0-9'’.\-]+"

    def _drop_report_only(m: re.Match) -> str:
        lookback = txt[max(0, m.start() - 120):m.start()]
        if re.search(case_name_nearby, lookback):
            return m.group(0)
        return ""

    for _pat in report_only_patterns:
        txt = re.sub(_pat, _drop_report_only, txt, flags=re.IGNORECASE)

    # Remove orphaned report-year parentheticals when introduced by dangling prepositions.
    txt = re.sub(
        r"(?i)\b(?:in|under|see|cf|per|as\s+noted\s+in|as\s+held\s+in|as\s+stated\s+in|as\s+established\s+in)\s+"
        r"\(\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})[^()]{0,120}\)",
        "",
        txt,
    )

    # Generic dangling citation lead-ins.
    txt = re.sub(r"(?i)\b(?:in|under|see|cf|per|from)\s+\.(?=\s|$)", "", txt)
    txt = re.sub(r"(?i)\b(?:as\s+noted\s+in|as\s+held\s+in|as\s+stated\s+in|as\s+established\s+in)\s+\.", "", txt)
    txt = re.sub(r"(?i)\b(?:in|under|see|cf|per|from)\s*,\s*", "", txt)
    txt = re.sub(r"(?i)\b(?:in|under|see|cf|per|from|as\s+noted\s+in|as\s+held\s+in|as\s+stated\s+in)\s*$", "", txt)

    # Second pass after citation removal to clear emptied wrappers.
    txt = re.sub(r"\(\s*\)", "", txt)
    txt = re.sub(r"\(\s*[,.;:]+\s*\)", "", txt)
    txt = re.sub(r"(?m)^\s*[.,;:]\s*$", "", txt)
    txt = re.sub(r"\s+([,.;:])", r"\1", txt)
    txt = re.sub(r"(?:\n[^\S\n]*){3,}", "\n\n", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    txt = _collapse_adjacent_duplicate_headings(txt)
    txt = _restore_paragraph_separation(txt).strip()
    return txt.strip()

def _problem_structure_issues(answer_text: str) -> List[str]:
    """
    Detect high-impact structural failures in problem-question outputs:
    - Part numbering regression / restart
    - Conclusion part appearing before later analytical parts
    - Missing A/B/C/D IRAC sub-headings in analytical parts
    - Standalone IRAC labels outside A/B/C/D pattern
    """
    issues: List[str] = []
    txt = (answer_text or "").strip()
    if not txt:
        return ["Empty answer text."]

    body = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    part_matches: List[Tuple[int, int, str]] = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*(.+?)\s*$", body):
        try:
            part_matches.append((m.start(), _roman_to_int(m.group(1)), (m.group(2) or "").strip().lower()))
        except Exception:
            continue

    if len(part_matches) >= 2:
        for idx in range(1, len(part_matches)):
            if part_matches[idx][1] <= part_matches[idx - 1][1]:
                issues.append("Part numbering regresses or repeats inside the same answer.")
                break

    conclusion_part_idx = -1
    for i, (_start, _num, heading_low) in enumerate(part_matches):
        if "conclusion" in heading_low:
            conclusion_part_idx = i
            break
    if conclusion_part_idx >= 0 and conclusion_part_idx != len(part_matches) - 1:
        issues.append("Conclusion appears before later Parts (structure ordering failure).")

    def _heading_keywords(s: str) -> List[str]:
        stop = {
            "the", "and", "of", "for", "to", "in", "on", "with", "a", "an", "by",
            "continued", "analysis", "introduction", "conclusion", "advice", "part",
        }
        toks = re.findall(r"[a-z]{4,}", s or "")
        return [t for t in toks if t not in stop]

    similar_heading_detected = False
    for i in range(1, len(part_matches)):
        prev_heading = part_matches[i - 1][2]
        curr_heading = part_matches[i][2]
        if ("conclusion" in prev_heading) or ("conclusion" in curr_heading):
            continue
        prev_kws = set(_heading_keywords(prev_heading))
        curr_kws = set(_heading_keywords(curr_heading))
        if not prev_kws or not curr_kws:
            continue
        overlap = len(prev_kws & curr_kws)
        denom = min(len(prev_kws), len(curr_kws))
        if denom and (overlap / denom) >= 0.67:
            similar_heading_detected = True
            break
    if similar_heading_detected:
        issues.append("Adjacent Parts substantially duplicate the same issue heading instead of progressing the analysis.")

    irac_pat = re.compile(r"(?im)^\s*([A-D])\.\s*(Issue|Rule|Application|Conclusion)\b")
    standalone_irac_pat = re.compile(r"(?im)^\s*(Issue|Rule|Application|Conclusion)\s*$")
    malformed_irac_detected = False
    standalone_detected = False
    open_final_part_detected = False
    has_any_irac = False

    for i, (start, _num, heading_low) in enumerate(part_matches):
        end = part_matches[i + 1][0] if (i + 1) < len(part_matches) else len(body)
        block = body[start:end]
        if "conclusion" in heading_low:
            continue

        labels = set()
        for m in irac_pat.finditer(block):
            labels.add((m.group(2) or "").strip().lower())
        if labels:
            has_any_irac = True
        if labels and len(labels) < 3:
            malformed_irac_detected = True

        if i == len(part_matches) - 1:
            if ("application" in labels) and ("conclusion" not in labels):
                open_final_part_detected = True

        if standalone_irac_pat.search(block):
            standalone_detected = True

    if malformed_irac_detected:
        issues.append("Problem-part IRAC subheadings are malformed; if A/B/C/D markers are used, they should form a complete Issue/Rule/Application/Conclusion sequence.")
    if standalone_detected:
        issues.append("Standalone IRAC headings found outside A/B/C/D pattern.")
    if open_final_part_detected:
        issues.append("Final analytical Part is left open after A/B/C without a closing D. Conclusion.")
    if (not has_any_irac) and _count_words(body) >= 180:
        bare_major_heading_detected = False
        for block in [b.strip() for b in re.split(r"\n\s*\n", body) if (b or "").strip()]:
            if "\n" in block:
                continue
            if re.match(r"(?i)^(question\s+\d+\s*:|part\s+[ivxlcdm0-9]+\s*:|title\s*:|\(end of answer\)|will continue\b)", block):
                continue
            if block.endswith((".", ":", ";", "?", "!")):
                continue
            words = re.findall(r"[A-Za-z][A-Za-z'’/-]*", block)
            if 1 <= len(words) <= 8:
                bare_major_heading_detected = True
                break
        if bare_major_heading_detected:
            issues.append("Problem-part IRAC subheadings are missing from analytical sections that need visible Issue/Rule/Application/Conclusion structure.")

    return issues


def _major_unheaded_section_issues(answer_text: str, *, is_problem_mode: bool = False) -> List[str]:
    """
    Detect long answers where major sections appear as bare standalone headings
    instead of Part-numbered sections.
    """
    issues: List[str] = []
    txt = (answer_text or "").strip()
    if not txt or _count_words(txt) < 220:
        return issues

    body = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    part_count = len(re.findall(r"(?im)^\s*Part\s+([IVXLCDM]+|\d+)\s*:\s*", body))

    candidate_headings: List[str] = []
    generic_bare_headings: List[str] = []
    for block in [b.strip() for b in re.split(r"\n\s*\n", body) if (b or "").strip()]:
        if "\n" in block:
            continue
        if re.match(r"(?i)^(question\s+\d+\s*:|part\s+[ivxlcdm0-9]+\s*:|[A-D]\.\s+|title\s*:|\(end of answer\)|will continue\b)", block):
            continue
        if block.endswith((".", ":", ";", "?", "!")):
            continue
        words = re.findall(r"[A-Za-z][A-Za-z'’/-]*", block)
        if not (2 <= len(words) <= 12):
            continue
        uppercase_initials = sum(1 for w in words if w[:1].isupper())
        if uppercase_initials < max(2, len(words) - 2):
            continue
        candidate_headings.append(block)
        if re.fullmatch(r"(?i)(introduction|conclusion|analysis|evaluation|application|remedies)", block):
            generic_bare_headings.append(block)

    if generic_bare_headings:
        issues.append("Contains bare generic headings outside the required Part structure.")
    if part_count <= 1 and candidate_headings:
        issues.append("Major sections appear as bare headings instead of Part-numbered sections.")
    elif len(candidate_headings) >= 3 and part_count <= 2:
        issues.append("Too many major standalone headings appear without matching Part-numbered sections.")

    return issues

def _essay_quality_issues(
    answer_text: str,
    prompt_text: str,
    is_short_single_essay: bool,
    is_problem_mode: bool = False,
) -> List[str]:
    """
    Validate essay structure/fluency for final outputs.
    Returns human-readable issue list.
    """
    issues: List[str] = []
    txt = (answer_text or "").strip()
    if not txt:
        return ["Empty answer text."]

    # Hard structural defects.
    if _is_abrupt_answer_ending(txt):
        issues.append("Answer ends abruptly or with an incomplete final sentence.")
    if not _has_visible_conclusion(txt):
        issues.append("No visible concluding section.")
    if re.search(r"(?im)^\s*(continuing the analysis|building directly upon|in continuation)\b", txt):
        issues.append("Contains draft continuation-style opening instead of final structured heading.")
    if _is_continuation_command(prompt_text or ""):
        for _m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*(.+?)\s*$", txt):
            try:
                _pn = _roman_to_int(_m.group(1))
            except Exception:
                _pn = 0
            _tail = (_m.group(2) or "").strip().lower()
            if _pn > 1 and _tail.startswith("introduction"):
                issues.append("Continuation output restarts with an 'Introduction' heading instead of progressing analysis.")
                break
    if re.search(
        r"(?is)^\s*Part\s+[IVXLC]+\s*:\s*Continued Analysis\s*\n\s*\n\s*Question\s+\d+\s*:",
        txt,
    ):
        issues.append("Uses 'Continued Analysis' immediately before a fresh question heading; a new question must start directly at its own heading.")
    if re.search(r"(?im)^\s*Answer\s+\d+\s*:\s*", txt):
        issues.append("Contains legacy 'Answer N:' wrapper; use only 'Question N:' headings.")
    if re.search(
        r"(?ims)^\s*(Question\s+\d+\s*:|Part\s+[IVXLC\d]+\s*:|[A-D]\.\s+.+?)\s*$"
        r"(?:\n\s*)+^\s*\1\s*$",
        txt,
    ):
        issues.append("Contains an immediately repeated structural heading line.")
    q_heading_re = re.compile(r"(?im)^\s*Question\s+\d+\s*:\s*.+?$")
    for _qm in q_heading_re.finditer(txt):
        _tail = txt[_qm.end():]
        _next_nonempty = next((ln.strip() for ln in _tail.splitlines() if ln.strip()), "")
        if _next_nonempty and not re.match(r"(?i)^part\s+([ivxlcdm]+|\d+)\s*:", _next_nonempty):
            issues.append("A question heading is not immediately followed by a Part heading.")
            break

    # Placeholder / malformed citation artefacts.
    if re.search(r"\(\s*[A-Za-z]\s*\)", txt):
        issues.append("Contains placeholder citation markers like '(J )'.")
    if re.search(r"\(\s*\[\s*\d+\s*\]\s*\)|\[\s*\d+\s*\]", txt):
        issues.append("Contains placeholder numeric citation markers instead of usable legal authorities.")
    if re.search(r"\(\s*\)", txt):
        issues.append("Contains empty parenthetical markers '()' created by citation stripping.")
    if re.search(r"\(\s*\(", txt):
        issues.append("Contains malformed double-parenthetical citation fragments.")
    if re.search(r"\(\s*(?:TEU|TFEU|Treaty on European Union|Treaty on the Functioning of the European Union)\s*\)", txt):
        issues.append("Contains bare treaty citations such as '(TFEU)' instead of a specific article reference.")
    if re.search(r"\(\s*C-\d+/\d+(?:\s*;\s*C-\d+/\d+)+\s*\)", txt, flags=re.IGNORECASE):
        issues.append("Contains EU case-number parentheticals without full case names or proper OSCOLA formatting.")
    if re.search(r"\([^)\n]*\.pdf[^)\n]*\)", txt, flags=re.IGNORECASE):
        issues.append("Contains file-name citations (for example '.pdf') instead of OSCOLA authorities.")
    if re.search(r"\([^)\n]*law\s+trove[^)\n]*\)", txt, flags=re.IGNORECASE):
        issues.append("Contains database-label citations (for example 'Law Trove') instead of OSCOLA authorities.")
    if re.search(
        r"(?im)^(?:\(\s*(?:[A-Z][^()\n]{2,90}\s+v\.?\s+[A-Z][^()\n]{2,90}|[^()\n]{0,40}\b(?:Act\s+\d{4}|Article\s+\d+|section\s+\d+|s\.?\s*\d+)[^()\n]{0,40})\s*\)[.;:]?)$",
        txt,
    ):
        issues.append("Contains a citation-only paragraph detached from the proposition sentence it should support.")
    if re.search(
        r"\([^)\n]*(?:\[[12]\d{3}\]|(?:\bact\s+\d{4}\b)|\bv\.?\b)[^)\n]*,\s*(?:by|which|where|because|as\s+noted|as\s+stated|as\s+held)\b[^)\n]*\)",
        txt,
        flags=re.IGNORECASE,
    ):
        issues.append("Contains malformed citation parentheticals with narrative text; OSCOLA brackets must contain citation-only content.")
    if re.search(
        r"\([A-Z][^()\n]{3,120}\bv\.?\b[^()\n]{3,120}[A-Za-z]\d{1,3}\)",
        txt,
    ):
        issues.append("Contains malformed case parentheticals with fused footnote-style digits instead of a proper OSCOLA citation.")
    if re.search(
        r"\([A-Za-z][^()\n]{0,50}\)\s+Act\s+\d{4}\)",
        txt,
    ):
        issues.append("Contains malformed statutory parentheticals with broken Act titles instead of a proper OSCOLA statute citation.")
    if re.search(r"\b[A-Z][^()\n]{2,120}\.\.\.[^()\n]{0,120}\[[12]\d{3}\]", txt):
        issues.append("Contains an incomplete case citation with ellipsis instead of a full OSCOLA case name.")
    if re.search(
        r"\([A-Z][^()\n]{3,160}\bv\.?\b[^()\n]{3,160}\)",
        txt,
    ):
        malformed_case_parens = []
        for _m in re.finditer(r"\(([A-Z][^()\n]{3,160}\bv\.?\b[^()\n]{3,160})\)", txt):
            inner = (_m.group(1) or "").strip()
            if not re.search(r"\[[12]\d{3}\]|\((?:19|20)\d{2}\)|\b(?:AC|QB|Ch|WLR|All\s+ER|UKSC|EWCA|EWHC|ICJ|HL)\b", inner):
                malformed_case_parens.append(inner)
        if malformed_case_parens:
            issues.append("Contains case citations in parentheses without a proper OSCOLA report/year reference.")

    report_only_patterns = [
        r"\(\s*\(?\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s*\)?\s*(?:\d{1,4}\s+)?(?:LR\s+\d{1,4}\s+)?(?:AC|QB|Ch|WLR|All\s+ER|Lloyd'?s\s+Rep|ECR|Bus\s+LR|ECC|KB|P)\b[^()]{0,60}\)",
        r"\(\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s+[A-Z][A-Za-z. ]{1,30}\s+\d{1,4}(?:\s+[A-Z]{1,4}\s+\d{1,4})?\s*\)",
        r"\(\s*(?:\[[12]\d{3}\]|\([12]\d{3}\)|[12]\d{3})\s+\d{1,4}\s+[A-Z]{1,8}(?:\s+[A-Z]{1,8}){0,2}\s+\d{1,5}\s*\)",
    ]
    case_name_nearby = r"\b[A-Z][A-Za-z0-9'’.\-]+(?:\s+[A-Z][A-Za-z0-9'’.\-]+){0,6}\s+v\.?\s+[A-Z][A-Za-z0-9'’.\-]+"
    has_bare_report_only = False
    for _pat in report_only_patterns:
        for _m in re.finditer(_pat, txt, flags=re.IGNORECASE):
            _lookback = txt[max(0, _m.start() - 120):_m.start()]
            if re.search(case_name_nearby, _lookback):
                continue
            has_bare_report_only = True
            break
        if has_bare_report_only:
            break
    if has_bare_report_only:
        issues.append("Contains report-only citation(s) without case name (not acceptable OSCOLA in-text style).")
    if re.search(r"(?im)\bsource\s+\d+\b", txt):
        issues.append("Contains internal source-label leakage (for example 'Source N').")
    if re.search(r"\[\s*RAG CONTEXT", txt, flags=re.IGNORECASE):
        issues.append("Contains leaked internal RAG context markers.")
    if _has_substantial_duplicate_blocks(txt):
        issues.append("Contains substantial duplicated blocks or repetitive paragraphs.")
    paragraph_blocks = [p.strip() for p in re.split(r"\n\s*\n", txt) if (p or "").strip()]
    oversized_paragraph_detected = False
    paragraph_word_cap = 150 if is_problem_mode else 165
    for para in paragraph_blocks:
        if re.match(r"(?im)^(?:Question\s+\d+|Part\s+[IVXLCDM0-9]+|Title)\s*:[^\n]+$", para):
            continue
        if re.match(r"(?im)^[A-D]\.\s*(?:Issue|Rule|Application|Conclusion)\b", para):
            continue
        if len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", para)) > paragraph_word_cap:
            oversized_paragraph_detected = True
            break
    if oversized_paragraph_detected:
        issues.append("Contains oversized paragraph blocks that should be split into smaller analytical paragraphs.")
    issues.extend(_empty_heading_structure_issues(txt))
    issues.extend(_major_unheaded_section_issues(txt, is_problem_mode=is_problem_mode))

    # Citation placement check: case names should generally carry immediate parenthetical refs.
    def _is_case_like_side(side: str) -> bool:
        tokens = re.findall(r"[A-Za-z][A-Za-z0-9'’.\-]*", side or "")
        if not tokens:
            return False
        connectors = {"and", "of", "the", "for", "in", "on", "to", "de", "la", "le", "du", "van", "von", "ex", "p"}
        substantive = [t for t in tokens if t.lower().strip(".") not in connectors and len(t.strip(".")) >= 3]
        if not substantive:
            return False
        return any((t[:1].isupper() or t.isupper()) for t in substantive)

    bare_case_count = 0
    case_counts: Counter[str] = Counter()
    case_display: Dict[str, str] = {}
    # Allow lowercase corporate suffixes/connectors inside case names
    # (for example "Axa ... plc v ... Ltd"), while still requiring case-like sides.
    case_pat = re.compile(
        r"\b([A-Z][A-Za-z0-9'’.\-]*(?:\s+[A-Za-z][A-Za-z0-9'’.\-]*){0,10})\s+v\.?\s+"
        r"([A-Z][A-Za-z0-9'’.\-]*(?:\s+[A-Za-z][A-Za-z0-9'’.\-]*){0,10})\b"
    )
    for m in case_pat.finditer(txt):
        left = (m.group(1) or "").strip()
        right = (m.group(2) or "").strip()
        if not (_is_case_like_side(left) and _is_case_like_side(right)):
            continue
        display_name = f"{left} v {right}"
        norm_name = re.sub(r"\s+", " ", display_name).strip().lower()
        case_counts[norm_name] += 1
        case_display.setdefault(norm_name, display_name)
        after = txt[m.end():m.end() + 14]
        if not re.match(r"^\s*\(", after):
            bare_case_count += 1
            if bare_case_count >= 1:
                issues.append("Case references are not followed by immediate OSCOLA-style parenthetical citations.")
                break

    total_case_mentions = sum(case_counts.values())
    if total_case_mentions >= 4 and case_counts:
        dominant_case, dominant_count = case_counts.most_common(1)[0]
        if dominant_count >= 4 and (dominant_count / total_case_mentions) >= 0.60:
            issues.append(
                f"Over-relies on a single authority ({case_display.get(dominant_case, dominant_case)}) instead of using a balanced spread of sources."
            )

    repeated_case_citations: Counter[str] = Counter()
    repeated_case_display: Dict[str, str] = {}
    for m in re.finditer(r"\(([^()\n]{12,260}\bv\.?\b[^()\n]{3,260})\)", txt):
        inner = re.sub(r"\s+", " ", (m.group(1) or "")).strip(" ,;:.")
        if not inner:
            continue
        key = inner.lower()
        repeated_case_citations[key] += 1
        repeated_case_display.setdefault(key, inner)
    if repeated_case_citations:
        rep_key, rep_count = repeated_case_citations.most_common(1)[0]
        if rep_count >= 3 and not any("over-relies on a single authority" in i.lower() for i in issues):
            issues.append(
                f"Over-relies on a single authority ({repeated_case_display.get(rep_key, rep_key)}) instead of using a balanced spread of sources."
            )

    citation_count = len(
        re.findall(
            r"\([^()\n]{3,260}(?:\[[12]\d{3}\]|Act\s+\d{4}|Article\s+\d+(?:\(\d+\))?|section\s+\d+|s\.?\s*\d+)",
            txt,
            flags=re.IGNORECASE,
        )
    )
    actual_words_now = _count_words(txt)
    min_citations = 0
    if actual_words_now >= 1800:
        min_citations = 9
    elif actual_words_now >= 1200:
        min_citations = 7
    elif actual_words_now >= 800:
        min_citations = 5
    elif actual_words_now >= 500:
        min_citations = 4
    elif actual_words_now >= 300:
        min_citations = 3
    if min_citations and citation_count < min_citations:
        issues.append(
            f"Inline OSCOLA citation density is too thin for the length of the answer ({citation_count} citation parentheticals; expected at least {min_citations})."
        )

    part_blocks = list(re.finditer(r"(?im)^\s*Part\s+([IVXLCDM]+|\d+)\s*:\s*(.+?)\s*$", txt))
    analytical_part_without_citation = False
    for i, m in enumerate(part_blocks):
        title_low = ((m.group(2) or "").strip().lower())
        if "introduction" in title_low or "conclusion" in title_low:
            continue
        block_start = m.end()
        block_end = part_blocks[i + 1].start() if (i + 1) < len(part_blocks) else len(txt)
        block = txt[block_start:block_end].strip()
        if _count_words(block) < 80:
            continue
        if not re.search(
            r"\([^()\n]{3,260}(?:\[[12]\d{3}\]|Act\s+\d{4}|Article\s+\d+(?:\(\d+\))?|section\s+\d+|s\.?\s*\d+)",
            block,
            flags=re.IGNORECASE,
        ):
            analytical_part_without_citation = True
            break
    if analytical_part_without_citation:
        issues.append("One or more analytical Parts contain no visible inline OSCOLA citation support.")

    # Proposition-level citation coverage (global legal-quality guard).
    # If a sentence contains a legal authority anchor but no parenthetical citation, flag it.
    missing_prop_cite = 0
    sentence_candidates = re.split(r"(?<=[.!?])\s+", txt)
    for sent in sentence_candidates:
        s = (sent or "").strip()
        if len(re.findall(r"\b\w+\b", s)) < 8:
            continue
        has_anchor = bool(
            case_pat.search(s)
            or re.search(r"\b[A-Z][A-Za-z ,&()'-]+ Act \d{4}\b", s)
            or re.search(r"\bArticle\s+\d+(?:\(\d+\))?\b", s, flags=re.IGNORECASE)
            or re.search(r"\b(?:section|s\.?)\s*\d+[a-z]?(?:\(\d+\))?\b", s, flags=re.IGNORECASE)
        )
        if not has_anchor:
            continue
        if not re.search(r"\([^()\n]{3,260}\)", s):
            missing_prop_cite += 1
            if missing_prop_cite >= 1:
                issues.append("Multiple legal proposition sentences contain authority anchors without immediate OSCOLA parenthetical citations.")
                break

    if is_problem_mode:
        issues.extend(_problem_structure_issues(txt))

    # Part-heading monotonicity: prevent restart from Part I late in answer.
    parts = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+)\s*:\s*", txt):
        try:
            parts.append(_roman_to_int(m.group(1)))
        except Exception:
            continue
    if len(parts) >= 2:
        for i in range(1, len(parts)):
            if parts[i] <= parts[i - 1]:
                issues.append("Part numbering regresses or repeats inside the same answer.")
                break
    elif actual_words_now >= 650:
        issues.append("Contains only one Part heading even though the answer is long enough to require further Part-numbered sections.")

    body_no_tail = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    question_heading_matches = list(re.finditer(r"(?im)^\s*Question\s+\d+\s*:\s*.+$", body_no_tail))
    block_bounds: List[Tuple[str, int, int]] = []
    if question_heading_matches:
        for idx, m in enumerate(question_heading_matches):
            start = m.start()
            end = question_heading_matches[idx + 1].start() if (idx + 1) < len(question_heading_matches) else len(body_no_tail)
            block_bounds.append(((m.group(0) or "").strip(), start, end))
    else:
        block_bounds.append(("Answer", 0, len(body_no_tail)))

    def _is_generic_body_part_title(title: str) -> bool:
        normalized = re.sub(r"[^a-z0-9]+", " ", (title or "").lower()).strip()
        if not normalized or any(marker in normalized for marker in ["introduction", "conclusion", "final conclusion"]):
            return False
        tokens = re.findall(r"[a-z]{3,}", normalized)
        if not tokens:
            return False
        generic_tokens = {
            "analysis", "discussion", "evaluation", "further", "more", "additional",
            "general", "doctrine", "doctrinal", "framework", "policy", "policies",
            "context", "background", "overview", "themes", "theme", "limits",
        }
        anchor_tokens = {
            "analysis", "discussion", "evaluation", "doctrine", "doctrinal",
            "framework", "policy", "context", "background", "overview", "limits",
        }
        return all(tok in generic_tokens for tok in tokens) and any(tok in anchor_tokens for tok in tokens)

    for block_label, start, end in block_bounds:
        block = body_no_tail[start:end].strip()
        if not block:
            continue
        block_words = _count_words(block)
        block_part_titles = [
            ((m.group(2) or "").strip().lower())
            for m in re.finditer(r"(?im)^\s*Part\s+([IVXLCDM]+|\d+)\s*:\s*(.+?)\s*$", block)
        ]
        if not block_part_titles:
            continue
        label_prefix = "" if block_label == "Answer" else f"{block_label} "
        intro_count = sum(1 for title in block_part_titles if "introduction" in title)
        conclusion_count = sum(1 for title in block_part_titles if "conclusion" in title)
        if intro_count > 1:
            issues.append(f"{label_prefix}contains more than one introduction heading.")
        if conclusion_count > 1:
            issues.append(f"{label_prefix}contains more than one conclusion heading.")

        generic_body_titles = [title for title in block_part_titles if _is_generic_body_part_title(title)]
        if len(generic_body_titles) >= 2 or any(
            re.fullmatch(
                r"(?:further|more|additional|general)\s+"
                r"(?:analysis|discussion|evaluation|doctrine|framework|policy|context|background|overview|limits)",
                title,
            )
            for title in generic_body_titles
        ):
            issues.append(
                f"{label_prefix}uses vague generic body Part headings instead of real issue/theme headings."
            )

        block_part_count = len(block_part_titles)
        if block_words >= 900:
            if block_words <= 2200:
                max_parts = 6 if not is_problem_mode else 7
            elif block_words <= 4500:
                max_parts = 7 if not is_problem_mode else 8
            else:
                max_parts = 9 if not is_problem_mode else 10
            if block_part_count > max_parts:
                issues.append(
                    f"{label_prefix}uses an over-fragmented Part structure for its length ({block_part_count} Parts over about {block_words} words). Merge overlapping sections and keep Parts to major issue clusters."
                )

    concl_words = _count_words(_extract_conclusion_section_text(txt))
    intro_words = _count_words(_extract_introduction_section_text(txt))
    is_short_single_problem = bool(is_problem_mode and _is_short_single_problem_prompt(prompt_text))
    part_titles_now = [
        ((m.group(2) or "").strip())
        for m in re.finditer(r"(?im)^\s*Part\s+([IVXLCDM]+|\d+)\s*:\s*(.+?)\s*$", txt)
    ]
    if is_short_single_essay:
        # For <=2000-word essays, require the same variable Part-numbered scaffold
        # as longer essays: Part I introduction, body Parts, and one final Conclusion Part.
        if not re.search(r"(?im)^\s*Part\s+I\s*:\s*Introduction\b", txt):
            issues.append("Short essay is missing 'Part I: Introduction'.")
        if not re.search(r"(?im)^\s*Part\s+[IVXLCDM]+\s*:\s*Conclusion\b", txt):
            issues.append("Short essay is missing a final Part-numbered conclusion heading.")
        if len(parts) < 3:
            issues.append("Short essay lacks sufficient Part-numbered structure (need introduction, body, and conclusion).")
        if concl_words < 90:
            issues.append(f"Conclusion is too short ({concl_words} words; require >=90).")
        short_conclusion_cap = min(260, max(170, int(actual_words_now * 0.16)))
        if concl_words > short_conclusion_cap:
            issues.append(
                f"Conclusion is disproportionately long ({concl_words} words; keep roughly <= {short_conclusion_cap} for this essay length)."
            )
    elif is_short_single_problem:
        normalized_part_titles_now = [re.sub(r"\s+", " ", (title or "")).strip().lower() for title in part_titles_now]
        if not re.search(r"(?im)^\s*Part\s+I\s*:\s*Introduction\b", txt):
            issues.append("Short problem answer is missing 'Part I: Introduction'.")
        if "remedies / liability" not in normalized_part_titles_now:
            issues.append("Short problem answer is missing the required Part-numbered 'Remedies / Liability' section.")
        if "final conclusion" not in normalized_part_titles_now:
            issues.append("Short problem answer must use the exact final heading 'Final Conclusion'.")
        if len(parts) < 4:
            issues.append("Short problem answer lacks sufficient Part-numbered structure (need introduction, issue analysis, remedies/liability, and final conclusion).")
        if intro_words > 150:
            issues.append(f"Short problem answer has an overlong introduction ({intro_words} words; keep Part I compact and usually <= 150 words).")
        if concl_words < 90:
            issues.append(f"Final Conclusion is too short ({concl_words} words; require >=90).")
    else:
        # For long-essay final parts, conclusion should still be substantive.
        if concl_words < 100:
            issues.append(f"Conclusion appears too thin ({concl_words} words; require >=100).")

    # If prompt explicitly requests a word target, enforce a practical quality floor.
    targets = _extract_word_targets(prompt_text)
    if len(targets) == 1:
        target = int(targets[0])
        floor = complete_word_count_floor(target)
        actual = _count_words(txt)
        if actual < max(350, floor):
            issues.append(f"Answer under-delivers on requested depth ({actual} words; expected at least {max(350, floor)}).")

    # Explicit issue-coverage guard for a)/b)/c)/d) and requirement bullets.
    # Deterministic + cheap keyword overlap check.
    reqs = _extract_explicit_issue_requirements(prompt_text)
    if reqs:
        low_txt = txt.lower()
        missing_labels: List[str] = []
        for item in reqs[:8]:
            kws = item.get("keywords") or []
            # Require at least one strong keyword hit.
            if not any(re.search(rf"\b{re.escape(k)}\b", low_txt) for k in kws):
                missing_labels.append(item.get("text") or "")
        if missing_labels:
            preview = "; ".join(missing_labels[:4])
            issues.append(f"May not explicitly answer required issue(s): {preview}")

    return issues

def _resolve_word_window_from_history(prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[tuple]:
    """
    Resolve current part word window (min,max) from the latest long-request anchor in history.
    This supports continuation messages like "continue" that do not contain explicit word counts.
    """
    # Fresh explicit request must reset to Part 1 sizing, regardless of earlier history.
    current_targets = _extract_word_targets(prompt_text or "")
    if current_targets and (not _is_continuation_command(prompt_text or "")):
        if len(current_targets) == 1 and int(current_targets[0]) <= 2000:
            target = int(current_targets[0])
            return complete_word_count_window(target)
        current_plan = detect_long_essay(prompt_text or "")
        if current_plan.get("is_long_essay"):
            current_deliverables = current_plan.get("deliverables") or []
            if current_deliverables:
                current_target = int(current_deliverables[0].get("target_words", 0) or 0)
            else:
                current_target = int(current_plan.get("words_per_part") or 0)
            if current_target > 0:
                return complete_word_count_window(current_target)

    if not messages:
        return None

    anchor_idx = -1
    anchor_text = ""
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if _looks_like_pasted_generation(txt):
            continue
        if _extract_word_targets(txt):
            anchor_idx = i
            anchor_text = txt
            break
    if anchor_idx < 0 or not anchor_text:
        return None

    plan = detect_long_essay(anchor_text)
    if not plan.get("is_long_essay"):
        anchor_targets = _extract_word_targets(anchor_text)
        if len(anchor_targets) == 1 and int(anchor_targets[0]) <= 2000:
            target = int(anchor_targets[0])
            return complete_word_count_window(target)
        return None

    assistant_messages = [
        m for m in messages[anchor_idx + 1:]
        if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
    ]
    assistants_after_anchor = len(assistant_messages)
    hold_same_part = _last_assistant_requires_same_logical_part(assistant_messages)
    current_part = assistants_after_anchor if hold_same_part else (assistants_after_anchor + 1)
    current_part = max(1, current_part)
    deliverables = plan.get("deliverables") or []
    total_requested = int(plan.get("requested_words") or 0)
    consumed_prior = sum(_count_words(m.get("text") or "") for m in assistant_messages)

    if deliverables:
        current_part = max(1, min(current_part, len(deliverables)))
        target = int(deliverables[current_part - 1].get("target_words", 0) or 0)
        cap = max(int(d.get("target_words", 0) or 0) for d in deliverables) if deliverables else target
        remaining_parts = max(1, len(deliverables) - current_part + 1)
    else:
        target = int(plan.get("words_per_part") or 0)
        cap = target
        remaining_parts = max(1, int(plan.get("suggested_parts") or 1) - current_part + 1)

    # Dynamic rebalance so cumulative total can still land in-range.
    if total_requested > 0 and cap > 0:
        remaining_total = max(1, total_requested - consumed_prior)
        if remaining_parts <= 1:
            target = min(cap, remaining_total)
        else:
            dynamic_share = int(math.ceil(remaining_total / remaining_parts))
            target = min(cap, max(1, dynamic_share))

    if target <= 0:
        return None
    # Strict per-part window: 99-100% of target.
    return complete_word_count_window(target)


def _active_question_block(answer_text: str) -> str:
    """
    When a response contains multiple Question N blocks, continuation logic
    should key off the last active question block instead of earlier ones.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt
    matches = list(re.finditer(r"(?im)^\s*Question\s+\d+\s*:\s*.+?$", txt))
    if not matches:
        return txt
    return txt[matches[-1].start():].strip()


def _last_assistant_requires_same_logical_part(assistant_messages: List[Dict[str, Any]]) -> bool:
    """
    If the latest assistant response ended with an open structural defect,
    keep the next continuation on the same logical part instead of advancing.
    """
    if not assistant_messages:
        return False
    last_text = (assistant_messages[-1].get("text") or "").strip()
    if not last_text:
        return False
    if re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", last_text):
        return False
    if re.search(r"\(End of Answer\)\s*$", last_text, flags=re.IGNORECASE):
        return False
    last_text = re.sub(r"\(End of Answer\)\s*$", "", last_text, flags=re.IGNORECASE).strip()
    last_text = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", last_text).strip()
    if not last_text:
        return False
    active = _active_question_block(last_text)
    empty_issues = _empty_heading_structure_issues(active)
    if any(
        ("bare lettered subheading" in issue.lower())
        or ("empty lettered subsection" in issue.lower())
        for issue in empty_issues
    ):
        return True
    open_part_issues = _problem_structure_issues(active)
    return any("Final analytical Part is left open" in issue for issue in open_part_issues)

def _expected_part_state_from_history(prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Determine whether the current response should be intermediate or final
    based on the latest anchored word-count request.
    """
    # Fresh explicit request must start at Part 1 (do not inherit stale anchors).
    current_targets = _extract_word_targets(prompt_text or "")
    if current_targets and (not _is_continuation_command(prompt_text or "")):
        current_plan = detect_long_essay(prompt_text or "")
        if current_plan.get("is_long_essay"):
            current_deliverables = current_plan.get("deliverables") or []
            current_total_parts = len(current_deliverables) if current_deliverables else int(current_plan.get("suggested_parts") or 0)
            if current_total_parts > 0:
                return {"current_part": 1, "total_parts": current_total_parts, "is_final": current_total_parts <= 1}

    if not messages:
        return None
    anchor_idx = -1
    anchor_text = ""
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if _looks_like_pasted_generation(txt):
            continue
        if _extract_word_targets(txt):
            anchor_idx = i
            anchor_text = txt
            break
    if anchor_idx < 0 or not anchor_text:
        return None
    plan = detect_long_essay(anchor_text)
    if not plan.get("is_long_essay"):
        return None
    assistant_messages = [
        m for m in messages[anchor_idx + 1:]
        if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
    ]
    assistants_after_anchor = len(assistant_messages)
    hold_same_part = _last_assistant_requires_same_logical_part(assistant_messages)
    current_part = assistants_after_anchor if hold_same_part else (assistants_after_anchor + 1)
    current_part = max(1, current_part)
    deliverables = plan.get("deliverables") or []
    total_parts = len(deliverables) if deliverables else int(plan.get("suggested_parts") or 0)
    if total_parts <= 0:
        return None
    is_final = current_part >= total_parts
    return {"current_part": current_part, "total_parts": total_parts, "is_final": is_final}

def _expected_internal_part_heading_from_history(prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[int]:
    """
    Derive the next in-answer Part heading number from the latest assistant part
    after the active word-count anchor.

    Example:
    - Previous response ended at "Part IV: ..."
    - Next continuation should start at "Part V: ...", regardless of chunk index.
    """
    if not messages:
        return None

    anchor_idx = -1
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        if _extract_word_targets(txt):
            anchor_idx = i
            break
    if anchor_idx < 0:
        return None

    prior_assistants = [
        m for m in messages[anchor_idx + 1:]
        if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
    ]
    if not prior_assistants:
        return None

    last_text = (prior_assistants[-1].get("text") or "").strip()
    if not last_text:
        return None
    last_text = re.sub(r"\(End of Answer\)\s*$", "", last_text, flags=re.IGNORECASE).strip()
    last_text = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", last_text).strip()
    if not last_text:
        return None
    last_text = _active_question_block(last_text)

    nums: List[int] = []
    for m in re.finditer(r"(?im)^\s*Part\s+([IVXLC]+|\d+)\s*:\s*", last_text):
        token = (m.group(1) or "").strip()
        try:
            if token.isdigit():
                n = int(token)
            else:
                n = _roman_to_int(token)
            if n > 0:
                nums.append(n)
        except Exception:
            continue
    if not nums:
        return None

    last_part_num = max(nums)
    open_part_issues = _problem_structure_issues(last_text)
    if (
        _is_abrupt_answer_ending(last_text)
        or _empty_heading_structure_issues(last_text)
        or any("Final analytical Part is left open" in issue for issue in open_part_issues)
    ):
        return last_part_num
    return last_part_num + 1


def _expected_unit_structure_state_from_history(prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Resolve the expected unit/question state for the current response so the
    deterministic structure validator can run in the backend post-processing path.
    """
    current_targets = _extract_word_targets(prompt_text or "")
    anchor_idx = -1
    anchor_text = ""

    if current_targets and (not _is_continuation_command(prompt_text or "")):
        anchor_text = prompt_text or ""
    else:
        for i in range(len(messages) - 1, -1, -1):
            msg = messages[i]
            if msg.get("role") != "user":
                continue
            txt = (msg.get("text") or "").strip()
            if not txt or _looks_like_pasted_generation(txt):
                continue
            if _extract_word_targets(txt):
                anchor_idx = i
                anchor_text = txt
                break

    if not anchor_text:
        return None

    plan = detect_long_essay(anchor_text)
    deliverables = plan.get("deliverables") or []
    if not plan.get("is_long_essay") or not deliverables:
        return None

    if anchor_idx >= 0:
        assistant_messages = [
            m for m in messages[anchor_idx + 1:]
            if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
        ]
    else:
        assistant_messages = []

    assistants_after_anchor = len(assistant_messages)
    hold_same_part = _last_assistant_requires_same_logical_part(assistant_messages)
    current_part = assistants_after_anchor if hold_same_part else (assistants_after_anchor + 1)
    current_part = max(1, min(current_part, len(deliverables)))
    d = deliverables[current_part - 1]

    current_q = int(d.get("starting_question_index", d.get("question_index", 0)) or 0) or None
    current_kind = str(d.get("starting_unit_kind", d.get("unit_kind")) or "").strip().lower() or None
    current_title = str((d.get("question_titles") or [""])[0] or "").strip()
    current_section_index = int(d.get("section_index", 0) or 0) or None
    current_subpart_index = int(
        d.get("subpart_index", d.get("part_in_section", 1)) or 1
    )
    current_subpart_total = int(
        d.get("subpart_total", d.get("parts_in_section", 1)) or 1
    )
    fragments = d.get("fragments") or []
    multi_question_part = len({
        int(f.get("question_index", 0) or 0)
        for f in fragments
        if int(f.get("question_index", 0) or 0) > 0
    }) >= 2
    overall_question_count = len({
        int(dd.get("question_index", 0) or 0)
        for dd in deliverables
        if int(dd.get("question_index", 0) or 0) > 0
    })
    total_requested = int(plan.get("requested_words") or 0)
    use_question_headings = bool(overall_question_count >= 2 and total_requested > LONG_ESSAY_THRESHOLD)
    current_heading = ""
    if current_q and use_question_headings:
        normalized_title = re.sub(r"(?im)^\s*(essay|problem)\s+question\s*[:\-]?\s*", "", current_title).strip()
        current_heading = f"Question {int(current_q)}: {normalized_title or ('Question ' + str(int(current_q)))}"
    use_global_part_sequence = False

    same_topic_continuation = False
    if current_part >= 2:
        prev_d = deliverables[current_part - 2]
        prev_q = int(prev_d.get("ending_question_index", prev_d.get("question_index", 0)) or 0) or None
        prev_section_index = int(prev_d.get("section_index", 0) or 0) or None
        same_topic_continuation = bool(
            (current_q and prev_q and current_q == prev_q)
            or ((not current_q) and current_section_index and prev_section_index and current_section_index == prev_section_index)
        )

    starts_new_question = bool(current_subpart_index == 1)
    expected_internal_part_number = _expected_internal_part_heading_from_history(prompt_text, messages)
    if starts_new_question:
        expected_internal_part_number = 1
    elif not expected_internal_part_number or expected_internal_part_number <= 0:
        expected_internal_part_number = 1 if assistants_after_anchor == 0 else int(current_part)

    return {
        "current_part": current_part,
        "total_parts": len(deliverables),
        "question_index": current_q,
        "question_title": current_title,
        "question_heading": current_heading,
        "question_text": str((d.get("unit_texts") or [""])[0] or "").strip(),
        "unit_kind": current_kind,
        "require_question_heading": bool(use_question_headings and current_q and (not multi_question_part)),
        "is_same_topic_continuation": same_topic_continuation,
        "expected_part_number": int(expected_internal_part_number),
        "starts_new_question": starts_new_question,
        "enforce_single_top_level_part": False,
        "question_final_part": bool(current_subpart_index >= current_subpart_total),
        "use_global_part_sequence": use_global_part_sequence,
    }


def _current_unit_mode_from_history(prompt_text: str, messages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Resolve whether the current planned response should be treated as essay or
    problem mode. For mixed prompts, prefer the active deliverable's unit kind
    over prompt-wide keyword detection.
    """
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    active_kind = str((state or {}).get("unit_kind") or "").strip().lower()
    prompt_problem = _is_problem_flow(prompt_text, messages)
    prompt_essay = _is_essay_flow(prompt_text, messages)
    return {
        "unit_kind": active_kind or None,
        "state": state,
        "is_problem_mode": bool((active_kind == "problem") or ((not active_kind) and prompt_problem)),
        "is_essay_mode": bool((active_kind == "essay") or ((not active_kind) and prompt_essay)),
    }


def _enforce_expected_question_heading(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> str:
    """
    Ensure long multi-question outputs expose exactly one correct Question heading
    at the top of the current response.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if not state or not state.get("require_question_heading"):
        return txt

    expected_heading = str(state.get("question_heading") or "").strip()
    if not expected_heading:
        q_num = int(state.get("question_index") or 0)
        if q_num <= 0:
            return txt
        expected_heading = f"Question {q_num}: Question {q_num}"

    body = re.sub(r"(?im)^\s*Question\s+\d+\s*:\s*.+\s*$", "", txt).strip()
    if not body:
        return expected_heading
    return f"{expected_heading}\n\n{body}".strip()


def _collapse_extra_top_level_parts(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> str:
    """
    In single-top-level-part mode, collapse any later Part headings into body
    subsections so legacy cleanup cannot create fake new top-level Parts.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if not state or not state.get("enforce_single_top_level_part"):
        return txt

    txt = _enforce_expected_question_heading(txt, prompt_text, messages)
    unit_kind = str(state.get("unit_kind") or "").strip().lower()
    rebuilt: List[str] = []
    part_seen = 0
    for line in txt.splitlines():
        m = re.match(r"(?im)^\s*Part\s+([IVXLCDM0-9]+)\s*:\s*(.+?)\s*$", line.strip())
        if not m:
            rebuilt.append(line)
            continue
        part_seen += 1
        if part_seen == 1:
            rebuilt.append(line.strip())
            continue
        title = (m.group(2) or "").strip()
        if not title:
            continue
        if "conclusion" in title.lower():
            rebuilt.append("D. Conclusion" if unit_kind == "problem" else "Conclusion")
        else:
            rebuilt.append(title)

    cleaned = "\n".join(rebuilt).strip()
    cleaned = _normalize_output_style(cleaned)
    cleaned = _restore_paragraph_separation(cleaned)
    return cleaned.strip()


def _continuation_repeat_issue(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[str]:
    """
    Detect when a continuation part substantially repeats the immediately
    previous assistant part instead of progressing.
    """
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if not state or not state.get("is_same_topic_continuation"):
        return None

    current_part = int(state.get("current_part") or 1)
    if current_part <= 1:
        return None

    anchor_idx = -1
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        if _extract_word_targets(txt):
            anchor_idx = i
            break
    if anchor_idx < 0:
        return None

    assistant_messages = [
        m for m in messages[anchor_idx + 1:]
        if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
    ]
    if not assistant_messages:
        return None

    prev_raw = (assistant_messages[-1].get("text") or "").strip()
    cur_raw = (answer_text or "").strip()
    if not prev_raw or not cur_raw:
        return None

    def _clean(raw: str) -> str:
        cleaned = _strip_generation_artifacts(raw)
        cleaned = re.sub(r"\(End of Answer\)\s*$", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", cleaned).strip()
        cleaned = _active_question_block(cleaned)
        return cleaned.strip()

    prev = _clean(prev_raw)
    cur = _clean(cur_raw)
    if not prev or not cur:
        return None

    def _norm(raw: str) -> str:
        cleaned = re.sub(r"(?im)^\s*(Question\s+\d+\s*:|Part\s+[IVXLC\d]+\s*:[^\n]*|[A-D]\.\s+\w+[^\n]*)\s*$", "", raw)
        cleaned = re.sub(r"\s+", " ", cleaned).strip().lower()
        return cleaned

    prev_n = _norm(prev)
    cur_n = _norm(cur)
    if len(prev_n.split()) < 60 or len(cur_n.split()) < 60:
        return None

    cur_probe = cur_n[: min(len(cur_n), 900)].strip()
    if len(cur_probe.split()) >= 40 and cur_probe in prev_n:
        return "Current continuation substantially repeats the previous part (copied body detected)."

    compare_len = min(len(prev_n), len(cur_n), 1800)
    if compare_len < 500:
        return None

    ratio = SequenceMatcher(None, prev_n[:compare_len], cur_n[:compare_len]).ratio()
    if ratio >= 0.72:
        return f"Current continuation substantially repeats the previous part (similarity {ratio:.2f})."

    def _ngrams(words: List[str], n: int) -> set:
        if len(words) < n:
            return set()
        return {" ".join(words[i:i + n]) for i in range(0, len(words) - n + 1)}

    prev_words = prev_n.split()[:1600]
    cur_words = cur_n.split()[:1600]
    prev_grams = _ngrams(prev_words, 6)
    cur_grams = _ngrams(cur_words, 6)
    if prev_grams and cur_grams:
        overlap = len(prev_grams & cur_grams) / max(1, min(len(prev_grams), len(cur_grams)))
    if overlap >= 0.28:
        return f"Current continuation substantially repeats the previous part (6-gram overlap {overlap:.2f})."

    return None


def _new_question_repeat_issue(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> Optional[str]:
    """
    Detect when a response should start a NEW question but instead substantially
    repeats the previous question's content.
    """
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if not state:
        return None
    if bool(state.get("is_same_topic_continuation")) or not bool(state.get("starts_new_question")):
        return None

    anchor_idx = -1
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        txt = (msg.get("text") or "").strip()
        if not txt or _looks_like_pasted_generation(txt):
            continue
        if _extract_word_targets(txt):
            anchor_idx = i
            break
    if anchor_idx < 0:
        return None

    assistant_messages = [
        m for m in messages[anchor_idx + 1:]
        if m.get("role") == "assistant" and _assistant_message_counts_as_part(m.get("text") or "")
    ]
    if not assistant_messages:
        return None

    prev_raw = (assistant_messages[-1].get("text") or "").strip()
    cur_raw = (answer_text or "").strip()
    if not prev_raw or not cur_raw:
        return None

    expected_heading = str(state.get("question_heading") or "").strip()
    actual_heading_match = re.search(r"(?im)^\s*Question\s+\d+\s*:\s*.+$", cur_raw)
    if expected_heading and actual_heading_match:
        actual_heading = actual_heading_match.group(0).strip()
        expected_num_m = re.search(r"(?i)\bquestion\s+(\d+)\b", expected_heading)
        actual_num_m = re.search(r"(?i)\bquestion\s+(\d+)\b", actual_heading)
        if expected_num_m and actual_num_m and expected_num_m.group(1) != actual_num_m.group(1):
            return "Current part appears to repeat the previous question instead of starting the planned new question."

        def _normalize_question_heading_core(s: str) -> str:
            core = re.sub(r"(?i)^\s*question\s+\d+\s*:\s*", "", s or "").strip()
            core = re.sub(r"\s*\([^)\n]*\)", "", core).strip()
            core = re.sub(r"\s+", " ", core).strip(" .:-").lower()
            return core

        expected_core = _normalize_question_heading_core(expected_heading)
        actual_core = _normalize_question_heading_core(actual_heading)
        if expected_core and actual_core:
            if (
                expected_core == actual_core
                or expected_core in actual_core
                or actual_core in expected_core
            ):
                pass
            # Otherwise, do not hard-fail on title wording alone. The numbered-heading
            # structure validator already catches wrong question numbers, and the body
            # similarity test below is better at detecting genuine repeat content.

    def _clean(raw: str) -> str:
        cleaned = _strip_generation_artifacts(raw)
        cleaned = re.sub(r"\(End of Answer\)\s*$", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", cleaned).strip()
        cleaned = re.sub(r"(?im)^\s*Question\s+\d+\s*:\s*.+$", "", cleaned).strip()
        cleaned = re.sub(r"(?im)^\s*Part\s+[IVXLC\d]+\s*:\s*.+$", "", cleaned).strip()
        cleaned = re.sub(r"(?im)^\s*[A-D]\.\s+\w+[^\n]*$", "", cleaned).strip()
        cleaned = re.sub(r"\s+", " ", cleaned).strip().lower()
        return cleaned

    prev_n = _clean(prev_raw)
    cur_n = _clean(cur_raw)
    if len(prev_n.split()) < 60 or len(cur_n.split()) < 60:
        return None

    def _content_token_set(s: str) -> set:
        stop = {
            "the", "and", "for", "that", "this", "with", "from", "into", "onto", "over",
            "under", "after", "before", "between", "about", "because", "where", "when",
            "which", "while", "would", "could", "should", "their", "there", "these",
            "those", "have", "has", "had", "were", "was", "will", "shall", "must", "may",
            "might", "such", "only", "also", "than", "then", "them", "they", "being",
            "part", "question", "issue", "issues", "essay", "problem", "introduction",
            "conclusion", "analysis", "application", "rule", "rules", "facts", "legal",
            "likely", "arguably", "court", "courts", "law", "laws", "claim", "claimant",
            "defendant", "duty", "care", "whether", "first", "second", "third",
        }
        words = [w.lower() for w in re.findall(r"[A-Za-z][A-Za-z'’-]*", s or "")]
        return {w for w in words if len(w) >= 4 and w not in stop}

    prev_tokens = _content_token_set(prev_n)
    cur_tokens = _content_token_set(cur_n)
    token_overlap = len(prev_tokens & cur_tokens) / max(1, min(len(prev_tokens), len(cur_tokens)))

    compare_len = min(len(prev_n), len(cur_n), 1800)
    ratio = SequenceMatcher(None, prev_n[:compare_len], cur_n[:compare_len]).ratio()
    if ratio >= 0.72 or (ratio >= 0.62 and token_overlap >= 0.20):
        return f"Current part appears to repeat the previous question instead of starting the planned new question (similarity {ratio:.2f})."

    return None


def _history_aware_structure_issues(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> List[str]:
    """
    Apply the core per-unit structure validator and continuation-repeat check in
    the backend post-processing layer.
    """
    issues: List[str] = []
    state = _expected_unit_structure_state_from_history(prompt_text, messages)
    if state:
        violates, reason = detect_unit_structure_policy_violation(
            answer_text,
            unit_kind=state.get("unit_kind"),
            require_question_heading=bool(state.get("require_question_heading")),
            expected_question_number=state.get("question_index"),
            is_same_topic_continuation=bool(state.get("is_same_topic_continuation")),
            expected_part_number=state.get("expected_part_number"),
            starts_new_question=state.get("starts_new_question"),
            enforce_single_top_level_part=bool(state.get("enforce_single_top_level_part")),
        )
        if violates and reason:
            issues.append(f"Unit-structure policy violation: {reason}.")
        txt = (answer_text or "").strip()
        question_heading_matches = list(re.finditer(r"(?im)^\s*Question\s+\d+\s*:\s*.+$", txt))
        if len(question_heading_matches) > 1:
            issues.append("Response contains more than one Question heading; the next question started too early in the same response.")

        conclusion_matches = list(
            re.finditer(
                r"(?im)^\s*(?:Part\s+(?:[IVXLC]+|\d+)\s*:\s*Conclusion(?:\s+and\s+Advice)?\b|D\.\s*Conclusion\b|Conclusion(?:\s+and\s+Advice)?\b)\s*$",
                txt,
            )
        )
        if conclusion_matches:
            last_conclusion = conclusion_matches[-1]
            trailing_after_conclusion = txt[last_conclusion.end():].strip()
            trailing_words = _count_words(trailing_after_conclusion)
            if trailing_words >= 60:
                issues.append("Content continues after a conclusion heading instead of ending the current question cleanly.")
            elif re.search(r"(?im)^\s*(Question\s+\d+\s*:|Part\s+(?:[IVXLC]+|\d+)\s*:\s*Introduction\b|Introduction\b)", trailing_after_conclusion):
                issues.append("A fresh introduction or new question begins after the conclusion of the current question.")
            if (not state.get("question_final_part")) and len(question_heading_matches) <= 1:
                issues.append("Current part is not the final planned part for this question but already contains a visible conclusion heading.")

        if state.get("question_final_part"):
            unit_kind = str(state.get("unit_kind") or "").strip().lower()
            structural_conclusion_heading = bool(
                re.search(
                    r"(?im)^\s*(?:Part\s+(?:[IVXLC]+|\d+)\s*:\s*Conclusion(?:\s+and\s+Advice)?\b|D\.\s*Conclusion\b|Conclusion(?:\s+and\s+Advice)?\b)",
                    txt,
                )
            )
            if not structural_conclusion_heading:
                issues.append("Current part is the final planned part for this question but has no visible conclusion heading.")
            elif conclusion_matches:
                trailing_after_conclusion = txt[conclusion_matches[-1].end():].strip()
                if _count_words(trailing_after_conclusion) >= 35:
                    issues.append("Current part is the final planned part for this question but continues with substantial material after the conclusion heading.")
            if _is_abrupt_answer_ending(txt):
                issues.append("Current part is the final planned part for this question but ends abruptly.")
            empty_issues = _empty_heading_structure_issues(txt)
            if any(
                ("bare lettered subheading" in issue.lower())
                or ("empty lettered subsection" in issue.lower())
                or ("part heading with no substantive content" in issue.lower())
                for issue in empty_issues
            ):
                issues.append("Current part is the final planned part for this question but still contains empty structural headings.")
            if unit_kind == "problem":
                open_problem = _problem_structure_issues(txt)
                if any("Final analytical Part is left open" in issue for issue in open_problem):
                    issues.append("Current problem-question final part leaves an analytical section open instead of finishing the advice.")
            if re.search(r"\(End of Answer\)\s*$", txt, flags=re.IGNORECASE):
                if (
                    _is_abrupt_answer_ending(txt)
                    or not structural_conclusion_heading
                    or any(
                        ("bare lettered subheading" in issue.lower())
                        or ("empty lettered subsection" in issue.lower())
                        or ("part heading with no substantive content" in issue.lower())
                        for issue in empty_issues
                    )
                ):
                    issues.append("End-of-answer marker appears even though the final planned question-part is incomplete.")

    repeat_issue = _continuation_repeat_issue(answer_text, prompt_text, messages)
    if repeat_issue:
        issues.append(repeat_issue)
    new_question_repeat = _new_question_repeat_issue(answer_text, prompt_text, messages)
    if new_question_repeat:
        issues.append(new_question_repeat)

    return issues

def _derive_deterministic_title_from_prompt(prompt_text: str) -> str:
    """
    Derive a stable title from the active prompt text for single-question 10,000+ Part 1 essays.
    """
    txt = (prompt_text or "").strip()
    if not txt:
        return "Legal Analysis"

    # Prefer explicit question headings.
    patterns = [
        r"(?im)^\s*QUESTION\s*\d+\s*[–\-:]\s*(.+?)\s*$",
        r"(?im)^\s*ESSAY QUESTION\s*[–\-:]\s*(.+?)\s*$",
        r"(?im)^\s*PROBLEM QUESTION\s*[–\-:]\s*(.+?)\s*$",
    ]
    for pat in patterns:
        m = re.search(pat, txt)
        if m:
            cand = re.sub(r"\s+", " ", (m.group(1) or "").strip()).strip(" .:-")
            if cand:
                return cand[:140]

    # Fallback: first meaningful line that is not control text.
    for ln in txt.splitlines():
        s = (ln or "").strip().strip("•-")
        if not s:
            continue
        low = s.lower()
        if any(k in low for k in ["words", "test", "part 1", "continue", "output"]):
            continue
        if len(s) >= 8:
            return re.sub(r"\s+", " ", s).strip(" .:-")[:140]
    return "Legal Analysis"

def _enforce_deterministic_title_policy(
    answer_text: str,
    prompt_text: str,
    messages: List[Dict[str, Any]],
) -> str:
    """
    Deterministic title policy:
    - single-question essay >= 10000 words: title allowed on Part 1 only.
    - multi-question requests: title forbidden.
    - standard essays under 10000 words: title forbidden.
    - problem questions: title forbidden.
    - Continuation parts: title forbidden.
    """
    txt = (answer_text or "").strip()
    if not txt:
        return txt

    # Resolve active requested total from latest anchored request, falling back to current prompt.
    requested_total = 0
    anchor_idx = -1
    for i in range(len(messages) - 1, -1, -1):
        msg = messages[i]
        if msg.get("role") != "user":
            continue
        u = (msg.get("text") or "").strip()
        if not u or _looks_like_pasted_generation(u):
            continue
        targets = _extract_word_targets(u)
        if targets:
            requested_total = sum(int(n) for n in targets)
            anchor_idx = i
            break
    if requested_total <= 0:
        requested_total = sum(_extract_word_targets(prompt_text or ""))

    active_prompt_for_title = prompt_text or ""
    if anchor_idx >= 0 and anchor_idx < len(messages):
        active_prompt_for_title = (messages[anchor_idx].get("text") or active_prompt_for_title)
    active_question_count = _count_active_question_units(active_prompt_for_title)
    is_problem_flow = _is_problem_flow(active_prompt_for_title, messages)

    state = _expected_part_state_from_history(prompt_text, messages)
    current_part = int(state.get("current_part")) if state else 1
    is_continuation_like = _is_continuation_command(prompt_text or "") or current_part > 1
    title_required = (
        (not is_problem_flow)
        and
        active_question_count == 1
        and requested_total >= 10000
        and (not is_continuation_like)
        and current_part == 1
    )

    # Normalize title lines first.
    title_line_re = re.compile(r"(?im)^\s*Title\s*:\s*.*$")
    lines = txt.splitlines()
    kept: List[str] = []
    seen_title = False
    for ln in lines:
        if title_line_re.match(ln):
            if title_required and (not seen_title):
                kept.append(re.sub(r"\s+", " ", ln).strip())
                seen_title = True
            # else drop repeated/forbidden titles
            continue
        kept.append(ln)
    txt = "\n".join(kept).strip()

    if active_question_count >= 2 or re.search(r"(?im)^\s*Question\s+\d+\s*:", txt):
        txt = re.sub(r"(?im)^\s*Title\s*:\s*.*(?:\n+)?", "", txt).strip()
        return txt

    if not title_required:
        # Strip standalone pre-Part heading before Part I when title is forbidden.
        m_part_i = re.search(r"(?im)^\s*Part\s+I\s*:\s*Introduction\b", txt)
        if m_part_i and m_part_i.start() > 0:
            prefix = txt[:m_part_i.start()].strip()
            if prefix:
                prefix_lines = [ln.strip() for ln in prefix.splitlines() if ln.strip()]
                # Drop short decorative/raw heading prefix blocks.
                if len(prefix_lines) <= 3 and not any(
                    re.search(r"(?i)\b(essay question|problem question|question\s*\d+)\b", ln)
                    for ln in prefix_lines
                ):
                    txt = txt[m_part_i.start():].lstrip()
        return txt.strip()

    # title required for Part 1: ensure exactly one title at top.
    if not re.search(r"(?im)^\s*Title\s*:\s*.+$", txt):
        title = _derive_deterministic_title_from_prompt(prompt_text)
        txt = f"Title: {title}\n\n{txt}".strip()
    else:
        # Ensure title is first visible line.
        m_first_title = re.search(r"(?im)^\s*Title\s*:\s*.+$", txt)
        if m_first_title and m_first_title.start() > 0:
            title_line = m_first_title.group(0).strip()
            body = (txt[:m_first_title.start()] + "\n" + txt[m_first_title.end():]).strip()
            txt = f"{title_line}\n\n{body}".strip()

    return txt.strip()

def _enforce_part_ending_by_history(answer_text: str, prompt_text: str, messages: List[Dict[str, Any]]) -> str:
    """
    Force correct part marker based on expected part state:
    - intermediate: must end with Will Continue
    - final: must end with (End of Answer)
    """
    state = _expected_part_state_from_history(prompt_text, messages)
    answer_text = _enforce_expected_question_heading(answer_text, prompt_text, messages)
    answer_text = _enforce_expected_part_heading(answer_text, prompt_text, messages)
    answer_text = _collapse_extra_top_level_parts(answer_text, prompt_text, messages)
    answer_text = _renumber_internal_part_sequence(answer_text)
    answer_text = _enforce_expected_question_heading(answer_text, prompt_text, messages)
    if not state:
        # Single-response (or unknown state): treat any "Will Continue..." marker as stray
        # and force a clean final ending.
        plan = detect_long_essay(prompt_text or "")
        if not plan.get("is_long_essay"):
            txt = _strip_generation_artifacts(answer_text or "").strip()
            txt = _strip_draft_continuation_opener(txt)
            txt = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", txt).strip()
            txt = _ensure_clean_terminal_sentence(txt, is_intermediate=False)
            return _enforce_end_of_answer(txt)
        # For true long responses, fall back to prior behaviour.
        return _enforce_end_of_answer(answer_text)

    txt = _strip_generation_artifacts(answer_text or "").strip()
    txt = _strip_draft_continuation_opener(txt)
    # Remove any accidental internal end markers; the app will add exactly one correct ending.
    txt = re.sub(r"\(End of Answer\)\s*", "", txt, flags=re.IGNORECASE).strip()
    txt = re.sub(r"\(End of Essay\)\s*", "", txt, flags=re.IGNORECASE).strip()
    txt = re.sub(r"\(End of Problem Question\)\s*", "", txt, flags=re.IGNORECASE).strip()
    txt = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", txt).strip()
    # Strip any trailing end marker variants (already removed globally above, but keep for safety).
    txt = re.sub(r"\(End of Answer\)\s*$", "", txt, flags=re.IGNORECASE).strip()
    if not txt:
        txt = "(No content generated.)"

    if state["is_final"]:
        txt = _ensure_clean_terminal_sentence(txt, is_intermediate=False)
        return txt + "\n\n(End of Answer)"
    txt = _ensure_clean_terminal_sentence(txt, is_intermediate=True)
    return txt + "\n\nWill Continue to next part, say continue"

def _truncate_to_word_cap(answer_text: str, max_words: int, min_words: int = 1) -> str:
    """
    Hard-cap output to max_words while preserving continuation/end markers.
    This is a final safety net for strict 99-100% word limits.
    """
    if max_words <= 0:
        return answer_text
    if min_words <= 0:
        min_words = 1
    if min_words > max_words:
        min_words = max_words

    text = answer_text or ""
    has_continue = bool(re.search(r"(?im)^\s*Will Continue to next part, say continue\s*$", text))
    has_end = bool(re.search(r"\(End of Answer\)\s*$", text, flags=re.IGNORECASE))

    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", text).strip()
    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()

    # Count "words" using the same token pattern as _count_words so hard-cap
    # truncation aligns with enforcement logic while preserving original layout.
    word_matches = list(re.finditer(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", body))
    if len(word_matches) <= max_words:
        return answer_text

    # Prefer a sentence boundary inside the required [min_words, max_words]
    # window so the response ends on a complete thought.
    word_ends = [m.end() for m in word_matches]
    sentence_end_re = re.compile(r'[.!?](?:["\')\]]+)?(?=\s|$)')
    sentence_ends = [m.end() for m in sentence_end_re.finditer(body)]

    cut_pos = None
    best_wc = -1
    for pos in sentence_ends:
        wc = bisect.bisect_right(word_ends, pos)
        if wc > max_words:
            break
        if wc >= min_words and wc > best_wc:
            best_wc = wc
            cut_pos = pos

    # If no sentence ending in range, use the latest sentence ending <= max_words.
    if cut_pos is None:
        for pos in sentence_ends:
            wc = bisect.bisect_right(word_ends, pos)
            if wc > max_words:
                break
            if wc > best_wc:
                best_wc = wc
                cut_pos = pos

    # Final fallback: exact max-word cut while preserving layout.
    # Also use this path if a sentence-boundary cut would violate the minimum.
    if cut_pos is None or best_wc < min_words:
        cut_pos = word_matches[max_words - 1].end()

    trimmed = body[:cut_pos].rstrip()
    if trimmed and trimmed[-1] not in ".!?":
        trimmed += "."

    if has_continue:
        return trimmed + "\n\nWill Continue to next part, say continue"
    if has_end:
        return trimmed + "\n\n(End of Answer)"
    return trimmed

def _is_abrupt_answer_ending(text: str) -> bool:
    """
    Detect likely truncation/abrupt stop (e.g., ending with 'bypass the').
    """
    body = (text or "")
    body = re.sub(r"(?im)^\s*Will Continue to next part, say continue\s*$", "", body).strip()
    body = re.sub(r"\(End of Answer\)\s*$", "", body, flags=re.IGNORECASE).strip()
    if not body:
        return False
    last_line = ""
    for ln in reversed(body.splitlines()):
        if ln.strip():
            last_line = ln.strip()
            break
    if not last_line:
        return False
    # Bare enumerators or bullets indicate likely truncation even if they end with punctuation.
    if re.match(r"(?i)^(?:\(?\d+\)?[.)]|[a-z][.)]|[ivxlcdm]+[.)]|[-*•])\s*$", last_line):
        return True
    # A trailing heading without content is likely abrupt.
    if re.match(
        r"(?i)^(?:part\s+[ivxlcdm]+\s*:|part\s+\d+\s*:|conclusion\s*:|conclusion and advice\s*:|advice to [^:]+:)\s*$",
        last_line,
    ):
        return True
    # Tiny sentence fragments at the end of an otherwise substantial answer are
    # usually truncation artefacts (for example "The perceived.").
    sentence_candidates = re.split(r"(?<=[.!?])\s+", body)
    tail_sentence = (sentence_candidates[-1] or "").strip() if sentence_candidates else last_line
    tail_words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", tail_sentence)
    if len(sentence_candidates) >= 2 and 0 < len(tail_words) <= 3:
        return True
    if re.search(r"(?i)\b(and|or|to|of|for|with|under|against|between|by|from|the|a|an)\s*[.!?](?:[\"')\]]+)?\s*$", body):
        return True
    if re.search(r"[.!?](?:[\"')\]]+)?\s*$", body):
        return False
    # If the tail ends on a connector/article, it's almost certainly cut.
    if re.search(r"(?i)\b(and|or|to|of|for|with|under|against|between|by|the|a|an)\s*$", body):
        return True
    # Very short trailing fragment without punctuation is also suspicious.
    if len(last_line.split()) <= 8:
        return True
    # Default: no terminal punctuation means likely abrupt cut.
    return True

# Custom CSS for legal styling with proper edge effects (NOT sticking to edges)
